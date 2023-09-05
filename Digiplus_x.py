from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from tkinter import scrolledtext as st
from tkinter import filedialog
from num2es import TextNumber
import sqlite3
import tkinter as tk
from docxtpl import DocxTemplate
from docx import Document 
import os
import pandas as pd
from datetime import date
from datetime import datetime
import calendar
import locale
import time
import re
from copy import deepcopy
from tkinter import PhotoImage



def inicio_sesion():
   
    global ventana2
    ventana2 = Toplevel(ventana1)
    ventana2.title("Inicio de sesión")
    ventana2.geometry("400x300")
    ventana2.iconbitmap("Logo.ico")
    ventana2.focus_set()
    ventana2.grab_set()
    ventana2.transient(master=ventana1)

    global User_verify
    global passUser_verify

    User_verify=StringVar()
    passUser_verify=StringVar()

    global nombre_entry
    global contraseña_entry
    
    Label(ventana2).pack()
    Label(ventana2, text="Usuario", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).pack()
    nombre_entry=Entry(ventana2, textvariable="User_verify", width=34, font=("calibri", 14))
    nombre_entry.focus()
    nombre_entry.pack(ipady=3)
    Label(ventana2).pack()
    

    Label(ventana2, text="Contraseña", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).pack()
    contraseña_entry=Entry(ventana2, textvariable="passUser_verify", show="*", width=34, font=("calibri", 14))
    contraseña_entry.pack(ipady=3)
    Label(ventana2).pack()

    ttk.Button(ventana2, text="Iniciar sesión", cursor="hand2", command=validacion_de_datos).pack()
    Label(ventana2).pack()

    ttk.Button(ventana2, text="Regresar", cursor="hand2", command=ventana2.destroy).pack()
    ventana2.wait_window(ventana2)
    


    

    

def registrarse():
    
    global ventana3
    ventana3 = Toplevel(ventana1)
    ventana3.title("Registrarse")
    ventana3.geometry("400x180")
    ventana3.iconbitmap("Logo.ico")
    ventana3.focus_set()
    ventana3.grab_set()
    ventana3.transient(master=ventana1)

    global administrador_entry
    global claveadm
    claveadm=StringVar()
    

    Label(ventana3).pack()
    Label(ventana3, text="Clave de administrador", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).pack()
    administrador_entry=Entry(ventana3, textvariable="claveadm", show="*", width=34, font=("calibri", 14))
    administrador_entry.focus()
    administrador_entry.pack(ipady=3)
    Label(ventana3).pack()

    ttk.Button(ventana3, text="Ingresar", cursor="hand2", command=clave_administrador).pack()
    Label(ventana3, text="").pack()

    ttk.Button(ventana3, text="Regresar", cursor="hand2", command=regresarclaveadmi).pack()
    ventana3.wait_window(ventana3)

def get_password():
    try:
        with open("pass.txt", "r") as f:
            return f.read()
    except FileNotFoundError:
        return "digiplus"

def set_password(new_password):
    with open("pass.txt", "w") as f:
        f.write(new_password)




def clave_administrador():
    if administrador_entry.get() == get_password():
        administrador_entry.delete(0, END)
        ventana3.destroy()
        global ventana4
        global tabla1
        ventana4 = Toplevel(ventana1)
        ventana4.title("Registrarse")
        ventana4.geometry("1000x700")
        ventana4.iconbitmap("Logo.ico")
        ventana4.focus_set()
        ventana4.grab_set()
        ventana4.transient(master=ventana1)

        
    

        barra_menus = tk.Menu()
        menu_archivo = tk.Menu(barra_menus, tearoff=False) 
        menu_archivo.add_command(label="Limpiar campos", command=limpiar_campos_tabla1)
        menu_archivo.add_command(label="Salir", command=salir_aplicacion)
        barra_menus.add_cascade(menu=menu_archivo, label="Inicio")
        ventana4.config(menu=barra_menus)


        global ID_nuevo
        global nombre_nuevo
        global apellidos_nuevo
        global nombre_usuario
        global contraseña_nuevo

        ID_nuevo=StringVar()
        nombre_nuevo=StringVar()
        apellidos_nuevo=StringVar()
        nombre_usuario=StringVar()
        contraseña_nuevo=StringVar()

        global nombreNuevo_entry
        global apellidosNuevo_entry
        global NombreUsuario_entry
        global ContraseñaNueva_entry
        global new_password_entry

        

        Label(ventana4).grid(row=0, column=0)
        Label(ventana4, text="Nombre", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).grid(row=2, column=1)
        nombreNuevo_entry=Entry(ventana4, textvariable=nombre_nuevo, width=34, font=("calibri", 14))
        nombreNuevo_entry.grid(row=3, column=1)
    

    
        Label(ventana4, text="Apellidos", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).grid(row=2, column=2)
        apellidosNuevo_entry=Entry(ventana4, textvariable=apellidos_nuevo, width=34, font=("calibri", 14))
        apellidosNuevo_entry.grid(row=3, column=2)
        
        
        Label(ventana4, text="Nombre de Usuario", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).grid(row=4, column=1)
        NombreUsuario_entry=Entry(ventana4, textvariable=nombre_usuario, width=34, font=("calibri", 14))
        NombreUsuario_entry.grid(row=5, column=1)
        

        
        Label(ventana4, text="Contraseña", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).grid(row=4, column=2)
        ContraseñaNueva_entry=Entry(ventana4, textvariable=contraseña_nuevo, show="*", width=34, font=("calibri", 14))
        ContraseñaNueva_entry.grid(row=5, column=2)

        Label(ventana4).grid(row=6, column=1)
        ttk.Button(ventana4, text="Registrar usuario", cursor="hand2", command=insertar_datos).place(relx=0.2, rely=0.9)
        ttk.Button(ventana4, text="Mostrar", cursor="hand2", command=mostrar_tabla1).place(relx=0.37, rely=0.9)
        ttk.Button(ventana4, text="Eliminar", cursor="hand2", command=borrar_tabla1).place(relx=0.5, rely=0.9)
        ttk.Button(ventana4, text="Regresar", cursor="hand2", command=ventana4.destroy).place(relx=0.7, rely=0.9)
        ttk.Button(ventana4, text="Cambiar contraseña", cursor="hand2", command=ventana_nueva_contraseña).place(relx=0.8, rely=0.9)

        tabla1=ttk.Treeview(ventana4, height=20,  columns=[f"#{n}" for n in range(1, 5)])
        tabla1.grid(row=8, columnspan=4)
        tabla1.heading("#0", text = "ID", anchor = CENTER)
        tabla1.heading("#1", text = "Nombre", anchor = CENTER)
        tabla1.heading("#2", text = "Apellido", anchor = CENTER)
        tabla1.heading("#3", text = "Nombre_usuario", anchor = CENTER)
        tabla1.heading("#4", text = "Contraseña", anchor = CENTER)

        tabla1.column("#0", anchor = CENTER)
        tabla1.column("#1", anchor = CENTER)
        tabla1.column("#2", anchor = CENTER)
        tabla1.column("#3", anchor = CENTER)
        tabla1.column("#4", anchor = CENTER)

        
        tabla1.bind("<ButtonRelease-1>", seleccionar_tabla1)
        

        ventana4.wait_window(ventana4)
            
    else:
        messagebox.showwarning(message="El nombre de usuario y contraseñas son incorrectas", title="Error")

def ventana_nueva_contraseña():
        global label
        top = tk.Toplevel(ventana4)
        top.iconbitmap("Logo.ico")
        top.focus_set()
        top.grab_set()
        top.transient(master=ventana4)
        label = tk.Label(top, text="Ingresa nueva contraseña:")
        label.pack()
        new_password_entry = tk.Entry(top, show="*")
        new_password_entry.pack()
        change_button = tk.Button(top, text="Cambiar contraseña", command=lambda: change_password(new_password_entry.get()))
        change_button.pack()

def change_password(new_password):
    set_password(new_password)
    label.config(text="Contraseña cambiada")


def regresarclaveadmi():
    administrador_entry.delete(0, END)
    ventana3.destroy()

    

def cerrar_sesión():
    ventana6.destroy()
    ventana1.deiconify()

def salir_aplicacion():
    valor=messagebox.askquestion(message= "Desea salir de la aplicación?")
    if valor=="yes":
        ventana1.destroy()
        
       
        

def limpiar_campos_tabla1():
    nombreNuevo_entry.delete(0, END)
    apellidosNuevo_entry.delete(0, END)
    NombreUsuario_entry.delete(0, END)
    ContraseñaNueva_entry.delete(0, END)





def seleccionar_tabla1(event):
    item= tabla1.identify('item', event.x, event.y)
    ID_nuevo.set(tabla1.item(item, "text"))
    nombre_nuevo.set(tabla1.item(item, "values")[0])
    apellidos_nuevo.set(tabla1.item(item, "values")[1])
    nombre_usuario.set(tabla1.item(item, "values")[2])
    contraseña_nuevo.set(tabla1.item(item, "values")[3])


def mostrar_tabla1():
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()
    registros=tabla1.get_children()
    for elemento in registros:
        tabla1.delete(elemento)

    
    mcursor.execute("SELECT * FROM UserNew")
    for row in mcursor:
        tabla1.insert("", 0, text=row[0], values=(row[1], row[2], row[3], row[4]))

def insertar_datos():
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()

          
    try:
        
        mcursor.execute("INSERT INTO UserNew  VALUES (NULL, ?, ?, ?, ?)", (nombreNuevo_entry.get(), apellidosNuevo_entry.get(), NombreUsuario_entry.get(), ContraseñaNueva_entry.get()))
        bd.commit()
        messagebox.showinfo(message="Registro exitoso", title="Aviso")

        
        nombreNuevo_entry.delete(0, END)
        apellidosNuevo_entry.delete(0, END)
        NombreUsuario_entry.delete(0, END)
        ContraseñaNueva_entry.delete(0, END)
        mostrar_tabla1()
    
    except:
        bd.rollback()
        messagebox.showinfo(message="No registrado", title="Aviso")

    bd.close()



    

def borrar_tabla1():
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()
    try:
        if messagebox.askyesno(message="Desea eliminar el registro?", title="Advertencia"):
            mcursor.execute("DELETE FROM UserNew WHERE ID="+ID_nuevo.get())
            bd.commit()
    except:
        messagebox.showinfo(message="Seleccione un registro")
        pass
    
    nombreNuevo_entry.delete(0, END)
    apellidosNuevo_entry.delete(0, END)
    NombreUsuario_entry.delete(0, END)
    ContraseñaNueva_entry.delete(0, END)
    mostrar_tabla1()



    

            

def validacion_de_datos():
    username=nombre_entry.get()
    pass1=contraseña_entry.get()
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()

    mcursor.execute("SELECT Contraseña FROM UserNew WHERE Usuario=? AND Contraseña=?" , (username, pass1))

    if mcursor.fetchall():
        escritura_datos(username)
       
    
    else:
        messagebox.showerror(title="Error", message="Usuario y/o contraseña incorrecta")
        nombre_entry.delete(0, END)
        contraseña_entry.delete(0, END)
        nombre_entry.focus()    
    
    
    




        
    
def escritura_datos(username):
    nombre_entry.delete(0, END)
    contraseña_entry.delete(0, END)
    ventana2.destroy()
    ventana1.withdraw()
    
    
    
    
    global tabla
    global ventana6
    ventana6=Toplevel(ventana1)
    ventana6.title("Sesion Iniciada")
    ventana6.geometry("1400x400")
    ventana6.iconbitmap("Logo.ico")
    ventana6.focus_set()
    





    global ID_registro
    global matriculaNuevo
    global cedulaNuevo
   
    global precioNuevo
    global radicado
    global nescrituras
    global valorLetras
    global ubicacion1
    global rural_urbano1
    
    global direccionNueva1
    global dia1
    global mes1
    global año1
    global añoLetras
    global Modo_adquirir1
    
   
    global notaria1
    global escritura1
    global escrLetras
    global notario1
    global notariod1 
    global departamento1
    global municipio1 
    global Notaria_municipio1
    global linderomatricula1
    global pazysalvomatricula1

    global ven1_cedulavendedor
    global ven1_primerApellido
    global ven1_segundoApellido
    global ven1_primerNombre
    global ven1_segundoNombre
    global ven1_sexo
    global ven1_fechadenacimiento
    global ven1_rh
    global ven1_domicilio
    global ven1_estadocivil 

    global cedulavendedor2
    global Ven2_primerApellido
    global Ven2_segundoApellido
    global Ven2_primerNombre
    global Ven2_segundoNombre
    global Ven2_sexo
    global Ven2_fechadenacimiento
    global Ven2_rh

    global Ven2_domicilio
    global Ven2_estadocivil


    global cedulavendedor3
    global Ven3_primerApellido
    global Ven3_segundoApellido
    global Ven3_primerNombre
    global Ven3_segundoNombre
    global Ven3_sexo
    global ven3_fechadenacimiento
    global Ven3_rh
    global Ven3_domicilio
    global Ven3_estadocivil


    global cedulavendedor4
    global Ven4_primerApellido
    global Ven4_segundoApellido
    global Ven4_primerNombre
    global Ven4_segundoNombre
    global Ven4_sexo
    global Ven4_fechanacimiento
    global Ven4_rh
    global Ven4_domicilio
    global Ven4_estadocivil

    global cedulavendedor5
    global Ven5_primerApellido
    global Ven5_segundoApellido
    global Ven5_primerNombre
    global Ven5_segundoNombre
    global Ven5_sexo
    global Ven5_fechanacimiento
    global Ven5_rh
    global Ven5_domicilio
    global Ven5_estadocivil

    global cedulavendedor6
    global Ven6_primerApellido
    global Ven6_segundoApellido
    global Ven6_primerNombre
    global Ven6_segundoNombre
    global Ven6_sexo
    global Ven6_fechanacimiento
    global Ven6_rh
    global Ven6_domicilio
    global Ven6_estadocivil



    global cedulacomprador1
    global con1_primer_apellido
    global con1_segundo_apellido
    global con1_primer_nombre
    global con1_segundo_nombre
    global con1_sexo
    global con1_fechanacimiento
    global con1_rh
    global con1_domicilio
    global con1_estadocivil 

    global cedulacomprador2
    global con2_primerApellido
    global con2_segundoApellido
    global con2_primerNombre
    global con2_segundoNombre
    global con2_sexo
    global con2_fechanacimiento
    global con2_rh
    global con2_domicilio
    global con2_estadocivil

    global cedulacomprador3
    global con3_primerApellido
    global con3_segundoApellido
    global con3_primerNombre
    global con3_segundoNombre
    global con3_sexo
    global con3_fechanacimiento
    global con3_rh
    global con3_domicilio
    global con3_estadocivil

    global cedulacomprador4
    global con4_primerApellido
    global con4_segundoApellido
    global con4_primerNombre
    global con4_segundoNombre
    global con4_sexo
    global con4_fechanacimiento
    global con4_rh
    global con4_domicilio
    global con4_estadocivil

    global cedulacomprador5
    global con5_primerApellido
    global con5_segundoApellido
    global con5_primerNombre
    global con5_segundoNombre
    global con5_sexo
    global con5_fechanacimiento
    global con5_rh
    global con5_domicilio
    global con5_estadocivil

    global cedulacomprador6
    global con6_primerApellido
    global con6_segundoApellido
    global con6_primerNombre
    global con6_segundoNombre
    global con6_sexo
    global con6_fechanacimiento
    global con6_rh
    global con6_domicilio
    global con6_estadocivil
    global boton_abrir
    global boton_notario
    global boton_lindero
    global boton_eliminar
    global boton_ventana_secundaria
   
    
    
    ID_registro=StringVar()
    matriculaNuevo=StringVar()
    cedulaNuevo=StringVar()
    precioNuevo= StringVar()
    valorLetras=StringVar()
    ubicacion1=StringVar()
    rural_urbano1=StringVar()
    radicado=StringVar()
    nescrituras=StringVar()
    variable_linderos=StringVar()
   

    direccionNueva1=StringVar()
    dia1=StringVar()
    mes1=StringVar()
    año1=StringVar()
    añoLetras=StringVar()
    Modo_adquirir1=StringVar()


    notaria1=StringVar()
    escritura1=StringVar()
    escrLetras=StringVar()
    notario1=StringVar()
    notariod1=StringVar()
    departamento1=StringVar()
    municipio1=StringVar() 
    Notaria_municipio1=StringVar()

    ven1_cedulavendedor=StringVar()
    ven1_primerApellido=StringVar()
    ven1_segundoApellido=StringVar()
    ven1_primerNombre=StringVar()
    ven1_segundoNombre=StringVar()
    ven1_sexo=StringVar()
    ven1_rh=StringVar()
    ven1_domicilio=StringVar()
    ven1_fechadenacimiento=StringVar()
    ven1_estadocivil=StringVar() 

    cedulavendedor2=StringVar()
    Ven2_primerApellido=StringVar()
    Ven2_segundoApellido=StringVar()
    Ven2_primerNombre=StringVar()
    Ven2_segundoNombre=StringVar()
    Ven2_sexo=StringVar()
    Ven2_fechadenacimiento=StringVar()
    Ven2_rh=StringVar()
    Ven2_domicilio=StringVar()
    Ven2_estadocivil=StringVar()


    cedulavendedor3=StringVar()
    Ven3_primerApellido=StringVar()
    Ven3_segundoApellido=StringVar()
    Ven3_primerNombre=StringVar()
    Ven3_segundoNombre=StringVar()
    Ven3_sexo=StringVar()
    ven3_fechadenacimiento=StringVar()
    Ven3_rh=StringVar()
    Ven3_domicilio=StringVar()
    Ven3_estadocivil=StringVar()


    cedulavendedor4=StringVar()
    Ven4_primerApellido=StringVar()
    Ven4_segundoApellido=StringVar()
    Ven4_primerNombre=StringVar()
    Ven4_segundoNombre=StringVar()
    Ven4_sexo=StringVar()
    Ven4_fechanacimiento=StringVar()
    Ven4_rh=StringVar()
    Ven4_domicilio=StringVar()
    Ven4_estadocivil=StringVar()

    cedulavendedor5=StringVar()
    Ven5_primerApellido=StringVar()
    Ven5_segundoApellido=StringVar()
    Ven5_primerNombre=StringVar()
    Ven5_segundoNombre=StringVar()
    Ven5_sexo=StringVar()
    Ven5_fechanacimiento=StringVar()
    Ven5_rh=StringVar()
    Ven5_domicilio=StringVar()
    Ven5_estadocivil=StringVar()

    cedulavendedor6=StringVar()
    Ven6_primerApellido=StringVar()
    Ven6_segundoApellido=StringVar()
    Ven6_primerNombre=StringVar()
    Ven6_segundoNombre=StringVar()
    Ven6_sexo=StringVar()
    Ven6_fechanacimiento=StringVar()
    Ven6_rh=StringVar()
    Ven6_domicilio=StringVar()
    Ven6_estadocivil=StringVar()



    cedulacomprador1=StringVar()
    con1_primer_apellido=StringVar()
    con1_segundo_apellido=StringVar()
    con1_primer_nombre=StringVar()
    con1_segundo_nombre=StringVar()
    con1_sexo=StringVar()
    con1_fechanacimiento=StringVar()
    con1_rh=StringVar()
    con1_domicilio=StringVar()
    con1_estadocivil=StringVar()

    cedulacomprador2=StringVar()
    con2_primerApellido=StringVar()
    con2_segundoApellido=StringVar()
    con2_primerNombre=StringVar()
    con2_segundoNombre=StringVar()
    con2_sexo=StringVar()
    con2_fechanacimiento=StringVar()
    con2_rh=StringVar()
    con2_domicilio=StringVar()
    con2_estadocivil=StringVar()

    cedulacomprador3=StringVar()
    con3_primerApellido=StringVar()
    con3_segundoApellido=StringVar()
    con3_primerNombre=StringVar()
    con3_segundoNombre=StringVar()
    con3_sexo=StringVar()
    con3_fechanacimiento=StringVar()
    con3_rh=StringVar()
    con3_domicilio=StringVar()
    con3_estadocivil=StringVar()

    cedulacomprador4=StringVar()
    con4_primerApellido=StringVar()
    con4_segundoApellido=StringVar()
    con4_primerNombre=StringVar()
    con4_segundoNombre=StringVar()
    con4_sexo=StringVar()
    con4_fechanacimiento=StringVar()
    con4_rh=StringVar()
    con4_domicilio=StringVar()
    con4_estadocivil=StringVar()

    cedulacomprador5=StringVar()
    con5_primerApellido=StringVar()
    con5_segundoApellido=StringVar()
    con5_primerNombre=StringVar()
    con5_segundoNombre=StringVar()
    con5_sexo=StringVar()
    con5_fechanacimiento=StringVar()
    con5_rh=StringVar()
    con5_domicilio=StringVar()
    con5_estadocivil=StringVar()

    cedulacomprador6=StringVar()
    con6_primerApellido=StringVar()
    con6_segundoApellido=StringVar()
    con6_primerNombre=StringVar()
    con6_segundoNombre=StringVar()
    con6_sexo=StringVar()
    con6_fechanacimiento=StringVar()
    con6_rh=StringVar()
    con6_domicilio=StringVar()
    con6_estadocivil=StringVar()

  

    def registrar_escritura():
        
        boton_abrir["state"]="disabled"
        boton_notario["state"]="disabled"
        boton_lindero["state"]="disabled"
        boton_ventana_secundaria["state"]="normal"
        global ventana5
        ventana5 = Toplevel(ventana6)
        ventana5.title("Registrar")
        ancho_ventana = ventana5.winfo_screenwidth()
        alto_ventana = ventana5.winfo_screenheight()
        ventana5.geometry(f"{ancho_ventana}x680")
        ventana5.iconbitmap("Logo.ico")
        ventana5.configure(bg="deep sky blue")
        ventana5.overrideredirect(True)

        
    
    
    
        global frame1
        global frame2
        global frame3
        global frame4
        global variable_linderos



        global ID_registro
        global matriculaNuevo
        global cedulaNuevo
    
        global precioNuevo
        global precioletras

        global valorLetras
        global ubicacion1
        global rural_urbano1
        global cajaUR1
        
        global direccionNueva1
        global dia1
        global mes1
        global año1
        global añoLetras
        global Modo_adquirir1
        
    
        global notaria1
        global notariod1
        global escritura1
        global escrLetras
        global notario1
        global departamento1
        global municipio1 
        global Notaria_municipio1
        global linderomatricula1
        global pazysalvomatricula1

        global ven1_cedulavendedor
        global ven1_primerApellido
        global ven1_segundoApellido
        global ven1_primerNombre
        global ven1_segundoNombre
        global ven1_sexo
        global ven1_fechadenacimiento
        global ven1_rh
        global ven1_domicilio
        global ven1_estadocivil 

        global cedulavendedor2
        global Ven2_primerApellido
        global Ven2_segundoApellido
        global Ven2_primerNombre
        global Ven2_segundoNombre
        global Ven2_sexo
        global Ven2_fechadenacimiento
        global Ven2_rh

        global Ven2_domicilio
        global Ven2_estadocivil


        global cedulavendedor3
        global Ven3_primerApellido
        global Ven3_segundoApellido
        global Ven3_primerNombre
        global Ven3_segundoNombre
        global Ven3_sexo
        global ven3_fechadenacimiento
        global Ven3_rh
        global Ven3_domicilio
        global Ven3_estadocivil


        global cedulavendedor4
        global Ven4_primerApellido
        global Ven4_segundoApellido
        global Ven4_primerNombre
        global Ven4_segundoNombre
        global Ven4_sexo
        global Ven4_fechanacimiento
        global Ven4_rh
        global Ven4_domicilio
        global Ven4_estadocivil

        global cedulavendedor5
        global Ven5_primerApellido
        global Ven5_segundoApellido
        global Ven5_primerNombre
        global Ven5_segundoNombre
        global Ven5_sexo
        global Ven5_fechanacimiento
        global Ven5_rh
        global Ven5_domicilio
        global Ven5_estadocivil

        global cedulavendedor6
        global Ven6_primerApellido
        global Ven6_segundoApellido
        global Ven6_primerNombre
        global Ven6_segundoNombre
        global Ven6_sexo
        global Ven6_fechanacimiento
        global Ven6_rh
        global Ven6_domicilio
        global Ven6_estadocivil



        global cedulacomprador1
        global con1_primer_apellido
        global con1_segundo_apellido
        global con1_primer_nombre
        global con1_segundo_nombre
        global con1_sexo
        global con1_fechanacimiento
        global con1_rh
        global con1_domicilio
        global con1_estadocivil 

        global cedulacomprador2
        global con2_primerApellido
        global con2_segundoApellido
        global con2_primerNombre
        global con2_segundoNombre
        global con2_sexo
        global con2_fechanacimiento
        global con2_rh
        global con2_domicilio
        global con2_estadocivil

        global cedulacomprador3
        global con3_primerApellido
        global con3_segundoApellido
        global con3_primerNombre
        global con3_segundoNombre
        global con3_sexo
        global con3_fechanacimiento
        global con3_rh
        global con3_domicilio
        global con3_estadocivil

        global cedulacomprador4
        global con4_primerApellido
        global con4_segundoApellido
        global con4_primerNombre
        global con4_segundoNombre
        global con4_sexo
        global con4_fechanacimiento
        global con4_rh
        global con4_domicilio
        global con4_estadocivil

        global cedulacomprador5
        global con5_primerApellido
        global con5_segundoApellido
        global con5_primerNombre
        global con5_segundoNombre
        global con5_sexo
        global con5_fechanacimiento
        global con5_rh
        global con5_domicilio
        global con5_estadocivil

        global cedulacomprador6
        global con6_primerApellido
        global con6_segundoApellido
        global con6_primerNombre
        global con6_segundoNombre
        global con6_sexo
        global con6_fechanacimiento
        global con6_rh
        global con6_domicilio
        global con6_estadocivil
        global diaactual
        global mesactual
        global añoactual
        global notariad1
        global notaria_actual1
        global notario_notaria1
        global municipio_not1
        global paginas
        global entry_chip1

        global boton_agregar_ven2
        global boton_quitar_ven2
        global boton_agregar_ven3
        global boton_quitar_ven3
        global boton_agregar_ven4
        global boton_quitar_ven4
        global boton_agregar_ven5
        global boton_quitar_ven5
        global boton_quitar_ven5
        global boton_quitar_ven6

        global buscarcedulav1
        global buscarcedulav2
        global buscarcedulav3
        global buscarcedulav4
        global buscarcedulav5
        global buscarcedulav6
        global buscarcedulac1
        global buscarcedulac2
        global buscarcedulac3
        global buscarcedulac4
        global buscarcedulac5
        global buscarcedulac6

        

        entry_chip1=StringVar()
        notariad1=StringVar()
        ID_registro=StringVar()
        matriculaNuevo=StringVar()
        cedulaNuevo=StringVar()
        precioNuevo= StringVar()
        precioletras=StringVar()
        valorLetras=StringVar()
        ubicacion1=StringVar()
        rural_urbano1=StringVar()
        cajaUR1 = StringVar()

        direccionNueva1=StringVar()
        dia1=StringVar()
        mes1=StringVar()
        año1=StringVar()
        añoLetras=StringVar()
        Modo_adquirir1=StringVar()


        notaria1=StringVar()
        notariod1=StringVar()
        notaria_actual1=StringVar()
        notario_notaria1=StringVar()
        municipio_not1=StringVar()
        escritura1=StringVar()
        escrLetras=StringVar()
        notario1=StringVar()
        departamento1=StringVar()
        municipio1=StringVar() 
        Notaria_municipio1=StringVar()
        linderomatricula1=StringVar()
        pazysalvomatricula1=StringVar()

        variable_linderos=tk.StringVar()

        ven1_cedulavendedor=StringVar()
        ven1_primerApellido=StringVar()
        ven1_segundoApellido=StringVar()
        ven1_primerNombre=StringVar()
        ven1_segundoNombre=StringVar()
        ven1_sexo=StringVar()
        ven1_rh=StringVar()
        ven1_domicilio=StringVar()
        ven1_fechadenacimiento=StringVar()
        ven1_estadocivil=StringVar() 

        cedulavendedor2=StringVar()
        Ven2_primerApellido=StringVar()
        Ven2_segundoApellido=StringVar()
        Ven2_primerNombre=StringVar()
        Ven2_segundoNombre=StringVar()
        Ven2_sexo=StringVar()
        Ven2_fechadenacimiento=StringVar()
        Ven2_rh=StringVar()
        Ven2_domicilio=StringVar()
        Ven2_estadocivil=StringVar()


        cedulavendedor3=StringVar()
        Ven3_primerApellido=StringVar()
        Ven3_segundoApellido=StringVar()
        Ven3_primerNombre=StringVar()
        Ven3_segundoNombre=StringVar()
        Ven3_sexo=StringVar()
        ven3_fechadenacimiento=StringVar()
        Ven3_rh=StringVar()
        Ven3_domicilio=StringVar()
        Ven3_estadocivil=StringVar()


        cedulavendedor4=StringVar()
        Ven4_primerApellido=StringVar()
        Ven4_segundoApellido=StringVar()
        Ven4_primerNombre=StringVar()
        Ven4_segundoNombre=StringVar()
        Ven4_sexo=StringVar()
        Ven4_fechanacimiento=StringVar()
        Ven4_rh=StringVar()
        Ven4_domicilio=StringVar()
        Ven4_estadocivil=StringVar()

        cedulavendedor5=StringVar()
        Ven5_primerApellido=StringVar()
        Ven5_segundoApellido=StringVar()
        Ven5_primerNombre=StringVar()
        Ven5_segundoNombre=StringVar()
        Ven5_sexo=StringVar()
        Ven5_fechanacimiento=StringVar()
        Ven5_rh=StringVar()
        Ven5_domicilio=StringVar()
        Ven5_estadocivil=StringVar()

        cedulavendedor6=StringVar()
        Ven6_primerApellido=StringVar()
        Ven6_segundoApellido=StringVar()
        Ven6_primerNombre=StringVar()
        Ven6_segundoNombre=StringVar()
        Ven6_sexo=StringVar()
        Ven6_fechanacimiento=StringVar()
        Ven6_rh=StringVar()
        Ven6_domicilio=StringVar()
        Ven6_estadocivil=StringVar()



        cedulacomprador1=StringVar()
        con1_primer_apellido=StringVar()
        con1_segundo_apellido=StringVar()
        con1_primer_nombre=StringVar()
        con1_segundo_nombre=StringVar()
        con1_sexo=StringVar()
        con1_fechanacimiento=StringVar()
        con1_rh=StringVar()
        con1_domicilio=StringVar()
        con1_estadocivil=StringVar()

        cedulacomprador2=StringVar()
        con2_primerApellido=StringVar()
        con2_segundoApellido=StringVar()
        con2_primerNombre=StringVar()
        con2_segundoNombre=StringVar()
        con2_sexo=StringVar()
        con2_fechanacimiento=StringVar()
        con2_rh=StringVar()
        con2_domicilio=StringVar()
        con2_estadocivil=StringVar()

        cedulacomprador3=StringVar()
        con3_primerApellido=StringVar()
        con3_segundoApellido=StringVar()
        con3_primerNombre=StringVar()
        con3_segundoNombre=StringVar()
        con3_sexo=StringVar()
        con3_fechanacimiento=StringVar()
        con3_rh=StringVar()
        con3_domicilio=StringVar()
        con3_estadocivil=StringVar()

        cedulacomprador4=StringVar()
        con4_primerApellido=StringVar()
        con4_segundoApellido=StringVar()
        con4_primerNombre=StringVar()
        con4_segundoNombre=StringVar()
        con4_sexo=StringVar()
        con4_fechanacimiento=StringVar()
        con4_rh=StringVar()
        con4_domicilio=StringVar()
        con4_estadocivil=StringVar()

        cedulacomprador5=StringVar()
        con5_primerApellido=StringVar()
        con5_segundoApellido=StringVar()
        con5_primerNombre=StringVar()
        con5_segundoNombre=StringVar()
        con5_sexo=StringVar()
        con5_fechanacimiento=StringVar()
        con5_rh=StringVar()
        con5_domicilio=StringVar()
        con5_estadocivil=StringVar()

        cedulacomprador6=StringVar()
        con6_primerApellido=StringVar()
        con6_segundoApellido=StringVar()
        con6_primerNombre=StringVar()
        con6_segundoNombre=StringVar()
        con6_sexo=StringVar()
        con6_fechanacimiento=StringVar()
        con6_rh=StringVar()
        con6_domicilio=StringVar()
        con6_estadocivil=StringVar()
        diaactual=StringVar()
        mesactual=StringVar()
        añoactual=StringVar()
        paginas=StringVar()

        buscarcedulav1=StringVar()
        buscarcedulav2=StringVar()
        buscarcedulav3=StringVar()
        buscarcedulav4=StringVar()
        buscarcedulav5=StringVar()
        buscarcedulav6=StringVar()
        buscarcedulac1=StringVar()
        buscarcedulac2=StringVar()
        buscarcedulac3=StringVar()
        buscarcedulac4=StringVar()
        buscarcedulac5=StringVar()
        buscarcedulac6=StringVar()

        global buscarcedulav1_combo
        global buscarcedulav2_combo
        global buscarcedulav3_combo
        global buscarcedulav4_combo
        global buscarcedulav5_combo
        global buscarcedulav6_combo
        global buscarcedulac1_combo
        global buscarcedulac2_combo
        global buscarcedulac3_combo
        global buscarcedulac4_combo
        global buscarcedulac5_combo
        global buscarcedulac6_combo
        global ec_entry
        global Ven2_ec_entry
        global Ven3_ec_entry
        global Ven4_ec_entry
        global Ven5_ec_entry
        global Ven6_ec_entry

        global ecc_entry
        global con2_ec_entry
        global con3_ec_entry
        global con4_ec_entry
        global con5_ec_entry
        global con6_ec_entry

        
        global matriculaNuevo_entry
        global cedulaNuevo_entry
        global precioNuevo_entry
        global cajaUR
        

        
        global direccionNueva_entry
        global dia_comb
        global mes_comb
        global año_comb
        global adquirir
        global ubicacion
        global rural_urbano
        global notaria
        global escritura
        global notario
        global notariod
        global notaria_actual
        global notario_notaria
        global municipio_not
        global departamento
        global municipio 
        global municipioNota
        global linderomatricula
        global pazysalvomatricula


       


        global cedulavendedor_entry
        global primerApellido1_entry
        global segundoApellido1_entry
        global primerNombre1_entry
        global segundoNombre1_entry
        global sexo1_entry
        global fecha_nacimiento1
        global rh1_entry
        global domicilio1_entry
        global estadocivil1_entry 

        global cedulavendedor2_entry
        global Ven2_primerApellido_entry
        global Ven2_segundoApellido_entry
        global Ven2_primerNombre_entry
        global Ven2_segundoNombre_entry
        global Ven2_sexo_entry
        global Ven2_fecha_nacimiento
        global Ven2_rh_entry
        global Ven2_domicilio_entry
        global Ven2_estadocivil_entry 


        global cedulavendedor3_entry
        global Ven3_primerApellido_entry
        global Ven3_segundoApellido_entry
        global Ven3_primerNombre_entry
        global Ven3_segundoNombre_entry
        global Ven3_sexo_entry
        global Ven3_fecha_nacimiento
        global Ven3_rh_entry
        global Ven3_domicilio_entry
        global Ven3_estadocivil_entry


        global cedulavendedor4_entry
        global Ven4_primerApellido_entry
        global Ven4_segundoApellido_entry
        global Ven4_primerNombre_entry
        global Ven4_segundoNombre_entry
        global Ven4_sexo_entry
        global Ven4_fecha_nacimiento
        global Ven4_rh_entry
        global Ven4_domicilio_entry
        global Ven4_estadocivil_entry

        global cedulavendedor5_entry
        global Ven5_primerApellido_entry
        global Ven5_segundoApellido_entry
        global Ven5_primerNombre_entry
        global Ven5_segundoNombre_entry
        global Ven5_sexo_entry
        global Ven5_fecha_nacimiento
        global Ven5_rh_entry
        global Ven5_domicilio_entry
        global Ven5_estadocivil_entry

        global cedulavendedor6_entry
        global Ven6_primerApellido_entry
        global Ven6_segundoApellido_entry
        global Ven6_primerNombre_entry
        global Ven6_segundoNombre_entry
        global Ven6_sexo_entry
        global Ven6_fecha_nacimiento
        global Ven6_rh_entry
        global Ven6_domicilio_entry
        global Ven6_estadocivil_entry



        global cedulacomprador_entry
        global primer_apellido2_entry
        global segundo_apellido2_entry
        global primer_nombre2_entry
        global segundo_nombre2_entry
        global sexo2_entry
        global fecha_nacimiento2
        global rh2_entry
        global domicilio2_entry
        global estadocivil2_entry 

        global cedulacomprador2_entry
        global con2_primerApellido_entry
        global con2_segundoApellido_entry
        global con2_primerNombre_entry
        global con2_segundoNombre_entry
        global con2_sexo_entry
        global con2_fecha_nacimiento
        global con2_rh_entry
        global con2_domicilio_entry
        global con2_estadocivil_entry

        global cedulacomprador3_entry
        global con3_primerApellido_entry
        global con3_segundoApellido_entry
        global con3_primerNombre_entry
        global con3_segundoNombre_entry
        global con3_sexo_entry
        global con3_fecha_nacimiento
        global con3_rh_entry
        global con3_domicilio_entry
        global con3_estadocivil_entry

        global cedulacomprador4_entry
        global con4_primerApellido_entry
        global con4_segundoApellido_entry
        global con4_primerNombre_entry
        global con4_segundoNombre_entry
        global con4_sexo_entry
        global con4_fecha_nacimiento
        global con4_rh_entry
        global con4_domicilio_entry
        global con4_estadocivil_entry

        global cedulacomprador5_entry
        global con5_primerApellido_entry
        global con5_segundoApellido_entry
        global con5_primerNombre_entry
        global con5_segundoNombre_entry
        global con5_sexo_entry
        global con5_fecha_nacimiento
        global con5_rh_entry
        global con5_domicilio_entry
        global con5_estadocivil_entry

        global cedulacomprador6_entry
        global con6_primerApellido_entry
        global con6_segundoApellido_entry
        global con6_primerNombre_entry
        global con6_segundoNombre_entry
        global con6_sexo_entry
        global con6_fecha_nacimiento
        global con6_rh_entry
        global con6_domicilio_entry
        global con6_estadocivil_entry

        global boton_agregar_ven2
        global boton_agregar_ven3
        global boton_agregar_ven4
        global boton_agregar_ven5
        global boton_agregar_ven6
        global boton_quitar_ven2
        global boton_quitar_ven3
        global boton_quitar_ven4
        global boton_quitar_ven5
        global boton_quitar_ven6


        global boton_agregar_con2
        global boton_agregar_con3
        global boton_agregar_con4
        global boton_agregar_con5
        global boton_agregar_con6
        global boton_quitar_con2
        global boton_quitar_con3
        global boton_quitar_con4
        global boton_quitar_con5
        global boton_quitar_con6
        global boton_insertar_datos
        global boton_actualizar_datos
        global cerrar_ventana5
        global entry_fecha
        global entry_fecha2
        global entry_fecha3
        global entry_nescrituras
        global boton_volver
        global entry_nradicado
        global entry_paginas
        
        global notaria_actual_entry
        global entry_chip
        global label_chip
        global button1
        global button2
        global button3
        global button4
        global button5
        global button6

        global button7
        global button8
        global button9
        global button10
        global button11
        global button12
        
          
            
           

    

        
        fecha_actual = datetime.now().strftime('%d-%m-%Y')
        fecha_actual1 = datetime.now().strftime('%d')
        fecha_actual2 = datetime.now().strftime('%m')
        fecha_actual3 = datetime.now().strftime('%Y')
        numeroescr= 0


        
        conn = sqlite3.connect('login1.db')
        cursor = conn.cursor()

        # Consulta para obtener los valores de la segunda columna
        cursor.execute("SELECT Matricula_lindero FROM linderos2")
        options1 = [row[0] for row in cursor.fetchall()]


    
        frame1=LabelFrame(ventana5, text="Datos de Innmueble", bg="deep sky blue", fg="white", width=140, height=1, font=("Arial Black", 14))
        frame1.place(x=20, y=2)                            
        Label(frame1, text="Matricula:", bg="deep sky blue",  font=("Arial Black", 11)).grid(row=1, column=0, sticky=E)
        matriculaNuevo_entry=ttk.Combobox(frame1, width=20, textvariable=matriculaNuevo, values=options1)
        matriculaNuevo_entry.focus()
        matriculaNuevo_entry.grid(row=1, column=1)

        labellinderomatricula=Label(ventana5, text="Lindero", bg="deep sky blue", font=("Arial Black", 11))
        labellinderomatricula.place(relx=0.22, rely=0.86)
        linderomatricula = tk.Entry(ventana5, textvariable=linderomatricula1, width=20)
        linderomatricula.place(relx=0.22, rely=0.9)

        labelpazsalvomatricula=Label(ventana5, text="Paz y salvo", bg="deep sky blue", font=("Arial Black", 11))
        labelpazsalvomatricula.place(relx=0.32, rely=0.86)
        pazysalvomatricula = tk.Entry(ventana5, textvariable=pazysalvomatricula1, width=20)
        pazysalvomatricula.place(relx=0.32, rely=0.9)

        def update_entry1(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valor = matriculaNuevo_entry.get()
            
            cursor.execute("SELECT lindero, paz_salvo FROM linderos2 WHERE Matricula_lindero=?", (seleccionar_valor,))
            result = cursor.fetchone()
            if result:
                # Actualizar el valor del Entry para la columna lindero
                linderomatricula.delete(0, END)
                linderomatricula.insert(0, result[0])
                # Actualizar el valor del Entry para la columna paz_salvo
                pazysalvomatricula.delete(0, END)
                pazysalvomatricula.insert(0, result[1])
        
        matriculaNuevo_entry.bind("<<ComboboxSelected>>", update_entry1)


        Label(frame1, text=" ", bg="deep sky blue").grid(row=1, column=3)


        Label(frame1, text="Cedula Catastral:", bg="deep sky blue", font=("Arial Black", 11)).grid(row=2, column=0)
        cedulaNuevo_entry=Entry(frame1, textvariable=cedulaNuevo)
        cedulaNuevo_entry.grid(row=2, column=1, sticky=W)



        Label(frame1, text="Precio:", bg="deep sky blue", font=("Arial Black", 11)).grid(row=3, column=0, sticky=E)
        precioNuevo_entry=Entry(frame1, textvariable=precioNuevo)
        precioNuevo_entry.grid(row=3, column=1, sticky=W)

    

        Label(frame1,text="Ubicación:", bg="deep sky blue", font=("Arial Black", 11)).grid(row=4, column=0, sticky=E)
        ubicacion=ttk.Combobox(frame1, width=10, textvariable=ubicacion1, state="roadonly", values=["","La ciudad", "El municipio"])
        ubicacion.grid(row=4, column=1, sticky=W)
        ubicacion.current(0)

        def seleccion_urbano_rural(event):
            seleccionar_valor = rural_urbano.get()
            tomar_valor = valorRU[seleccionar_valor]
            cajaUR1.set(tomar_valor)

        Label(frame1,text="Rural o Urbano:", bg="deep sky blue", font=("Arial Black", 11)).grid(row=5, column=0, sticky=E)
        valorRU = {"":"", 'Rural': 'en la vereda', 'Urbano': ' en el casco urbano'}
        rural_urbano=ttk.Combobox(frame1, width=10, textvariable=rural_urbano1, state="roadonly", values=list(valorRU.keys()))
        rural_urbano.bind("<<ComboboxSelected>>", seleccion_urbano_rural)
        rural_urbano.grid(row=5, column=1, sticky=W)
        rural_urbano.current(0)

        
        cajaUR = Entry(frame1, textvariable=cajaUR1)
        cajaUR.grid(row=6, column=1, sticky=W)

        opciones = {

                "":(),
                "Amazonas": ("Leticia", "El Encanto", "La Chorrera", "La Pedrera","La Victoria",  "Miriti-Parana", 
                "Puerto alegria", "Puerto Arica", "Puerto Nariño", "Puero santander","Tarapaca"), 

                "Antioquia": ("Medellin", "Abejorral", "Abriaquí", "Alejandría", "Amagá", "Amalfi", "Andes", "Angelópolis", "Angostura", "Anorí",
                "Anza", "Apartadó", "Arboletes", "Argelia", "Armenia", "Barbosa", "Bello", "Belmira", "Betania", "Betulia",
                "Briceño", "Buriticá", "Cáceres", "Caicedo", "Caldas", "Campamento", "Cañasgordas", "Caracolí", "Caramanta", "Carepa",
                "Carolina", "Caucasia", "Chigorodó", "Cisneros", "Ciudad Bolívar", "Cocorná", "Concepción", "Concordia", "	Copacabana", "Dabeiba",
                "Don Matías", "Ebéjico", "El Bagre", "El Carmen de Viboral", "El Peñol", "El Retiro", "El Santuario", "Entrerrios", "Envigado", "Fredonia",
                "Frontino", "Giraldo", "Girardota", "Gómez Plata", "Granada", "Guadalupe", "Guarne", "Guatapé", "Heliconia", "Hispania",
                "Itagui", "Ituango", "Jardín", "Jericó", "La Ceja", "La Estrella", "La Pintada", "La Unión", "Liborina", "Maceo",
                "Marinilla", "Montebello", "Murindó", "Mutatá", "Nariño", "Nechí", "Necoclí", "Olaya", "Peque", "Pueblorrico",
                "Puerto Berrío", "Puerto Nare", "Puerto Triunfo", "Remedios", "Rionegro", "Sabanalarga", "Sabaneta", "Salgar", "San Andrés de Cuerquía", "San Carlos",
                "San Francisco", "San Jerónimo", "San José de La Montaña", "San Juan de Urabá", "San Luis", "San Pedro", "San Pedro de Uraba", "San Rafael", "San Roque", "San Vicente",
                "Santa Bárbara", "	Santa Rosa de Osos", "Santafé de Antioquia", "Santo Domingo", "Segovia", "Sonson", "Sopetrán", "9	Támesis", "Tarazá", "Tarso",
                "Titiribí", "Toledo", "Turbo", "Uramita", "Urrao", "Valdivia", "Valparaíso", "Vegachí", "Venecia", "Vigía del Fuerte",
                "Yalí", "	Yarumal", "Yolombó", "Yondó", "Zaragoza"),

                "Arauca":("Arauca", "Arauquita", "Cravo Norte", "Fortul", "Puerto Rodón", "Saravena", "Tame"),

                "Atlantico":("Barranquilla", "Baranoa", "Campo de la Cruz", "Candelaria", "Galapa", "Juan de Acosta", "Luruaco", "Malambo", "Manatí", "Palmar de Varela",
                "Piojó", "Polonuevo", "Ponedera", "Puerto Colombia", "Repelón", "Sabanagrande", "Sabanalarga", "Santa Lucía", "Santo Tomás", "Soledad",
                "Suán", "Tubará", "Usiacurí"),

                
                "Bolivar": ("Cartagena", "Achí", "Altos del Rosario", "Arenal",
                "Arjona", "Arroyohondo", "Barranco de Loba", "Brazuelo de Papayal", "Calamar", "Cantagallo", "El Carmen de Bolívar", "El Carmen de Bolívar", "Cicuco", "Clemencia",
                "Córdoba", "El Guamo", "El Peñón", "Hatillo de Loba", "Magangué", "Mahates", "Margarita", "María La Baja", "Montecristo", "Morales",
                "Norosí", "Pinillos", "Regidor", "Río Viejo", "San Cristóbal", "San Estanislao", "San Fernando", "San Jacinto", "San Jacinto del Cauca", "San Juan Nepomuceno",
                "San Martín de Loba", "San Pablo", "Santa Catalina", "Santa Cruz de Mompox", "Santa Rosa", "Santa Rosa del Sur", "Simití", "Soplaviento", "Talaigua Nuevo", "Tiquisio",
                "Turbaco", "Turbaná", "Villanueva", "Zambrano"),

                "Boyaca":("Tunja", "Almeida", "Aquitania", "Arcabuco", "Belén", "Berbeo", "Betéitiva", "Boavita", "Boyacá", "Briceño",
                "Buenavista", "Busbanzá", "Caldas", "Campohermoso", "Cerinza", "Chinavita", "Chiquinquirá", "Chíquiza", "Chiscas", "Chita",
                "Chitaraque", "Chivatá", "Chivor", "Ciénega", "Cómbita", "Coper", "Corrales", "Covarachía", "Cubará", "Cucaita",
                "Cuítiva", "Duitama", "El Cocuy", "El Espino", "Firavitoba", "Floresta", "Gachantivá", "Gameza", "Garagoa", "Guacamayas",
                "Guateque", "	Guayatá", "Güicán", "Iza", "Jenesano", "Jericó", "La Capilla", "La Uvita", "La Victoria", "Labranzagrande",
                "Macanal", "Maripí", "Miraflores", "Mongua", "Monguí", "Moniquirá", "Motavita", "Muzo", "Nobsa", "Nuevo Colón",
                "Oicatá", "Otanche", "Pachavita", "Páez", "Paipa", "Pajarito", "Panqueba", "Pauna", "Paya","Paz de Río", "Pesca", "Pisba", 
                "Puerto Boyacá", "Quípama", "Ramiriquí", "	Ráquira", "Rondón", "Saboyá", "Sáchica", "Samacá", "San Eduardo", "San José de Pare", 
                "San Luis de Gaceno", "San Mateo", "San Miguel de Sema", "San Pablo de Borbur", "Santa María", "Santa Rosa de Viterbo", "Santa Sofía",
                "Santana", "Sativanorte", "Sativasur", "Siachoque", "Soatá", "Socha", "Socotá", "Sogamoso", "Somondoco", "Sora","Soracá", "Sotaquirá", 
                "Susacón", "Sutamarchán", "Sutatenza", "Tasco", "Tenza", "Tibaná", "Tibasosa", "Tinjacá","Tipacoque", "Toca", "Togüí", "Tópaga", "Tota",
                "Tununguá", "Turmequé", "Tuta", "Tutazá", "Umbita","Ventaquemada", "Villa de Leyva", "Viracachá", "Zetaquira"),

                "Caldas": ("Manizales", "Aguadas", "Anserma", "Aranzazu", "Belalcazar", "Chinchiná", "Filadelfia", "La Dorada", "La Merced", "Manzanares",
                "Marmato", "Marquetalia", "Marulanda", "Neira", "Norcasia", "Pacora", "Palestina", "Pensilvania", "Riosucio", "Risaralda",
                "Salamina", "Samana", "San Jose", "Supía", "Victoria", "Villamaría", "Viterbo"),

                "Caqueta": ("Florencia", "Albania", " Cartagena del Chairá", "Curillo", "El Doncello", "El Paujil", "La Montañita", "Morelia", 
                "Puerto Milán", "Puerto Rico", "San José del Fragua", "San Vicente del Caguán", "Solano", "Solita", "Valparaíso"),

                "Casanare": ("Yopal", "Aguazul", "Chámeza", "Hato Corozal", "La Salina", "Maní", "Monterrey", "Nunchía", "Orocué", "Paz de Ariporo", 
                "Pore", "Recetor", "Sabanalarga", "Sácama", "San Luis de Palenque", "Támara", "Tauramena", "Trinidad", "Villanueva"),

                "Cauca": ("Popayán", "Almaguer", "Argelia", "Balboa", "Bolívar", "Buenos Aires", "Cajibio", "Caldono", "Caloto", "Corinto", "El Tambo", 
                "Florencia", "Guapi", "Inza", "Jambaló", "La Sierra", "La Vega", "López", "Mercaderes", "Miranda", "Morales", "Padilla", 
                "Páez", "Patia (El Bordo)", "Piamonte", "Piendamo",  "Puerto Tejada", "Purace", "Rosas", "San Sebastián", "Santa Rosa",
                    "Santander de Quilichao", "Silvia", "Sotara", "Suárez", "Sucre", "Timbío", "Timbiquí", "Toribio", "Totoro", "Villa Rica"),

                "Cesar": ("Valledupar", "Aguachica", "Codazzi", "Astrea", "Becerril", "Bosconia", "Chimichagua", "Curumaní",
                "El Copey", "El Paso", "Gamarra", "González", "La Gloria", "La Jagua Ibirico", "Manaure Balcón Del Cesar", "Pailitas", "Pelaya", 
                "Pueblo Bello","Río De Oro", "Robles (La Paz)", "San Alberto", "San Diego", "San Martín", "Tamalameque"),

                "Choco": ("Quibdó", " Acandi", "Alto Baudo (pie de pato)","Atrato", "Bagado", "Bahia Solano (mutis)", "Bajo Baudo (pizarro)", 
                "Bojaya (bellavista)", "Canton de San Pablo", "Carmen del Darien", "Certegui", "Condoto", "El Carmen", "Istmina", "Jurado", "Litoral del san juan",
                "Lloro","Medio Atrato", "Medio Baudo (boca de pepe)", "Medio San Juan", "Novita", "Nuqui", "Rio iro", "Rio Quito", "Riosucio", 
                "San Jose del Palmar","Sipi", "Tado", "Unguia", "Unión Panamericana"),

                "Córdoba": ("Montería", "Ayapel", "Buenavista", "Canalete", "Cereté", "Chima", "Chinú", "Cienaga De Oro", "Cotorra", "La Apartada", "Lorica",
                "Los Córdobas", "Momil", "Montelíbano",  "Moñitos", "Planeta Rica", "Pueblo Nuevo", "Puerto Escondido", "Puerto Libertador", 
                "Purísima","Sahagún", "San Andrés de Sotavento", "San Antero", "San Bernardo del Viento", "San Carlos", "San Pelayo", "Tierralta", 
                "Valencia"),

                "Cundinamarca": ("Bogotá_D.C","Agua de Dios", "Albán", "Anapoima", "Anolaima", "Apulo", "Arbeláez", "Beltrán", "Bituima", "Bojacá", "Cabrera",
                "Cachipay", "Cajicá", "Caparrapí", "Cáqueza", "Carmen de Carupa", "Chaguaní", "Chía", "Chipaque", "Choachí", 
                "Chocontá","Cogua", "Cota", "Cucunubá", "El Colegio", "El Peñón", "El Rosal", "Facatativá", "Fómeque", "Fosca", "Funza",
                "Fúquene", "Fusagasugá", "Gachalá", "Gachancipá", "Gachetá", "Gama", "Girardot", "Granada", "Guachetá", "Guaduas", "Guasca",
                "Guataquí","Guatavita", "Guayabal de Síquima", "Guayabetal", "Gutiérrez", "Jerusalén", "Junín", "La Calera", "La Mesa", "La Palma",
                "La Peña", "La Vega", "Lenguazaque", "Machetá", "Madrid", "Manta", "Medina", "Mosquera", "Nariño", "Nemocón", "Nilo", "Nimaima",
                "Nocaima", "Pacho", "Paime", "Pandi", "Paratebueno", "Pasca", "Puerto Salgar", "Pulí", "Quebradanegra", "Quetame", "Quipile",
                "Ricaurte","San Antonio del Tequendama", "San Bernardo", "San Cayetano", "San Francisco", "San Juan de Rioseco", "Sasaima", 
                "Sesquilé", "Sibaté", "Silvania", "Simijaca", "Soacha", "Sopó", "Subachoque", "Suesca", "Supatá", "Susa", "Sutatausa", "Tabio",
                "Tausa", "Tena", "Tenjo", "Tibacuy", "Tibirita", "Tocaima", "Tocancipá", "Topaipí", "Ubalá", "Ubaque", "Ubaté", "Une", "Útica",
                "Venecia", "Vergara", "Vianí", "Villagómez", "Villapinzón", "Villeta", "Viotá", "Yacopí", "Zipacón", "Zipaquirá"),

                "Guainía": ("Inírida", "Barrancominas", "Cacahual", "La Guadalupe", "Mapiripana", "Morichal Nuevo", "Pana Pana", 
                "Puerto Colombia", "San Felipe"),

                "Guaviare": ("San Jose del Guaviare", "Calamar", "El Retorno", "Miraflorez"),

                "Huila": ("Neiva", "Acevedo", "Aipe", "Algeciras", "Altamira", "Baraya", "Campoalegre", "Colombia", "Elías", "El Agrado","Garzón",
                "Gigante", "Guadalupe", "Hobo", "Íquira", "Isnos", "La Argentina", "La Plata", "Nátaga", "Oporapa", "Paicol", "Palermo", "Palestina",
                "Pital", "Pitalito", "Rivera", "Saladoblanco", "Santa María", "San Agustín", "Suaza", "Tarqui", "Tello", "Teruel", "Tesalia", "Timaná",
                "Villavieja", "Yaguará"), 

                "La Guajira": ("Riohacha","Albania", "Barrancas", "Dibulla", "Distracción", "El Molino", "Fonseca", "Hatonuevo", "La Jagua del Pilar", "Maicao", 
                "Manaure", "San Juan del Cesar", "Uribia", "Urumita", "Villanueva"),

                "Magdalena": ("Santa Martha", "Algarrobo", "Aracataca", "Ariguaní", "Cerro de San Antonio", "Chibolo", "Ciénaga", "Concordia", "El Banco",
                "El Piñon","El Retén", "Fundación", "Guamal", "Nueva Granada", "Pedraza", "Pijino del Carmen", "Pivijai", "Plato", "Pueblo Viejo", 
                "Remolino", "Sabanas de San Ángel", "Salamina", "San Sebastián de Buenavista", "Santa Ana", "Santa Bárbara de Pinto", "San Zenón", 
                "Sitionuevo", "Tenerife", "Zapayán", "Zona Bananera"),

                "Meta": ("Villavicencio", "Acacías", "Barranca de Upía", "Cabuyaro", "Castilla La Nueva", "	Cubarral", "Cumaral", "El Calvario", "El Castillo",
                "El Dorado","Fuente de Oro", "Granada", "Guamal", "La Macarena", "Lejanías", "Mapiripán", "Mesetas", "Puerto Concordia", "Puerto Gaitán", 
                "Puerto Lleras","Puerto López", "Puerto Rico", "Restrepo", "San Carlos de Guaroa", "San Juan de Arama", "San Juanito", "San Martín", "Uribe", "Vista Hermosa"),

                "Nariño": ("Pasto", "Alban", "Aldaña", "Ancuya", "Arboleda", "Barbacoas", "Belen", "Buesaco", "Chachagui", "Colon(genova)","Consaca", 
                "Contadero", "Cordoba", "Cuaspud", "Cumbal", "Cumbitara", "El Charco", "El Peñol", "El Rosario", "El tablón", "El Tambo", "Funes", "Guachucal",
                "Guaitarilla", "Gualmatan", "Iles", "Imues", "Ipiales", "La Cruz", "La florida", "La llanada","La Tola", "La Unión", "Leiva", "Leiva", "Linares",
                "Los Andes", "Magui", "Mallama", "Mosquera", "Nariño", "Olaya Herrera", "Ospina", "Pizarro", "Policarpa", "Potosi", "Providencia", "Puerres", 
                "Pupiales", "Ricaurte", "Roberto Payan", "Samaniego", "San Bernardo", "San Lorenzo", "San Pablo", "Nariño", "San Pedro de Cartago", "Ospina", 
                "Sandona", "Santa Barbara", "Santacruz", "Sapuyes", "Taminango", "Tangua", "Tumaco", "Tuquerres", "Yacuanquer"),

                "Norte de Santander": ("Cucuta", "Ábrego", "Arboledas", "Bochalema", "Bochalema", "Cáchira", "Cácota", "Chinácota", "Chitagá", "",
                "Convención", "Cucutilla", "Durania", "El Carmen", "El Tarra", "El Zulia", "Gramalote", "Hacarí", "Herrán", "La Esperanza", "La Playa de Belén",
                "Labateca", "Los Patios", "	Lourdes", "Mutiscua", "Ocaña", "Pamplona", "Pamplonita", "Puerto Santander", "Ragonvalia", "Salazar de Las Palmas",
                "San Calixto", "San Cayetano", "Santiago", "Santo Domingo de Silos", "Sardinata", "Teorama", "Tibú", "Toledo", "Villa Caro", "Villa del Rosario"),

                "Putumayo": ("Mocoa", "Colón", "Orito", "Puerto Asís", "Puerto Caicedo", "Puerto Guzmán", "Puerto Leguízamo", "San Francisco", "San Miguel", 
                "Santiago", "Sibundoy", "Valle del Guamuez", "Villagarzón"),

                "Quindio": ("Armenia", "Buenavista", "Calarcá", "Circasia", "Córdoba", "Filandia", "Génova", "La Tebaida", "Montenegro", "Pijao", "Quimbaya","Salento"),

                "Risaralda": ("Pereira", "Apía", "Balboa", "Belén de Umbría", "Dosquebradas", "Guática", "La Celia", "La Virginia", "Marsella", "Mistrató",
                "Pueblo Rico", "Quinchía", "Santa Rosa de Cabal", "Santuario"),

                "San Andres": ("San_Andres"),

                "Santander": ("Bucaramanga", "Aguada", "Albania", "Aratoca", "Barbosa", "Barichara", "Barrancabermeja", "Betulia", "Bolívar", "Cabrera",
                " California", "Capitanejo", "Carcasí", "Cepitá", "Cerrito", "Charalá", "Charta", "Chima", "Chipatá", "Cimitarra", "Concepción", "Confines",
                "Contratación", "Coromoro", "Curití", "El Carmen de Chucurí", "El Guacamayo", "El Peñón", "El Playón", "Encino", "Enciso","Florián", "Floridablanca",
                "Galán", "Gámbita", "Girón", "Guaca", "Guadalupe", "Guapotá", "Guavatá", "Güepsa", "Hato","Jesús María", "Jordán", "La Belleza", "La Paz", "Landázuri",
                "Lebrija", "Los Santos", "Macaravita", "Málaga", "Matanza", "Mogotes", "Molagavita", " Ocamonte", "Oiba", "Onzaga", "Palmar", "Palmas del Socorro",
                "Páramo", "Piedecuesta", "Pinchote", "Puente Nacional", "Puerto Parra", "Puerto Wilches", "Rionegro", "Sabana de Torres", "San Andrés", "San Benito", 
                "San Gil", "San Joaquín", "San José de Miranda", "San Miguel", "San Vicente de Chucurí", "Santa Bárbara", "Santa Helena del Opón", "Simacota",
                "Socorro", "Suaita", "Sucre", "Suratá", "Tona", "Valle de San José", "Vélez", "Vetas", "Villanueva", "Zapatoca"),

                "Sucre": ("Sincelejo", "Buenavista", "Caimito", "Chalán", "Colosó", "Corozal", "Coveñas", "El Roble", "Galeras", "Guaranda", "La Unión", 
                "Los Palmitos", "Majagual", "Morroa", "Ovejas", "Palmito", "Sampués", "San Benito Abad", "San Juan de Betulia", "San Marcos", "San Onofre", 
                "San Pedro", "Santiago de Tolú", "Sincé", "Sucre", "Tolúviejo"),

                "Tolima":("Ibagué", "Alpujarra", "Alvarado", "Ambalema", "Anzoátegui", "Armero guayabal", "Ataco", "cunday", "Cajamarca", "Carmen de Apicalá", "Casabianca",
                "Chaparral", "Coello", "Coyaima", "Dolores", "Espinal", "Falan", "Flandes", "Fresno ", "Guamo", "Herveo", "Honda", "Icononzo", "Lérida ", "Líbano",
                "Mariquita", "Melgar", "Murillo", "Natagaima", "Ortega", "Palocabildo", "Piedras", "Planadas", "Prado", "Purificación", "Rioblanco", "Roncesvalles",
                "Rovira", "Saldaña", "San Antonio", "San Luis", "Santa Isabel", "Suárez", "Valle de San Juan", "Venadillo", "Villahermosa", "Villarrica"),

                "Valle del Cauca": ("Cali", "Alcalá", "Andalucía", "Ansermanuevo", "Argelia", "Bolívar", "Buenaventura", "Buga", "Bugalagrande", "Caicedonia",
                "Calima - El Darién", "Candelaria", "Cartago", "Dagua", "El Águila", "El Cairo", "El Cerrito", "El Dovio", "Florida", "Ginebra", "Guacarí", "Jamundí",
                "La Cumbre", "La Unión", "La Victoria", "Obando", "Palmira", "Pradera", "Restrepo", "Riofrío", "Roldanillo", "San Pedro", "Sevilla", "Toro", "Trujillo", 
                "Tuluá", "Ulloa", "Versalles", "Vijes", "Yotoco", "Yumbo", "Zarzal"),

                "Vaupés": ("Mitú", "Caruru", "Pacoa", "Taraira", "Papunaua", "Yavaraté"),

                "Vichada": ("Puerto Carreño", "Cumaribo", "La Primavera", "Santa Rosalía"),
                }
            
        def lugar_seleccion(event):

                
            municipio.set("")
            municipio.config(values=opciones[departamento.get()])


        departamento = ttk.Combobox(frame1, width="15", textvariable=departamento1, state="readonly", values=tuple(opciones.keys()))
        Label(frame1, text="Departamento: ", bg="deep sky blue", font=("Arial Black", 11)).grid(row=4, column=2, sticky=E)
        departamento.grid(row=4, column=3, sticky=W)
        Label(frame1, text="Municipio: ", bg="deep sky blue",  font=("Arial Black", 11)).grid(row=5, column=2, sticky=E)
        municipio = ttk.Combobox(frame1, width="15", textvariable=municipio1, state="readonly")
        municipio.grid(row=5, column=3, sticky=W)
        departamento.bind("<<ComboboxSelected>>", lugar_seleccion)
        departamento.current(0)

        def chip(event):
            sel_valor = municipio.get()
            if sel_valor == "Bogotá_D.C":
                label_chip.grid(row=1, column=2, sticky=E)
                entry_chip.grid(row=1, column=3, sticky=W)
            else:
                label_chip.grid_remove()
                entry_chip.grid_remove()
        
        municipio.bind("<<ComboboxSelected>>", chip)

        label_chip=Label(frame1, text="Chip:", bg="deep sky blue",  font=("Arial Black", 11))
        label_chip.grid(row=1, column=2, sticky=E)
        label_chip.grid_remove()
        entry_chip=Entry(frame1, textvariable=entry_chip1)
        entry_chip.grid(row=1, column=3, sticky=W)
        entry_chip.grid_remove()

        Label(frame1, text="Nombre o Dirección:", bg="deep sky blue", font=("Arial Black", 11)).grid(row=6, column=2, sticky=E)
        direccionNueva_entry=Entry(frame1, textvariable=direccionNueva1)
        direccionNueva_entry.grid(row=6, column=3, sticky=E+W)

        Label(frame1, text="", bg="deep sky blue").grid(row=6, column=4)




        

        

        frame2=LabelFrame(ventana5, text="Tradición", bg="deep sky blue", fg="white", width=140, height=1, font=("Arial Black", 14))
        frame2.place(x=625, y=2)

        Label(frame2, text="Modo de Adquirir:", bg="deep sky blue", font=("Arial Black", 11)).grid(row=1, column=0, sticky=E)
        adquirir=ttk.Combobox(frame2, textvariable=Modo_adquirir1, width=48, state="roadonly", values=["","Compra-Venta", "Dación en Pago", "Transferencia a Titulo de Beneficio en Fiducia Mercantil", "Adjudicación en Sucesión", "Adjuducación en Liquidación de sociedad conyugal", "Adjudicación en Sociedad comercial", "Adjudicacion de Baldios", "Adjudicación en Remate", "Adjudicación en sociedad de comunidad", "Adjudicacion sucesión participación adicional", "Permuta"])
        adquirir.grid(row=1, column=1, columnspan=3, sticky=E+W)
        adquirir.current(0)

        Label(frame2, text="Escritura publica:",  bg="deep sky blue", font=("Arial Black", 11)).grid(row=2, column=0, sticky=E)
        escritura=Entry(frame2, textvariable=escritura1)
        escritura.grid(row=2, column=1, sticky=W)

        Label(frame2, text="Fecha:Dia/Mes/Año ",  bg="deep sky blue", font=("Arial Black", 11)).grid(row=3, column=0)
        dia_comb=ttk.Combobox(frame2, width=15, textvariable=dia1, state="readonly", values=["","Uno (1)", "Dos (2)", "Tres (3)",
        "Cuatro (4)", "Cinco (5)", "Seis (6)", "Siete (7)", "Ocho (8)", "Nueve (9)", "Diez (10)", "Once (11)", 
        "Doce (12)", "Trece (13)", "Catorce (14)", "Quince (15)", "Dieciseis (16)", "Dicisiete (17)", "Dieciocho (18)", 
        "Diecinueve (19)", "Veinte (20)", "Ventiuno (21)", "Ventidos (22)", "Veintitres (23)", "Veinticuatro (24)", 
        "Veinticinco (25)", "Veintiseis (26)", "Veintisiete (27)", "Veintiocho (28)", "Veintinueve (29)", 
        "Treinta (30)", "Treinta y uno (31)"])
        dia_comb.grid(row=3, column=1, sticky=E+W)
        dia_comb.current(0)

        
        mes_comb=ttk.Combobox(frame2, width=12, textvariable=mes1, state="roadonly", values=["","Enero", "Febrero", "Marzo", "Abril", "Mayo", 
        "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"] )
        mes_comb.grid(row=3, column=2, padx=10)
        mes_comb.current(0)
            
        año_comb=Entry(frame2, textvariable=año1)
        año_comb.grid(row=3, column=3)

        
        Label(frame2, text="Notaria:",  bg="deep sky blue", font=("Arial Black", 11)).grid(row=4, column=0, sticky=E)
        notaria=ttk.Combobox(frame2, width=30, state="roadonly", textvariable=notaria1, values=["","Unica", "Primera (1ra)", "segunda (2da)", "Tercera (3ra)", "Cuarta (4a)", "Quinta (5a)", "Sexta (6a)", "Séptima (7a)", "Octava(8a)", "Novena(9a)", "Decima (10)", "Once (11)", "Doce (12)", "Trece (13)", "Catorce (14)", "Quince (15)", "Dieciseis (16)", "Diecisiete (17)", "Dieciocho(18)", "Diecinueve(19)", "Veinte (20)", "Veintiuno (21)", "Venitidós (22)", "Veintitres (23)", "veinticuatro (24)", "Veinticinco (25)", "Veintiseis (26)", "Veintisiete (27)", "Ventiocho(28)", "Ventinueve(29)", "Treinta (30)", "Treinta y uno (31)", "Treinta y dós (32)", "Treinta y tres (33)", "Treinta y cuatro (34)", "Treinta y cinco (35)", "Treinta y seis (36)", "Treinta y siete (37)", "Treinta y ocho (38)", "Treinta y nueve (39)", "Cuarenta (40)", "Cuarenta y uno (41)", "Cuarenta y dós (42)", "Cuarenta y tres (43)", "Cuarenta y cuatro (44)", "Cuarenta y cinco (45)", "Cuarenta y seis (46)", "Cuarenta y siete (47)", "Cuarenta y ocho (48)", "Cuarenta y nueve (49)", "Cincuenta (50)", "Cincuenta y uno (51)", "Cincuenta y dós (52)", "Cincuenta y tres (53)", "Cincuenta y cuatro (54)", "Cincuenta y cinco (55)", "Cincuenta y seis (56)", "Cincuenta y siete (57)", "Cincuenta y ocho (58)", "Cincuenta y nueve (59)", "Sesenta (60)", "Sesenta y uno (61)", "Sesenta y dós (62)", "Sesenta y tres (63)", "Sesenta y cuatro (64)", "Sesenta y cinco (65)", "Sesenta y seis (66)", "Sesenta y siete (67)", "Sesenta y ocho (68)", "Sesenta y nueve (69)", "Setenta(70)", "Setenta y uno (71)", "Setenta y dós (72)", "Setenta y tres (73)", "Setenta y cuatro (74)", "Setenta y cinco (75)", "Setenta y seis (76)", "Setenta y siete (77)", "Setenta y ocho (78)", "Setenta y nueve (79)", "Ochenta(80)", "Ochenta y uno(81)", "Ochenta y dós(82)"])
        notaria.grid(row=4, column=1, sticky=W)
        notaria.current(0)

        Label(frame2, text="Municipio Notaria:",  bg="deep sky blue", font=("Arial Black", 11)).grid(row=5, column=0, sticky=E)
        municipioNota=Entry(frame2, textvariable=Notaria_municipio1)
        municipioNota.grid(row=5, column=1, sticky=E+W)

        conn = sqlite3.connect('login1.db')
        cursor = conn.cursor()

        # Consulta para obtener los valores de la segunda columna
        cursor.execute("SELECT Nombre_Notario FROM registro_notario")
        options = [row[0] for row in cursor.fetchall()]

        Label(frame2, text="Notario:",  bg="deep sky blue", font=("Arial Black", 11)).grid(row=6, column=0, sticky=E)
        notario=ttk.Combobox(frame2, textvariable=notario1, values=options)
        notario.grid(row=6, column=1, sticky=E+W)

        
        notariod = Entry(frame2, textvariable=notariad1)
        notariod.grid(row=6, column=2, sticky=E+W)

        notaria_actual = Entry(frame2, textvariable=notaria_actual1)
        notaria_actual.grid(row=6, column=3, sticky=E+W)

        labeldepartamento=Label(ventana5, text="Dep. Notaria", bg="deep sky blue", font=("Arial Black", 11))
        labeldepartamento.place(relx=0.02, rely=0.86)
        notario_notaria = Entry(ventana5, textvariable=notario_notaria1)
        notario_notaria.place(relx=0.02, rely=0.9)

        labelmunicipio=Label(ventana5, text="Mun. Notaria", bg="deep sky blue", font=("Arial Black", 11))
        labelmunicipio.place(relx=0.12, rely=0.86)
        municipio_not= Entry(ventana5, textvariable=municipio_not1)
        municipio_not.place(relx=0.12, rely=0.9)


        def update_entry(event):
            # Obtener el valor seleccionado en el combobox
            selected_value = notario.get()
            # Consulta para obtener los valores de las columnas Declaración y notaria_actual correspondientes al valor seleccionado
            cursor.execute("SELECT Declaración, notaria_actual, notario_notaria, municipio1 FROM registro_notario WHERE Nombre_Notario=?", (selected_value,))
            result = cursor.fetchone()
            if result:
                # Actualizar el valor del Entry para la columna Declaración
                notariod.delete(0, END)
                notariod.insert(0, result[0])
                # Actualizar el valor del Entry para la columna notaria_actual
                notaria_actual.delete(0, END)
                notaria_actual.insert(0, result[1])

                notario_notaria.delete(0, END)
                notario_notaria.insert(0, result[2])

                municipio_not.delete(0, END)
                municipio_not.insert(0, result[3])

        # Vincular la función al evento de selección del combobox
        notario.bind("<<ComboboxSelected>>", update_entry)

            
        Label(frame2, text=" ", bg="deep sky blue").grid (row=1, column=4, sticky=E)

        conn = sqlite3.connect('login1.db')
        cursor = conn.cursor()
        # Consulta para obtener los valores 
        cursor.execute("SELECT Cedula FROM Ingresar_usuario1")
        opcionesu1 = [row[0] for row in cursor.fetchall()]


        frame3=LabelFrame(ventana5, text="D:", bg="deep sky blue", fg="white", width=140, height=1, font=("Arial Black", 14))
        frame3.place(x=3, y=203)
        Label(frame3, text="Cedula", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=0)
        buscarcedulav1_combo=ttk.Combobox(frame3, width=12, textvariable=buscarcedulav1, values=opcionesu1)
        buscarcedulav1_combo.grid(row=1, column=0)

        def actualizar_combobox(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valor = buscarcedulav1_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valor,))
            result1 = cursor.fetchone()
            if result1:
                # Actualizar el valor del Entry para la column 
                cedulavendedor_entry.delete(0, END)
                cedulavendedor_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                primerApellido1_entry.delete(0, END)
                primerApellido1_entry.insert(0, result1[1])

                segundoApellido1_entry.delete(0, END)
                segundoApellido1_entry.insert(0, result1[2])

                primerNombre1_entry.delete(0, END)
                primerNombre1_entry.insert(0, result1[3])

                segundoNombre1_entry.delete(0, END)
                segundoNombre1_entry.insert(0, result1[4])
                
                sexo1_entry.delete(0, END)
                sexo1_entry.insert(0, result1[5])

                domicilio1_entry.delete(0, END)
                domicilio1_entry.insert(0, result1[6])

                estadocivil1_entry.delete(0, END)
                estadocivil1_entry.insert(0, result1[7])

                 

        # Vincular la función al evento de selección del combobox
        buscarcedulav1_combo.bind("<<ComboboxSelected>>", actualizar_combobox)


        Label(frame3, text="C.C", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=1)
        cedulavendedor_entry=Entry(frame3, textvariable=ven1_cedulavendedor)
        cedulavendedor_entry.grid(row=1, column=1)

        
        Label(frame3, text="Primer Apellido", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=2)
        primerApellido1_entry=Entry(frame3, textvariable=ven1_primerApellido)
        primerApellido1_entry.grid(row=1, column=2)

        Label(frame3, text="Segundo Apellido", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=3)
        segundoApellido1_entry=Entry(frame3, textvariable=ven1_segundoApellido)
        segundoApellido1_entry.grid(row=1, column=3)

        Label(frame3, text="Primer Nombre", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=4)
        primerNombre1_entry=Entry(frame3, textvariable=ven1_primerNombre)
        primerNombre1_entry.grid(row=1, column=4)

        Label(frame3, text="Segundo Nombre", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=5)
        segundoNombre1_entry=Entry(frame3, textvariable=ven1_segundoNombre)
        segundoNombre1_entry.grid(row=1, column=5)

        
        
        fecha_nacimiento1=Entry(frame3, width=1, textvariable=ven1_fechadenacimiento)
        fecha_nacimiento1.grid(row=1, column=7)
        
        
        

        rh1_entry=Entry(frame3, width=1, textvariable=ven1_rh)
        rh1_entry.grid(row=1, column=8)
        
        ec_entry=Entry(frame3, width=1)
        ec_entry.grid(row=1, column=9)

        Label(frame3, text="Domicilio", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=10)
        domicilio1_entry=Entry(frame3, textvariable=ven1_domicilio)
        domicilio1_entry.grid(row=1, column=10)

        opcionesV1 = {
            "":(),
            "F" : ("soltera sin unión marital de hecho", "soltera con unión marital de hecho", "casada con sociedad conyugal vigente", "casada con sociedad conyugal disuelta y liquidada", "viuda con sociedad conyugal disuelta y liquidada", "viuda con sociedad conyugal pendiente de liquidación"),
            "M" : ("soltero sin unión marital de hecho", "soltero con unión marital de hecho", "casado con sociedad conyugal vigente", "casado con sociedad conyugal disuelta y liquidada", "viudo con sociedad conyugal disuelta y liquidada", "viudo con sociedad conyugal pendiente de liquidación"),
            }
        
        def escoger_gv1(event):

                
            estadocivil1_entry.set("")
            estadocivil1_entry.config(values=opcionesV1[sexo1_entry.get()])
        
        def escoger_gc1(event):
            estadocivil2_entry.set("")
            estadocivil2_entry.config(values=opcionesV1[sexo2_entry.get()])

        def escoger_gv2(event):
            Ven2_estadocivil_entry.set("")
            Ven2_estadocivil_entry.config(values=opcionesV1[Ven2_sexo_entry.get()])
        
        def escoger_gv3(event):
            Ven3_estadocivil_entry.set("")
            Ven3_estadocivil_entry.config(values=opcionesV1[Ven3_sexo_entry.get()])
        
        def escoger_gv4(event):
            Ven4_estadocivil_entry.set("")
            Ven4_estadocivil_entry.config(values=opcionesV1[Ven4_sexo_entry.get()])
        
        def escoger_gv5(event):
            Ven5_estadocivil_entry.set("")
            Ven5_estadocivil_entry.config(values=opcionesV1[Ven5_sexo_entry.get()])
        
        def escoger_gv6(event):
            Ven6_estadocivil_entry.set("")
            Ven6_estadocivil_entry.config(values=opcionesV1[Ven6_sexo_entry.get()])
        
        def escoger_gc2(event):
            con2_estadocivil_entry.set("")
            con2_estadocivil_entry.config(values=opcionesV1[con2_sexo_entry.get()])
        
        def escoger_gc3(event):
            con3_estadocivil_entry.set("")
            con3_estadocivil_entry.config(values=opcionesV1[con3_sexo_entry.get()])

        def escoger_gc4(event):
            con4_estadocivil_entry.set("")
            con4_estadocivil_entry.config(values=opcionesV1[con4_sexo_entry.get()])
        
        def escoger_gc5(event):
            con5_estadocivil_entry.set("")
            con5_estadocivil_entry.config(values=opcionesV1[con5_sexo_entry.get()])

        def escoger_gc6(event):
            con6_estadocivil_entry.set("")
            con6_estadocivil_entry.config(values=opcionesV1[con6_sexo_entry.get()])



        Label(frame3, text="Sexo", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=6)
        sexo1_entry=ttk.Combobox(frame3, width=5, values=tuple(opcionesV1.keys()))
        sexo1_entry.grid(row=1, column=6)

       

        Label(frame3, text="Estado Civil", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=11)
        estadocivil1_entry=ttk.Combobox(frame3, state="roadonly", textvariable=ven1_estadocivil, width=30)
        estadocivil1_entry.grid(row=1, column=11)
        sexo1_entry.bind("<<ComboboxSelected>>", escoger_gv1)
        sexo1_entry.current(0)

        conn = sqlite3.connect('login1.db')
        cursor = conn.cursor()
        # Consulta para obtener los valores 
        cursor.execute("SELECT Cedula FROM Ingresar_usuario1")
        opcionesu6 = [row[0] for row in cursor.fetchall()]   

        



        frame4=LabelFrame(ventana5, text="A:", bg="deep sky blue", fg="white", width=140, height=1, font=("Arial Black", 14))
        frame4.place(x=3, y=400)

        Label(frame4, text="Cedula", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=0)
        buscarcedulac1_combo=ttk.Combobox(frame4, width=12, textvariable=buscarcedulac1, values=opcionesu6)
        buscarcedulac1_combo.grid(row=1, column=0)

        def actualizar_comboboxc1(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valor = buscarcedulac1_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valor,))
            result1 = cursor.fetchone()
            if result1:
                                # Actualizar el valor del Entry para la column 
                cedulacomprador_entry.delete(0, END)
                cedulacomprador_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                primer_apellido2_entry.delete(0, END)
                primer_apellido2_entry.insert(0, result1[1])

                segundo_apellido2_entry.delete(0, END)
                segundo_apellido2_entry.insert(0, result1[2])

                primer_nombre2_entry.delete(0, END)
                primer_nombre2_entry.insert(0, result1[3])

                segundo_nombre2_entry.delete(0, END)
                segundo_nombre2_entry.insert(0, result1[4])
                
                sexo2_entry.delete(0, END)
                sexo2_entry.insert(0, result1[5])

                domicilio2_entry.delete(0, END)
                domicilio2_entry.insert(0, result1[6])

                estadocivil2_entry.delete(0, END)
                estadocivil2_entry.insert(0, result1[7])

   

        # Vincular la función al evento de selección del combobox
        buscarcedulac1_combo.bind("<<ComboboxSelected>>", actualizar_comboboxc1)

        

        Label(frame4, text="C.C", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=1)
        cedulacomprador_entry=Entry(frame4, textvariable=cedulacomprador1)
        cedulacomprador_entry.grid(row=1, column=1)

        Label(frame4, text="Primer Apellido", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=2)
        primer_apellido2_entry=Entry(frame4, textvariable=con1_primer_apellido)
        primer_apellido2_entry.grid(row=1, column=2)

        Label(frame4, text="Segundo Apellido", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=3)
        segundo_apellido2_entry=Entry(frame4, textvariable=con1_segundo_apellido)
        segundo_apellido2_entry.grid(row=1, column=3)

        Label(frame4, text="Primer Nombre", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=4)
        primer_nombre2_entry=Entry(frame4, textvariable=con1_primer_nombre)
        primer_nombre2_entry.grid(row=1, column=4)

        Label(frame4, text="Segundo Nombre", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=5)
        segundo_nombre2_entry=Entry(frame4, textvariable=con1_segundo_nombre)
        segundo_nombre2_entry.grid(row=1, column=5)


    

        
        fecha_nacimiento2=Entry(frame4, width=1, textvariable=con1_fechanacimiento)
        fecha_nacimiento2.grid(row=1, column=7)

        
        rh2_entry=Entry(frame4, width=1, textvariable=con1_rh)
        rh2_entry.grid(row=1, column=8)

        ecc_entry=Entry(frame4, width=1)
        ecc_entry.grid(row=1, column=9)


        Label(frame4, text="Domicilio", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=10)
        domicilio2_entry=Entry(frame4, textvariable=con1_domicilio)
        domicilio2_entry.grid(row=1, column=10)

        

    


        Label(frame4, text="Sexo", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=6)
        sexo2_entry=ttk.Combobox(frame4, textvariable=con1_sexo, width=5,  values=tuple(opcionesV1.keys()))
        sexo2_entry.grid(row=1, column=6)

        Label(frame4, text="Estado Civil", bg="deep sky blue", font=("Arial Black", 9)).grid(row=0, column=11)
        estadocivil2_entry=ttk.Combobox(frame4, state="roadonly", width=30, textvariable=con1_estadocivil)
        estadocivil2_entry.grid(row=1, column=11)
        sexo2_entry.bind("<<ComboboxSelected>>", escoger_gc1)
        sexo2_entry.current(0)
        

                    

        def ingresarvendedor2():
            buscarcedulav2_combo.grid()
            cedulavendedor2_entry.grid()
            cedulavendedor2_entry.focus()
            Ven2_primerApellido_entry.grid()
            Ven2_segundoApellido_entry.grid()
            Ven2_primerNombre_entry.grid()
            Ven2_segundoNombre_entry.grid()
            Ven2_sexo_entry.grid()
            Ven2_fecha_nacimiento.grid()
            Ven2_rh_entry.grid()
            Ven2_domicilio_entry.grid()
            Ven2_estadocivil_entry.grid()
            button2.grid()
            boton_agregar_ven3.grid()
            boton_quitar_ven2.grid()
            Ven2_ec_entry.grid()
            boton_agregar_ven2.grid_remove()
            button1.grid_remove()


        def eliminarvendedor2():
            buscarcedulav2_combo.grid_remove()
            cedulavendedor2_entry.grid_remove()
            Ven2_primerApellido_entry.grid_remove()
            Ven2_segundoApellido_entry.grid_remove()
            Ven2_primerNombre_entry.grid_remove()
            Ven2_segundoNombre_entry.grid_remove()
            Ven2_sexo_entry.grid_remove()
            Ven2_fecha_nacimiento.grid_remove()
            Ven2_rh_entry.grid_remove()
            Ven2_domicilio_entry.grid_remove()
            Ven2_estadocivil_entry.grid_remove()
            boton_agregar_ven3.grid_remove()
            boton_quitar_ven2.grid_remove()
            button2.grid_remove()
            Ven2_ec_entry.grid_remove()
            boton_agregar_ven2.grid()
            button1.grid()

            cedulavendedor2_entry.delete(0, END)
            Ven2_primerApellido_entry.delete(0, END)
            Ven2_segundoApellido_entry.delete(0, END)
            Ven2_primerNombre_entry.delete(0, END)
            Ven2_segundoNombre_entry.delete(0, END)
            Ven2_sexo_entry.delete(0, END)
            Ven2_fecha_nacimiento.delete(0, END)
            Ven2_rh_entry.delete(0, END)
            Ven2_domicilio_entry.delete(0, END)
            Ven2_estadocivil_entry.delete(0, END)

            

        def ingresarvendedor3():
            buscarcedulav3_combo.grid()
            cedulavendedor3_entry.grid()
            cedulavendedor3_entry.focus()
            Ven3_primerApellido_entry.grid()
            Ven3_segundoApellido_entry.grid()
            Ven3_primerNombre_entry.grid()
            Ven3_segundoNombre_entry.grid()
            Ven3_sexo_entry.grid()
            Ven3_fecha_nacimiento.grid()
            Ven3_rh_entry.grid()
            Ven3_ec_entry.grid()
            Ven3_domicilio_entry.grid()
            Ven3_estadocivil_entry.grid()
            boton_agregar_ven4.grid()
            boton_quitar_ven3.grid()
            button3.grid()
            boton_agregar_ven3.grid_remove()
            boton_quitar_ven2.grid_remove()
            button2.grid_remove()

        def eliminarvendedor3():
            buscarcedulav3_combo.grid_remove()
            cedulavendedor3_entry.grid_remove()
            Ven3_primerApellido_entry.grid_remove()
            Ven3_segundoApellido_entry.grid_remove()
            Ven3_primerNombre_entry.grid_remove()
            Ven3_segundoNombre_entry.grid_remove()
            Ven3_sexo_entry.grid_remove()
            Ven3_fecha_nacimiento.grid_remove()
            Ven3_rh_entry.grid_remove()
            Ven3_ec_entry.grid_remove()
            Ven3_domicilio_entry.grid_remove()
            Ven3_estadocivil_entry.grid_remove()
            boton_agregar_ven4.grid_remove()
            boton_quitar_ven3.grid_remove()
            button3.grid_remove()
            boton_agregar_ven3.grid()
            boton_quitar_ven2.grid()
            button2.grid()

            cedulavendedor3_entry.delete(0, END)
            Ven3_primerApellido_entry.delete(0, END)
            Ven3_segundoApellido_entry.delete(0, END)
            Ven3_primerNombre_entry.delete(0, END)
            Ven3_segundoNombre_entry.delete(0, END)
            Ven3_sexo_entry.delete(0, END)
            Ven3_fecha_nacimiento.delete(0, END)
            Ven3_rh_entry.delete(0, END)
            Ven3_domicilio_entry.delete(0, END)
            Ven3_estadocivil_entry.delete(0, END)


        def ingresarvendedor4():
            buscarcedulav4_combo.grid()
            cedulavendedor4_entry.grid()
            cedulavendedor4_entry.focus()
            Ven4_primerApellido_entry.grid()
            Ven4_segundoApellido_entry.grid()
            Ven4_primerNombre_entry.grid()
            Ven4_segundoNombre_entry.grid()
            Ven4_sexo_entry.grid()
            Ven4_fecha_nacimiento.grid()
            Ven4_rh_entry.grid()
            Ven4_ec_entry.grid()
            button4.grid()
            
            Ven4_domicilio_entry.grid()
            Ven4_estadocivil_entry.grid()
            boton_agregar_ven5.grid()
            boton_quitar_ven4.grid()
            boton_agregar_ven4.grid_remove()
            boton_quitar_ven3.grid_remove()
            button3.grid_remove()

        def eliminarvendedor4():
            buscarcedulav4_combo.grid_remove()
            cedulavendedor4_entry.grid_remove()
            Ven4_primerApellido_entry.grid_remove()
            Ven4_segundoApellido_entry.grid_remove()
            Ven4_primerNombre_entry.grid_remove()
            Ven4_segundoNombre_entry.grid_remove()
            Ven4_sexo_entry.grid_remove()
            Ven4_fecha_nacimiento.grid_remove()
            Ven4_rh_entry.grid_remove()
            Ven4_ec_entry.grid_remove()
            button4.grid_remove()
            Ven4_domicilio_entry.grid_remove()
            Ven4_estadocivil_entry.grid_remove()
            boton_agregar_ven5.grid_remove()
            boton_quitar_ven4.grid_remove()
            boton_agregar_ven4.grid()
            boton_quitar_ven3.grid()
            button3.grid()

            cedulavendedor4_entry.delete(0, END)
            Ven4_primerApellido_entry.delete(0, END)
            Ven4_segundoApellido_entry.delete(0, END)
            Ven4_primerNombre_entry.delete(0, END)
            Ven4_segundoNombre_entry.delete(0, END)
            Ven4_sexo_entry.delete(0, END)
            Ven4_fecha_nacimiento.delete(0, END)
            Ven4_rh_entry.delete(0, END)
            Ven4_domicilio_entry.delete(0, END)
            Ven4_estadocivil_entry.delete(0, END)

        def ingresarvendedor5():
            buscarcedulav5_combo.grid()
            cedulavendedor5_entry.grid()
            cedulavendedor5_entry.focus()
            Ven5_primerApellido_entry.grid()
            Ven5_segundoApellido_entry.grid()
            Ven5_primerNombre_entry.grid()
            Ven5_segundoNombre_entry.grid()
            Ven5_sexo_entry.grid()
            Ven5_fecha_nacimiento.grid()
            Ven5_rh_entry.grid()
            Ven5_ec_entry.grid()
            Ven5_domicilio_entry.grid()
            Ven5_estadocivil_entry.grid()
            boton_agregar_ven6.grid()
            boton_quitar_ven5.grid()
            button5.grid()
            boton_agregar_ven5.grid_remove()
            boton_quitar_ven4.grid_remove()
            button4.grid_remove()
        
        def eliminarvendedor5():
            buscarcedulav5_combo.grid_remove()
            cedulavendedor5_entry.grid_remove()
            Ven5_primerApellido_entry.grid_remove()
            Ven5_segundoApellido_entry.grid_remove()
            Ven5_primerNombre_entry.grid_remove()
            Ven5_segundoNombre_entry.grid_remove()
            Ven5_sexo_entry.grid_remove()
            Ven5_fecha_nacimiento.grid_remove()
            Ven5_rh_entry.grid_remove()
            Ven5_ec_entry.grid_remove()
            Ven5_domicilio_entry.grid_remove()
            Ven5_estadocivil_entry.grid_remove()
            boton_agregar_ven6.grid_remove()
            boton_quitar_ven5.grid_remove()
            button5.grid_remove()
            boton_agregar_ven5.grid()
            boton_quitar_ven4.grid()
            button4.grid()

            cedulavendedor5_entry.delete(0, END)
            Ven5_primerApellido_entry.delete(0, END)
            Ven5_segundoApellido_entry.delete(0, END)
            Ven5_primerNombre_entry.delete(0, END)
            Ven5_segundoNombre_entry.delete(0, END)
            Ven5_sexo_entry.delete(0, END)
            Ven5_fecha_nacimiento.delete(0, END)
            Ven5_rh_entry.delete(0, END)
            Ven5_domicilio_entry.delete(0, END)
            Ven5_estadocivil_entry.delete(0, END)

        def ingresarvendedor6():
            buscarcedulav6_combo.grid()
            cedulavendedor6_entry.grid()
            cedulavendedor6_entry.focus()
            Ven6_primerApellido_entry.grid()
            Ven6_segundoApellido_entry.grid()
            Ven6_primerNombre_entry.grid()
            Ven6_segundoNombre_entry.grid()
            Ven6_sexo_entry.grid()
            Ven6_fecha_nacimiento.grid()
            Ven6_rh_entry.grid()
            Ven6_ec_entry.grid()
            Ven6_domicilio_entry.grid()
            Ven6_estadocivil_entry.grid()
            boton_quitar_ven6.grid()
            button6.grid()
            boton_agregar_ven6.grid_remove()
            boton_quitar_ven5.grid_remove()
            button5.grid_remove()
        
        def eliminarvendedor6():
            buscarcedulav6_combo.grid_remove()
            cedulavendedor6_entry.grid_remove()
            Ven6_primerApellido_entry.grid_remove()
            Ven6_segundoApellido_entry.grid_remove()
            Ven6_primerNombre_entry.grid_remove()
            Ven6_segundoNombre_entry.grid_remove()
            Ven6_sexo_entry.grid_remove()
            Ven6_fecha_nacimiento.grid_remove()
            Ven6_rh_entry.grid_remove()
            Ven6_ec_entry.grid_remove()
            Ven6_domicilio_entry.grid_remove()
            Ven6_estadocivil_entry.grid_remove()
            boton_quitar_ven6.grid_remove()
            button6.grid_remove()
            boton_agregar_ven6.grid()
            boton_quitar_ven5.grid()
            button5.grid()

            cedulavendedor6_entry.delete(0, END)
            Ven6_primerApellido_entry.delete(0, END)
            Ven6_segundoApellido_entry.delete(0, END)
            Ven6_primerNombre_entry.delete(0, END)
            Ven6_segundoNombre_entry.delete(0, END)
            Ven6_sexo_entry.delete(0, END)
            Ven6_fecha_nacimiento.delete(0, END)
            Ven6_rh_entry.delete(0, END)
            Ven6_domicilio_entry.delete(0, END)
            Ven6_estadocivil_entry.delete(0, END)
        
        def ingresarcomprador2():
            buscarcedulac2_combo.grid()
            cedulacomprador2_entry.grid()
            cedulacomprador2_entry.focus()
            con2_primerApellido_entry.grid()
            con2_segundoApellido_entry.grid()
            con2_primerNombre_entry.grid()
            con2_segundoNombre_entry.grid()
            con2_sexo_entry.grid()
            con2_fecha_nacimiento.grid()
            con2_rh_entry.grid()
            con2_ec_entry.grid()
            con2_domicilio_entry.grid()
            con2_estadocivil_entry.grid()
            boton_agregar_con3.grid()
            boton_quitar_con2.grid()
            button8.grid()
            button7.grid_remove()
            boton_agregar_con2.grid_remove()

            cedulavendedor3_entry.grid_remove()
            Ven3_primerApellido_entry.grid_remove()
            Ven3_segundoApellido_entry.grid_remove()
            Ven3_primerNombre_entry.grid_remove()
            Ven3_segundoNombre_entry.grid_remove()
            Ven3_sexo_entry.grid_remove()
            Ven3_fecha_nacimiento.grid_remove()
            Ven3_rh_entry.grid_remove()
            Ven3_domicilio_entry.grid_remove()
            Ven3_estadocivil_entry.grid_remove()
            boton_agregar_ven4.grid_remove()
            boton_quitar_ven3.grid_remove()


        def eliminarcomprador2():
            buscarcedulac2_combo.grid_remove()
            cedulacomprador2_entry.grid_remove()
            con2_primerApellido_entry.grid_remove()
            con2_segundoApellido_entry.grid_remove()
            con2_primerNombre_entry.grid_remove()
            con2_segundoNombre_entry.grid_remove()
            con2_sexo_entry.grid_remove()
            con2_fecha_nacimiento.grid_remove()
            con2_rh_entry.grid_remove()
            con2_domicilio_entry.grid_remove()
            con2_estadocivil_entry.grid_remove()
            boton_agregar_con3.grid_remove()
            boton_quitar_con2.grid_remove()
            con2_ec_entry.grid_remove()
            boton_agregar_con2.grid()
            button7.grid()
            button8.grid_remove()
            
            cedulacomprador2_entry.delete(0, END) 
            con2_primerApellido_entry.delete(0, END)
            con2_segundoApellido_entry.delete(0, END)
            con2_primerNombre_entry.delete(0, END)
            con2_segundoNombre_entry.delete(0, END)
            con2_sexo_entry.delete(0, END)
            con2_fecha_nacimiento.delete(0, END)
            con2_rh_entry.delete(0, END)
            con2_domicilio_entry.delete(0, END)
            con2_estadocivil_entry.delete(0, END)

            



        def ingresarcomprador3():
            buscarcedulac3_combo.grid()
            cedulacomprador3_entry.grid()
            cedulacomprador3_entry.focus()
            con3_primerApellido_entry.grid()
            con3_segundoApellido_entry.grid()
            con3_primerNombre_entry.grid()
            con3_segundoNombre_entry.grid()
            con3_sexo_entry.grid()
            con3_fecha_nacimiento.grid()
            con3_rh_entry.grid()
            con3_ec_entry.grid()
            con3_domicilio_entry.grid()
            con3_estadocivil_entry.grid()
            boton_agregar_con4.grid()
            boton_quitar_con3.grid()
            button9.grid()
            boton_agregar_con3.grid_remove()
            boton_quitar_con2.grid_remove()
            button8.grid_remove()

        def eliminarcomprador3():
            buscarcedulac3_combo.grid_remove()
            cedulacomprador3_entry.grid_remove()
            con3_primerApellido_entry.grid_remove()
            con3_segundoApellido_entry.grid_remove()
            con3_primerNombre_entry.grid_remove()
            con3_segundoNombre_entry.grid_remove()
            con3_sexo_entry.grid_remove()
            con3_fecha_nacimiento.grid_remove()
            con3_rh_entry.grid_remove()
            con3_ec_entry.grid_remove()
            con3_domicilio_entry.grid_remove()
            con3_estadocivil_entry.grid_remove()
            boton_agregar_con4.grid_remove()
            boton_quitar_con3.grid_remove()
            button9.grid_remove()
            button8.grid()
            boton_agregar_con3.grid()
            boton_quitar_con2.grid()

            
            cedulacomprador3_entry.delete(0, END)
            con3_primerApellido_entry.delete(0, END)
            con3_segundoApellido_entry.delete(0, END)
            con3_primerNombre_entry.delete(0, END)
            con3_segundoNombre_entry.delete(0, END)
            con3_sexo_entry.delete(0, END)
            con3_fecha_nacimiento.delete(0, END)
            con3_rh_entry.delete(0, END)
            con3_domicilio_entry.delete(0, END)
            con3_estadocivil_entry.delete(0, END)

        def ingresarcomprador4():
            buscarcedulac4_combo.grid()
            cedulacomprador4_entry.grid()
            cedulacomprador4_entry.focus()
            con4_primerApellido_entry.grid()
            con4_segundoApellido_entry.grid()
            con4_primerNombre_entry.grid()
            con4_segundoNombre_entry.grid()
            con4_sexo_entry.grid()
            con4_fecha_nacimiento.grid()
            con4_rh_entry.grid()
            con4_ec_entry.grid()
            con4_domicilio_entry.grid()
            con4_estadocivil_entry.grid()
            boton_agregar_con5.grid()
            boton_quitar_con4.grid()
            boton_agregar_con4.grid_remove()
            boton_quitar_con3.grid_remove()
            button9.grid_remove()
            button10.grid()

        def eliminarcomprador4():
            buscarcedulac4_combo.grid_remove()
            cedulacomprador4_entry.grid_remove()
            con4_primerApellido_entry.grid_remove()
            con4_segundoApellido_entry.grid_remove()
            con4_primerNombre_entry.grid_remove()
            con4_segundoNombre_entry.grid_remove()
            con4_sexo_entry.grid_remove()
            con4_fecha_nacimiento.grid_remove()
            con4_rh_entry.grid_remove()
            con4_ec_entry.grid_remove()
            con4_domicilio_entry.grid_remove()
            con4_estadocivil_entry.grid_remove()
            boton_agregar_con5.grid_remove()
            boton_quitar_con4.grid_remove()
            boton_agregar_con4.grid()
            boton_quitar_con3.grid()
            button9.grid()
            button10.grid_remove()


            
            cedulacomprador4_entry.delete(0, END)
            con4_primerApellido_entry.delete(0, END)
            con4_segundoApellido_entry.delete(0, END)
            con4_primerNombre_entry.delete(0, END)
            con4_segundoNombre_entry.delete(0, END)
            con4_sexo_entry.delete(0, END)
            con4_fecha_nacimiento.delete(0, END)
            con4_rh_entry.delete(0, END)
            con4_domicilio_entry.delete(0, END)
            con4_estadocivil_entry.delete(0, END)

        def ingresarcomprador5():
            buscarcedulac5_combo.grid()
            cedulacomprador5_entry.grid()
            cedulacomprador5_entry.focus()
            con5_primerApellido_entry.grid()
            con5_segundoApellido_entry.grid()
            con5_primerNombre_entry.grid()
            con5_segundoNombre_entry.grid()
            con5_sexo_entry.grid()
            con5_fecha_nacimiento.grid()
            con5_rh_entry.grid()
            con5_ec_entry.grid()
            con5_domicilio_entry.grid()
            con5_estadocivil_entry.grid()
            boton_agregar_con6.grid()
            boton_quitar_con5.grid()
            boton_agregar_con5.grid_remove()
            boton_quitar_con4.grid_remove()
            button10.grid_remove()
            button11.grid()
        
        def eliminarcomprador5():
            buscarcedulac5_combo.grid_remove()
            cedulacomprador5_entry.grid_remove()
            con5_primerApellido_entry.grid_remove()
            con5_segundoApellido_entry.grid_remove()
            con5_primerNombre_entry.grid_remove()
            con5_segundoNombre_entry.grid_remove()
            con5_sexo_entry.grid_remove()
            con5_fecha_nacimiento.grid_remove()
            con5_rh_entry.grid_remove()
            con5_ec_entry.grid_remove()
            con5_domicilio_entry.grid_remove()
            con5_estadocivil_entry.grid_remove()
            boton_agregar_con6.grid_remove()
            boton_quitar_con5.grid_remove()
            boton_agregar_con5.grid()
            boton_quitar_con4.grid()
            button10.grid()
            button11.grid_remove()

            
            cedulacomprador5_entry.delete(0, END)
            con5_primerApellido_entry.delete(0, END)
            con5_segundoApellido_entry.delete(0, END)
            con5_primerNombre_entry.delete(0, END)
            con5_segundoNombre_entry.delete(0, END)
            con5_sexo_entry.delete(0, END)
            con5_fecha_nacimiento.delete(0, END)
            con5_rh_entry.delete(0, END)
            con5_domicilio_entry.delete(0, END)
            con5_estadocivil_entry.delete(0, END)

        def ingresarcomprador6():
            buscarcedulac6_combo.grid()
            cedulacomprador6_entry.grid()
            cedulacomprador6_entry.focus()
            con6_primerApellido_entry.grid()
            con6_segundoApellido_entry.grid()
            con6_primerNombre_entry.grid()
            con6_segundoNombre_entry.grid()
            con6_sexo_entry.grid()
            con6_fecha_nacimiento.grid()
            con6_rh_entry.grid()
            con6_ec_entry.grid()
            con6_domicilio_entry.grid()
            con6_estadocivil_entry.grid()
            boton_quitar_con6.grid()
            boton_agregar_con6.grid_remove()
            boton_agregar_con5.grid_remove()
            boton_quitar_con5.grid_remove()
            button11.grid_remove()
            button12.grid()

        def eliminarcomprador6():
            buscarcedulac6_combo.grid_remove()
            cedulacomprador6_entry.grid_remove()
            con6_primerApellido_entry.grid_remove()
            con6_segundoApellido_entry.grid_remove()
            con6_primerNombre_entry.grid_remove()
            con6_segundoNombre_entry.grid_remove()
            con6_sexo_entry.grid_remove()
            con6_fecha_nacimiento.grid_remove()
            con6_rh_entry.grid_remove()
            con6_ec_entry.grid_remove()
            con6_domicilio_entry.grid_remove()
            con6_estadocivil_entry.grid_remove()
            boton_quitar_con6.grid_remove()
            boton_agregar_con6.grid()
            boton_quitar_con5.grid()
            button11.grid()
            button12.grid_remove()
            
            cedulacomprador6_entry.delete(0, END)
            con6_primerApellido_entry.delete(0, END)
            con6_segundoApellido_entry.delete(0, END)
            con6_primerNombre_entry.delete(0, END)
            con6_segundoNombre_entry.delete(0, END)
            con6_sexo_entry.delete(0, END)
            con6_fecha_nacimiento.delete(0, END)
            con6_rh_entry.delete(0, END)
            con6_domicilio_entry.delete(0, END)
            con6_estadocivil_entry.delete(0, END)
        
        buscarcedulac2_combo=ttk.Combobox(frame4, width=12, values=opcionesu1)
        buscarcedulac2_combo.grid(row=2, column=0)
        buscarcedulac2_combo.grid_remove()


        buscarcedulav2_combo=ttk.Combobox(frame3, width=12, textvariable=buscarcedulav2, values=opcionesu1)
        buscarcedulav2_combo.grid(row=2, column=0)
        buscarcedulav2_combo.grid_remove()

        

        def actualizar_comboboxv2(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valor2 = buscarcedulav2_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valor2,))
            result1 = cursor.fetchone()
            if result1:

                
           
                # Actualizar el valor del Entry para la column 
                cedulavendedor2_entry.delete(0, END)
                cedulavendedor2_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                Ven2_primerApellido_entry.delete(0, END)
                Ven2_primerApellido_entry.insert(0, result1[1])

                Ven2_segundoApellido_entry.delete(0, END)
                Ven2_segundoApellido_entry.insert(0, result1[2])

                Ven2_primerNombre_entry.delete(0, END)
                Ven2_primerNombre_entry.insert(0, result1[3])

                Ven2_segundoNombre_entry.delete(0, END)
                Ven2_segundoNombre_entry.insert(0, result1[4])
                
                Ven2_sexo_entry.delete(0, END)
                Ven2_sexo_entry.insert(0, result1[5])

                Ven2_domicilio_entry.delete(0, END)
                Ven2_domicilio_entry.insert(0, result1[6])

                Ven2_estadocivil_entry.delete(0, END)
                Ven2_estadocivil_entry.insert(0, result1[7])

   

        # Vincular la función al evento de selección del combobox
        buscarcedulav2_combo.bind("<<ComboboxSelected>>", actualizar_comboboxv2)


       

        cedulavendedor2_entry=Entry(frame3, textvariable=cedulavendedor2)
        cedulavendedor2_entry.grid(row=2, column=1)
        cedulavendedor2_entry.grid_remove()
        

        Ven2_primerApellido_entry=Entry(frame3, textvariable=Ven2_primerApellido)
        Ven2_primerApellido_entry.grid(row=2, column=2)
        Ven2_primerApellido_entry.grid_remove()

        Ven2_segundoApellido_entry=Entry(frame3, textvariable=Ven2_segundoApellido)
        Ven2_segundoApellido_entry.grid(row=2, column=3)
        Ven2_segundoApellido_entry.grid_remove()

        Ven2_primerNombre_entry=Entry(frame3, textvariable=Ven2_primerNombre)
        Ven2_primerNombre_entry.grid(row=2, column=4)
        Ven2_primerNombre_entry.grid_remove()

        Ven2_segundoNombre_entry=Entry(frame3, textvariable=Ven2_segundoNombre)
        Ven2_segundoNombre_entry.grid(row=2, column=5)
        Ven2_segundoNombre_entry.grid_remove()


        


        Ven2_fecha_nacimiento=Entry(frame3, width=1, textvariable=Ven2_fechadenacimiento)
        Ven2_fecha_nacimiento.grid(row=2, column=7)
        Ven2_fecha_nacimiento.grid_remove()

        Ven2_rh_entry=Entry(frame3, width=1, textvariable=Ven2_rh)
        Ven2_rh_entry.grid(row=2, column=8)
        Ven2_rh_entry.grid_remove()

        Ven2_ec_entry=Entry(frame3, width=1)
        Ven2_ec_entry.grid(row=2, column=9)
        Ven2_ec_entry.grid_remove()
   

        
        Ven2_domicilio_entry=Entry(frame3, textvariable=Ven2_domicilio)
        Ven2_domicilio_entry.grid(row=2, column=10)
        Ven2_domicilio_entry.grid_remove()
        
        Ven2_sexo_entry=ttk.Combobox(frame3, textvariable=Ven2_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()))
        Ven2_sexo_entry.grid(row=2, column=6)
        Ven2_sexo_entry.grid_remove()

        Ven2_estadocivil_entry=ttk.Combobox(frame3, state="roadonly", textvariable=Ven2_estadocivil, width=30)
        Ven2_estadocivil_entry.grid(row=2, column=11)
        Ven2_estadocivil_entry.grid_remove()
        Ven2_sexo_entry.bind("<<ComboboxSelected>>", escoger_gv2)
        Ven2_sexo_entry.current(0)

        
               

        boton_agregar_ven2= ttk.Button(frame3, text="+", cursor="hand2", command=ingresarvendedor2)
        boton_agregar_ven2.grid(row=1, column=13)
        boton_agregar_ven2.configure(width=5)

        




        buscarcedulav3_combo=ttk.Combobox(frame3, width=12, textvariable=buscarcedulav3, values=opcionesu1) 
        buscarcedulav3_combo.grid(row=3, column=0)
        buscarcedulav3_combo.grid_remove()

        def actualizar_comboboxv3(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valor3 = buscarcedulav3_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valor3,))
            result1 = cursor.fetchone()
            if result1:

                # Actualizar el valor del Entry para la column 
                cedulavendedor3_entry.delete(0, END)
                cedulavendedor3_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                Ven3_primerApellido_entry.delete(0, END)
                Ven3_primerApellido_entry.insert(0, result1[1])

                Ven3_segundoApellido_entry.delete(0, END)
                Ven3_segundoApellido_entry.insert(0, result1[2])

                Ven3_primerNombre_entry.delete(0, END)
                Ven3_primerNombre_entry.insert(0, result1[3])

                Ven3_segundoNombre_entry.delete(0, END)
                Ven3_segundoNombre_entry.insert(0, result1[4])
                
                Ven3_sexo_entry.delete(0, END)
                Ven3_sexo_entry.insert(0, result1[5])

                Ven3_domicilio_entry.delete(0, END)
                Ven3_domicilio_entry.insert(0, result1[6])

                Ven3_estadocivil_entry.delete(0, END)
                Ven3_estadocivil_entry.insert(0, result1[7])

   

        # Vincular la función al evento de selección del combobox
        buscarcedulav3_combo.bind("<<ComboboxSelected>>", actualizar_comboboxv3)




        cedulavendedor3_entry=Entry(frame3, textvariable=cedulavendedor3)
        cedulavendedor3_entry.grid(row=3, column=1)
        cedulavendedor3_entry.grid_remove()
        

        Ven3_primerApellido_entry=Entry(frame3, textvariable=Ven3_primerApellido)
        Ven3_primerApellido_entry.grid(row=3, column=2)
        Ven3_primerApellido_entry.grid_remove()

        Ven3_segundoApellido_entry=Entry(frame3, textvariable=Ven3_segundoApellido)
        Ven3_segundoApellido_entry.grid(row=3, column=3)
        Ven3_segundoApellido_entry.grid_remove()

        Ven3_primerNombre_entry=Entry(frame3, textvariable=Ven3_primerNombre)
        Ven3_primerNombre_entry.grid(row=3, column=4)
        Ven3_primerNombre_entry.grid_remove()

        Ven3_segundoNombre_entry=Entry(frame3, textvariable=Ven3_segundoNombre)
        Ven3_segundoNombre_entry.grid(row=3, column=5)
        Ven3_segundoNombre_entry.grid_remove()


    


        Ven3_fecha_nacimiento=Entry(frame3, width=1, textvariable=ven3_fechadenacimiento)
        Ven3_fecha_nacimiento.grid(row=3, column=7)
        Ven3_fecha_nacimiento.grid_remove()

        Ven3_rh_entry=Entry(frame3, width=1, textvariable=Ven3_rh)
        Ven3_rh_entry.grid(row=3, column=8)
        Ven3_rh_entry.grid_remove()

        Ven3_ec_entry=Entry(frame3, width=1)
        Ven3_ec_entry.grid(row=3, column=9)
        Ven3_ec_entry.grid_remove()
            

        
        Ven3_domicilio_entry=Entry(frame3, textvariable=Ven3_domicilio)
        Ven3_domicilio_entry.grid(row=3, column=10)
        Ven3_domicilio_entry.grid_remove()
        
        Ven3_sexo_entry=ttk.Combobox(frame3, textvariable=Ven3_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()))
        Ven3_sexo_entry.grid(row=3, column=6)
        Ven3_sexo_entry.grid_remove()
        Ven3_estadocivil_entry=ttk.Combobox(frame3, state="roadonly", textvariable=Ven3_estadocivil, width=30)
        Ven3_estadocivil_entry.grid(row=3, column=11)
        Ven3_estadocivil_entry.grid_remove()
        Ven3_sexo_entry.bind("<<ComboboxSelected>>", escoger_gv3)
        Ven3_sexo_entry.current(0)
        Ven3_estadocivil_entry.grid_remove()
        

        
        
        boton_agregar_ven3= ttk.Button(frame3, text="+", cursor="hand2", command=ingresarvendedor3)
        boton_agregar_ven3.grid(row=2, column=13)
        boton_agregar_ven3.configure(width=5)
        boton_agregar_ven3.grid_remove()

        boton_quitar_ven2= ttk.Button(frame3, text="-", cursor="hand2", command=eliminarvendedor2)
        boton_quitar_ven2.grid(row=2, column=14)
        boton_quitar_ven2.configure(width=5)
        boton_quitar_ven2.grid_remove()

        buscarcedulav4_combo=ttk.Combobox(frame3, width=12, textvariable=buscarcedulav4, values=opcionesu1) 
        buscarcedulav4_combo.grid(row=4, column=0)
        buscarcedulav4_combo.grid_remove()

        def actualizar_comboboxv4(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valor4 = buscarcedulav4_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valor4,))
            result1 = cursor.fetchone()
            if result1:

                # Actualizar el valor del Entry para la column 
                cedulavendedor4_entry.delete(0, END)
                cedulavendedor4_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                Ven4_primerApellido_entry.delete(0, END)
                Ven4_primerApellido_entry.insert(0, result1[1])

                Ven4_segundoApellido_entry.delete(0, END)
                Ven4_segundoApellido_entry.insert(0, result1[2])

                Ven4_primerNombre_entry.delete(0, END)
                Ven4_primerNombre_entry.insert(0, result1[3])

                Ven4_segundoNombre_entry.delete(0, END)
                Ven4_segundoNombre_entry.insert(0, result1[4])
                
                Ven4_sexo_entry.delete(0, END)
                Ven4_sexo_entry.insert(0, result1[5])

                Ven4_domicilio_entry.delete(0, END)
                Ven4_domicilio_entry.insert(0, result1[6])

                Ven4_estadocivil_entry.delete(0, END)
                Ven4_estadocivil_entry.insert(0, result1[7])

   

        # Vincular la función al evento de selección del combobox
        buscarcedulav4_combo.bind("<<ComboboxSelected>>", actualizar_comboboxv4)


        cedulavendedor4_entry=Entry(frame3, textvariable=cedulavendedor4)
        cedulavendedor4_entry.grid(row=4, column=1)
        cedulavendedor4_entry.grid_remove()
        

        Ven4_primerApellido_entry=Entry(frame3, textvariable=Ven4_primerApellido)
        Ven4_primerApellido_entry.grid(row=4, column=2)
        Ven4_primerApellido_entry.grid_remove()

        Ven4_segundoApellido_entry=Entry(frame3, textvariable=Ven4_segundoApellido)
        Ven4_segundoApellido_entry.grid(row=4, column=3)
        Ven4_segundoApellido_entry.grid_remove()

        Ven4_primerNombre_entry=Entry(frame3, textvariable=Ven4_primerNombre)
        Ven4_primerNombre_entry.grid(row=4, column=4)
        Ven4_primerNombre_entry.grid_remove()

        Ven4_segundoNombre_entry=Entry(frame3, textvariable=Ven4_segundoNombre)
        Ven4_segundoNombre_entry.grid(row=4, column=5)
        Ven4_segundoNombre_entry.grid_remove()


        
        Ven4_fecha_nacimiento=Entry(frame3, width=1, textvariable=Ven4_fechanacimiento) 
        Ven4_fecha_nacimiento.grid(row=4, column=7)
        Ven4_fecha_nacimiento.grid_remove()

        Ven4_rh_entry=Entry(frame3, width=1, textvariable=Ven4_rh)
        Ven4_rh_entry.grid(row=4, column=8)
        Ven4_rh_entry.grid_remove()

        Ven4_ec_entry=Entry(frame3, width=1)
        Ven4_ec_entry.grid(row=4, column=9)
        Ven4_ec_entry.grid_remove()
            


        
        Ven4_domicilio_entry=Entry(frame3, textvariable=Ven4_domicilio)
        Ven4_domicilio_entry.grid(row=4, column=10)
        Ven4_domicilio_entry.grid_remove()
        

        Ven4_sexo_entry=ttk.Combobox(frame3, textvariable=Ven4_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()))
        Ven4_sexo_entry.grid(row=4, column=6)
        Ven4_sexo_entry.grid_remove()

        Ven4_estadocivil_entry=ttk.Combobox(frame3, state="roadonly", textvariable=Ven4_estadocivil, width=30)
        Ven4_estadocivil_entry.grid(row=4, column=11)
        Ven4_estadocivil_entry.grid_remove()
        Ven4_sexo_entry.bind("<<ComboboxSelected>>", escoger_gv4)
        Ven4_sexo_entry.current(0)
        
      

        boton_agregar_ven4= ttk.Button(frame3, text="+", cursor="hand2", command=ingresarvendedor4)
        boton_agregar_ven4.grid(row=3, column=13)
        boton_agregar_ven4.configure(width=5)
        boton_agregar_ven4.grid_remove()

        boton_quitar_ven3= ttk.Button(frame3, text="-", cursor="hand2", command=eliminarvendedor3)
        boton_quitar_ven3.grid(row=3, column=14)
        boton_quitar_ven3.configure(width=5)
        boton_quitar_ven3.grid_remove()


        buscarcedulav5_combo=ttk.Combobox(frame3, width=12, textvariable=buscarcedulav5, values=opcionesu1) 
        buscarcedulav5_combo.grid(row=5, column=0)
        buscarcedulav5_combo.grid_remove()

        def actualizar_comboboxv5(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valor5 = buscarcedulav5_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valor5,))
            result1 = cursor.fetchone()
            if result1:

                # Actualizar el valor del Entry para la column 
                cedulavendedor5_entry.delete(0, END)
                cedulavendedor5_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                Ven5_primerApellido_entry.delete(0, END)
                Ven5_primerApellido_entry.insert(0, result1[1])

                Ven5_segundoApellido_entry.delete(0, END)
                Ven5_segundoApellido_entry.insert(0, result1[2])

                Ven5_primerNombre_entry.delete(0, END)
                Ven5_primerNombre_entry.insert(0, result1[3])

                Ven5_segundoNombre_entry.delete(0, END)
                Ven5_segundoNombre_entry.insert(0, result1[4])
                
                Ven5_sexo_entry.delete(0, END)
                Ven5_sexo_entry.insert(0, result1[5])

                Ven5_domicilio_entry.delete(0, END)
                Ven5_domicilio_entry.insert(0, result1[6])

                Ven5_estadocivil_entry.delete(0, END)
                Ven5_estadocivil_entry.insert(0, result1[7])

   

        # Vincular la función al evento de selección del combobox
        buscarcedulav5_combo.bind("<<ComboboxSelected>>", actualizar_comboboxv5)



        cedulavendedor5_entry=Entry(frame3, textvariable=cedulavendedor5)
        cedulavendedor5_entry.grid(row=5, column=1)
        cedulavendedor5_entry.grid_remove()
        

        Ven5_primerApellido_entry=Entry(frame3, textvariable=Ven5_primerApellido)
        Ven5_primerApellido_entry.grid(row=5, column=2)
        Ven5_primerApellido_entry.grid_remove()

        Ven5_segundoApellido_entry=Entry(frame3, textvariable=Ven5_segundoApellido)
        Ven5_segundoApellido_entry.grid(row=5, column=3)
        Ven5_segundoApellido_entry.grid_remove()

        Ven5_primerNombre_entry=Entry(frame3, textvariable=Ven5_primerNombre)
        Ven5_primerNombre_entry.grid(row=5, column=4)
        Ven5_primerNombre_entry.grid_remove()

        Ven5_segundoNombre_entry=Entry(frame3, textvariable=Ven5_segundoNombre)
        Ven5_segundoNombre_entry.grid(row=5, column=5)
        Ven5_segundoNombre_entry.grid_remove()


    


        Ven5_fecha_nacimiento=Entry(frame3, width=1, textvariable=Ven5_fechanacimiento)
        Ven5_fecha_nacimiento.grid(row=5, column=7)
        Ven5_fecha_nacimiento.grid_remove()

        Ven5_rh_entry=Entry(frame3, width=1, textvariable=Ven5_rh)
        Ven5_rh_entry.grid(row=5, column=8)
        Ven5_rh_entry.grid_remove()
        
        Ven5_ec_entry=Entry(frame3, width=1)
        Ven5_ec_entry.grid(row=5, column=9)
        Ven5_ec_entry.grid_remove()



        
        Ven5_domicilio_entry=Entry(frame3, textvariable=Ven5_domicilio)
        Ven5_domicilio_entry.grid(row=5, column=10)
        Ven5_domicilio_entry.grid_remove()
        

        Ven5_sexo_entry=ttk.Combobox(frame3, textvariable=Ven5_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()))
        Ven5_sexo_entry.grid(row=5, column=6)
        Ven5_sexo_entry.grid_remove()



        Ven5_estadocivil_entry=ttk.Combobox(frame3, state="roadonly", textvariable=Ven5_estadocivil, width=30)
        Ven5_estadocivil_entry.grid(row=5, column=11)
        Ven5_estadocivil_entry.grid_remove()
        Ven5_sexo_entry.bind("<<ComboboxSelected>>", escoger_gv5)
        Ven5_sexo_entry.current(0)
        
        
        boton_agregar_ven5= ttk.Button(frame3, text="+", cursor="hand2", command=ingresarvendedor5)
        boton_agregar_ven5.grid(row=4, column=13)
        boton_agregar_ven5.configure(width=5)
        boton_agregar_ven5.grid_remove()

        
        boton_quitar_ven4= ttk.Button(frame3, text="-", cursor="hand2", command=eliminarvendedor4)
        boton_quitar_ven4.grid(row=4, column=14)
        boton_quitar_ven4.configure(width=5)
        boton_quitar_ven4.grid_remove()

        buscarcedulav6_combo=ttk.Combobox(frame3, width=12, textvariable=buscarcedulav6, values=opcionesu1) 
        buscarcedulav6_combo.grid(row=6, column=0)
        buscarcedulav6_combo.grid_remove()

        def actualizar_comboboxv6(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valor6 = buscarcedulav6_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valor6,))
            result1 = cursor.fetchone()
            if result1:

                # Actualizar el valor del Entry para la column 
                cedulavendedor6_entry.delete(0, END)
                cedulavendedor6_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                Ven6_primerApellido_entry.delete(0, END)
                Ven6_primerApellido_entry.insert(0, result1[1])

                Ven6_segundoApellido_entry.delete(0, END)
                Ven6_segundoApellido_entry.insert(0, result1[2])

                Ven6_primerNombre_entry.delete(0, END)
                Ven6_primerNombre_entry.insert(0, result1[3])

                Ven6_segundoNombre_entry.delete(0, END)
                Ven6_segundoNombre_entry.insert(0, result1[4])
                
                Ven6_sexo_entry.delete(0, END)
                Ven6_sexo_entry.insert(0, result1[5])

                Ven6_domicilio_entry.delete(0, END)
                Ven6_domicilio_entry.insert(0, result1[6])

                Ven6_estadocivil_entry.delete(0, END)
                Ven6_estadocivil_entry.insert(0, result1[7])

        buscarcedulav6_combo.bind("<<ComboboxSelected>>", actualizar_comboboxv6)


        cedulavendedor6_entry=Entry(frame3, textvariable=cedulavendedor6)
        cedulavendedor6_entry.grid(row=6, column=1)
        cedulavendedor6_entry.grid_remove()
        

        Ven6_primerApellido_entry=Entry(frame3, textvariable=Ven6_primerApellido)
        Ven6_primerApellido_entry.grid(row=6, column=2)
        Ven6_primerApellido_entry.grid_remove()

        Ven6_segundoApellido_entry=Entry(frame3, textvariable=Ven6_segundoApellido)
        Ven6_segundoApellido_entry.grid(row=6, column=3)
        Ven6_segundoApellido_entry.grid_remove()

        Ven6_primerNombre_entry=Entry(frame3, textvariable=Ven6_primerNombre)
        Ven6_primerNombre_entry.grid(row=6, column=4)
        Ven6_primerNombre_entry.grid_remove()

        Ven6_segundoNombre_entry=Entry(frame3, textvariable=Ven6_segundoNombre)
        Ven6_segundoNombre_entry.grid(row=6, column=5)
        Ven6_segundoNombre_entry.grid_remove()


    


        Ven6_fecha_nacimiento=Entry(frame3, width=1, textvariable=Ven6_fechanacimiento)
        Ven6_fecha_nacimiento.grid(row=6, column=7)
        Ven6_fecha_nacimiento.grid_remove()

        Ven6_rh_entry=Entry(frame3, width=1, textvariable=Ven6_rh)
        Ven6_rh_entry.grid(row=6, column=8)
        Ven6_rh_entry.grid_remove()

        Ven6_ec_entry=Entry(frame3, width=1)
        Ven6_ec_entry.grid(row=6, column=9)
        Ven6_ec_entry.grid_remove()
            

        
        Ven6_domicilio_entry=Entry(frame3, textvariable=Ven6_domicilio)
        Ven6_domicilio_entry.grid(row=6, column=10)
        Ven6_domicilio_entry.grid_remove()
        

        Ven6_sexo_entry=ttk.Combobox(frame3, textvariable=Ven6_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()))
        Ven6_sexo_entry.grid(row=6, column=6)
        Ven6_sexo_entry.grid_remove()

        Ven6_estadocivil_entry=ttk.Combobox(frame3, state="roadonly", textvariable=Ven6_estadocivil, width=30)
        Ven6_estadocivil_entry.grid(row=6, column=11)
        Ven6_estadocivil_entry.grid_remove()
        Ven6_sexo_entry.bind("<<ComboboxSelected>>", escoger_gv6)
        Ven6_sexo_entry.current(0)

        
        boton_agregar_ven6= ttk.Button(frame3, text="+", cursor="hand2", command=ingresarvendedor6)
        boton_agregar_ven6.grid(row=5, column=13)
        boton_agregar_ven6.configure(width=5)
        boton_agregar_ven6.grid_remove()



        boton_quitar_ven5= ttk.Button(frame3, text="-", cursor="hand2", command=eliminarvendedor5)
        boton_quitar_ven5.grid(row=5, column=14)
        boton_quitar_ven5.configure(width=5)
        boton_quitar_ven5.grid_remove()

        boton_quitar_ven6=ttk.Button(frame3, text="-", cursor="hand2", command=eliminarvendedor6)
        boton_quitar_ven6.grid(row=6, column=13)
        boton_quitar_ven6.configure(width=5)
        boton_quitar_ven6.grid_remove()









       

        def actualizar_comboboxc2(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valorc2 = buscarcedulac2_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valorc2,))
            result1 = cursor.fetchone()
            if result1:

              

                # Actualizar el valor del Entry para la column 
                cedulacomprador2_entry.delete(0, END)
                cedulacomprador2_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                con2_primerApellido_entry.delete(0, END)
                con2_primerApellido_entry.insert(0, result1[1])

                con2_segundoApellido_entry.delete(0, END)
                con2_segundoApellido_entry.insert(0, result1[2])

                con2_primerNombre_entry.delete(0, END)
                con2_primerNombre_entry.insert(0, result1[3])

                con2_segundoNombre_entry.delete(0, END)
                con2_segundoNombre_entry.insert(0, result1[4])
                
                con2_sexo_entry.delete(0, END)
                con2_sexo_entry.insert(0, result1[5])

                con2_domicilio_entry.delete(0, END)
                con2_domicilio_entry.insert(0, result1[6])

                con2_estadocivil_entry.delete(0, END)
                con2_estadocivil_entry.insert(0, result1[7])

   

        # Vincular la función al evento de selección del combobox
        buscarcedulac2_combo.bind("<<ComboboxSelected>>", actualizar_comboboxc2)

        cedulacomprador2_entry=Entry(frame4, textvariable=cedulacomprador2)
        cedulacomprador2_entry.grid(row=2, column=1)
        cedulacomprador2_entry.grid_remove()

        

        con2_primerApellido_entry=Entry(frame4, textvariable=con2_primerApellido)
        con2_primerApellido_entry.grid(row=2, column=2)
        con2_primerApellido_entry.grid_remove()

        con2_segundoApellido_entry=Entry(frame4, textvariable=con2_segundoApellido)
        con2_segundoApellido_entry.grid(row=2, column=3)
        con2_segundoApellido_entry.grid_remove()

        con2_primerNombre_entry=Entry(frame4, textvariable=con2_primerNombre)
        con2_primerNombre_entry.grid(row=2, column=4)
        con2_primerNombre_entry.grid_remove()

        con2_segundoNombre_entry=Entry(frame4, textvariable=con2_segundoNombre)
        con2_segundoNombre_entry.grid(row=2, column=5)
        con2_segundoNombre_entry.grid_remove()


        con2_fecha_nacimiento=Entry(frame4, width=1, textvariable=con2_fechanacimiento )
        con2_fecha_nacimiento.grid(row=2, column=7)
        con2_fecha_nacimiento.grid_remove()

        con2_rh_entry=Entry(frame4, width=1, textvariable= con2_rh)
        con2_rh_entry.grid(row=2, column=8)
        con2_rh_entry.grid_remove()

        con2_ec_entry=Entry(frame4, width=1)
        con2_ec_entry.grid(row=2, column=9)
        con2_ec_entry.grid_remove()
            

        
        con2_domicilio_entry=Entry(frame4, textvariable=con2_domicilio)
        con2_domicilio_entry.grid(row=2, column=10)
        con2_domicilio_entry.grid_remove()
        
        con2_sexo_entry=ttk.Combobox(frame4, textvariable=con2_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()) )
        con2_sexo_entry.grid(row=2, column=6)
        con2_sexo_entry.grid_remove()

        con2_estadocivil_entry=ttk.Combobox(frame4, state="roadonly", textvariable=con2_estadocivil, width=30)
        con2_estadocivil_entry.grid(row=2, column=11)
        con2_estadocivil_entry.grid_remove()
        con2_sexo_entry.bind("<<ComboboxSelected>>", escoger_gc2)
        con2_sexo_entry.current(0)
        
        
        

        boton_agregar_con2= ttk.Button(frame4, text="+", cursor="hand2", command=ingresarcomprador2)
        boton_agregar_con2.grid(row=1, column=13)
        boton_agregar_con2.configure(width=5)

        buscarcedulac3_combo=ttk.Combobox(frame4, width=12, values=opcionesu1)
        buscarcedulac3_combo.grid(row=3, column=0)
        buscarcedulac3_combo.grid_remove()

        def actualizar_comboboxc3(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valorc3 = buscarcedulac3_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valorc3,))
            result1 = cursor.fetchone()
            if result1:

              

                # Actualizar el valor del Entry para la column 
                cedulacomprador3_entry.delete(0, END)
                cedulacomprador3_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                con3_primerApellido_entry.delete(0, END)
                con3_primerApellido_entry.insert(0, result1[1])

                con3_segundoApellido_entry.delete(0, END)
                con3_segundoApellido_entry.insert(0, result1[2])

                con3_primerNombre_entry.delete(0, END)
                con3_primerNombre_entry.insert(0, result1[3])

                con3_segundoNombre_entry.delete(0, END)
                con3_segundoNombre_entry.insert(0, result1[4])
                
                con3_sexo_entry.delete(0, END)
                con3_sexo_entry.insert(0, result1[5])

                con3_domicilio_entry.delete(0, END)
                con3_domicilio_entry.insert(0, result1[6])

                con3_estadocivil_entry.delete(0, END)
                con3_estadocivil_entry.insert(0, result1[7])

   

        # Vincular la función al evento de selección del combobox
        buscarcedulac3_combo.bind("<<ComboboxSelected>>", actualizar_comboboxc3)

        cedulacomprador3_entry=Entry(frame4, textvariable=cedulacomprador3)
        cedulacomprador3_entry.grid(row=3, column=1)
        cedulacomprador3_entry.grid_remove()
        

        con3_primerApellido_entry=Entry(frame4, textvariable=con3_primerApellido)
        con3_primerApellido_entry.grid(row=3, column=2)
        con3_primerApellido_entry.grid_remove()

        con3_segundoApellido_entry=Entry(frame4, textvariable=con3_segundoApellido)
        con3_segundoApellido_entry.grid(row=3, column=3)
        con3_segundoApellido_entry.grid_remove()

        con3_primerNombre_entry=Entry(frame4, textvariable=con3_primerNombre)
        con3_primerNombre_entry.grid(row=3, column=4)
        con3_primerNombre_entry.grid_remove()

        con3_segundoNombre_entry=Entry(frame4, textvariable=con3_segundoNombre)
        con3_segundoNombre_entry.grid(row=3, column=5)
        con3_segundoNombre_entry.grid_remove()


        con3_fecha_nacimiento=Entry(frame4, width=1, textvariable=con3_fechanacimiento)
        con3_fecha_nacimiento.grid(row=3, column=7)
        con3_fecha_nacimiento.grid_remove()

        con3_rh_entry=Entry(frame4, width=1, textvariable=con3_rh)
        con3_rh_entry.grid(row=3, column=8 )
        con3_rh_entry.grid_remove()

        con3_ec_entry=Entry(frame4, width=1)
        con3_ec_entry.grid(row=3, column=9 )
        con3_ec_entry.grid_remove()
            
            

        
        con3_domicilio_entry=Entry(frame4, textvariable=con3_domicilio)
        con3_domicilio_entry.grid(row=3, column=10)
        con3_domicilio_entry.grid_remove()
        

        con3_sexo_entry=ttk.Combobox(frame4, textvariable=con3_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()) )
        con3_sexo_entry.grid(row=3, column=6)
        con3_sexo_entry.grid_remove()

        con3_estadocivil_entry=ttk.Combobox(frame4, state="roadonly", textvariable=con3_estadocivil, width=30)
        con3_estadocivil_entry.grid(row=3, column=11)
        con3_estadocivil_entry.grid_remove()
        con3_sexo_entry.bind("<<ComboboxSelected>>", escoger_gc3)
        con3_sexo_entry.current(0)
        
        
        

        boton_agregar_con3= ttk.Button(frame4, text="+", cursor="hand2", command=ingresarcomprador3)
        boton_agregar_con3.grid(row=2, column=13)
        boton_agregar_con3.configure(width=5)
        boton_agregar_con3.grid_remove()

        boton_quitar_con2= ttk.Button(frame4, text="-", cursor="hand2", command=eliminarcomprador2)
        boton_quitar_con2.grid(row=2, column=14)
        boton_quitar_con2.configure(width=5)
        boton_quitar_con2.grid_remove()


        buscarcedulac4_combo=ttk.Combobox(frame4, width=12, values=opcionesu1)
        buscarcedulac4_combo.grid(row=4, column=0)
        buscarcedulac4_combo.grid_remove()

        def actualizar_comboboxc4(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valorc4 = buscarcedulac4_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valorc4,))
            result1 = cursor.fetchone()
            if result1:

              

                # Actualizar el valor del Entry para la column 
                cedulacomprador4_entry.delete(0, END)
                cedulacomprador4_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                con4_primerApellido_entry.delete(0, END)
                con4_primerApellido_entry.insert(0, result1[1])

                con4_segundoApellido_entry.delete(0, END)
                con4_segundoApellido_entry.insert(0, result1[2])

                con4_primerNombre_entry.delete(0, END)
                con4_primerNombre_entry.insert(0, result1[3])

                con4_segundoNombre_entry.delete(0, END)
                con4_segundoNombre_entry.insert(0, result1[4])
                
                con4_sexo_entry.delete(0, END)
                con4_sexo_entry.insert(0, result1[5])

                con4_domicilio_entry.delete(0, END)
                con4_domicilio_entry.insert(0, result1[6])

                con4_estadocivil_entry.delete(0, END)
                con4_estadocivil_entry.insert(0, result1[7])
        
        buscarcedulac4_combo.bind("<<ComboboxSelected>>", actualizar_comboboxc4)

        cedulacomprador4_entry=Entry(frame4, textvariable=cedulacomprador4)
        cedulacomprador4_entry.grid(row=4, column=1)
        cedulacomprador4_entry.grid_remove()
        

        con4_primerApellido_entry=Entry(frame4, textvariable=con4_primerApellido)
        con4_primerApellido_entry.grid(row=4, column=2)
        con4_primerApellido_entry.grid_remove()

        con4_segundoApellido_entry=Entry(frame4, textvariable=con4_segundoApellido)
        con4_segundoApellido_entry.grid(row=4, column=3)
        con4_segundoApellido_entry.grid_remove()

        con4_primerNombre_entry=Entry(frame4, textvariable=con4_primerNombre)
        con4_primerNombre_entry.grid(row=4, column=4)
        con4_primerNombre_entry.grid_remove()

        con4_segundoNombre_entry=Entry(frame4, textvariable=con4_segundoNombre)
        con4_segundoNombre_entry.grid(row=4, column=5)
        con4_segundoNombre_entry.grid_remove()


        


        con4_fecha_nacimiento=Entry(frame4, width=1, textvariable=con4_fechanacimiento)
        con4_fecha_nacimiento.grid(row=4, column=7)
        con4_fecha_nacimiento.grid_remove()

        con4_rh_entry=Entry(frame4, width=1, textvariable=con4_rh)
        con4_rh_entry.grid(row=4, column=8)
        con4_rh_entry.grid_remove()

        con4_ec_entry=Entry(frame4, width=1)
        con4_ec_entry.grid(row=4, column=9)
        con4_ec_entry.grid_remove()
            
            

        
        con4_domicilio_entry=Entry(frame4, textvariable=con4_domicilio)
        con4_domicilio_entry.grid(row=4, column=10)
        con4_domicilio_entry.grid_remove()


        con4_sexo_entry=ttk.Combobox(frame4, textvariable=con4_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()) )
        con4_sexo_entry.grid(row=4, column=6)
        con4_sexo_entry.grid_remove()
        
        con4_estadocivil_entry=ttk.Combobox(frame4, state="roadonly", textvariable=con4_estadocivil, width=30)
        con4_estadocivil_entry.grid(row=4, column=11)
        con4_estadocivil_entry.grid_remove()
        con4_sexo_entry.bind("<<ComboboxSelected>>", escoger_gc4)
        con4_sexo_entry.current(0)
        
        

        boton_agregar_con4= ttk.Button(frame4, text="+", cursor="hand2", command=ingresarcomprador4)
        boton_agregar_con4.grid(row=3, column=13)
        boton_agregar_con4.configure(width=5)
        boton_agregar_con4.grid_remove()

        boton_quitar_con3= ttk.Button(frame4, text="-", cursor="hand2", command=eliminarcomprador3)
        boton_quitar_con3.grid(row=3, column=14)
        boton_quitar_con3.configure(width=5)
        boton_quitar_con3.grid_remove()

        buscarcedulac5_combo=ttk.Combobox(frame4, width=12, values=opcionesu1)
        buscarcedulac5_combo.grid(row=5, column=0)
        buscarcedulac5_combo.grid_remove()

        def actualizar_comboboxc5(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valorc5 = buscarcedulac5_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valorc5,))
            result1 = cursor.fetchone()
            if result1:

              

                # Actualizar el valor del Entry para la column 
                cedulacomprador5_entry.delete(0, END)
                cedulacomprador5_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                con5_primerApellido_entry.delete(0, END)
                con5_primerApellido_entry.insert(0, result1[1])

                con5_segundoApellido_entry.delete(0, END)
                con5_segundoApellido_entry.insert(0, result1[2])

                con5_primerNombre_entry.delete(0, END)
                con5_primerNombre_entry.insert(0, result1[3])

                con5_segundoNombre_entry.delete(0, END)
                con5_segundoNombre_entry.insert(0, result1[4])
                
                con5_sexo_entry.delete(0, END)
                con5_sexo_entry.insert(0, result1[5])

                con5_domicilio_entry.delete(0, END)
                con5_domicilio_entry.insert(0, result1[6])

                con5_estadocivil_entry.delete(0, END)
                con5_estadocivil_entry.insert(0, result1[7])

        buscarcedulac5_combo.bind("<<ComboboxSelected>>", actualizar_comboboxc5)

        cedulacomprador5_entry=Entry(frame4, textvariable=cedulacomprador5)
        cedulacomprador5_entry.grid(row=5, column=1)
        cedulacomprador5_entry.grid_remove()
        

        con5_primerApellido_entry=Entry(frame4, textvariable=con5_primerApellido)
        con5_primerApellido_entry.grid(row=5, column=2)
        con5_primerApellido_entry.grid_remove()

        con5_segundoApellido_entry=Entry(frame4, textvariable=con5_segundoApellido)
        con5_segundoApellido_entry.grid(row=5, column=3)
        con5_segundoApellido_entry.grid_remove()

        con5_primerNombre_entry=Entry(frame4, textvariable=con5_primerNombre)
        con5_primerNombre_entry.grid(row=5, column=4)
        con5_primerNombre_entry.grid_remove()

        con5_segundoNombre_entry=Entry(frame4, textvariable=con5_segundoNombre)
        con5_segundoNombre_entry.grid(row=5, column=5)
        con5_segundoNombre_entry.grid_remove()


        con5_fecha_nacimiento=Entry(frame4, width=1, textvariable=con5_fechanacimiento)
        con5_fecha_nacimiento.grid(row=5, column=7)
        con5_fecha_nacimiento.grid_remove()

        con5_rh_entry=Entry(frame4, width=1, textvariable=con5_rh)
        con5_rh_entry.grid(row=5, column=8)
        con5_rh_entry.grid_remove()

        con5_ec_entry=Entry(frame4, width=1 )
        con5_ec_entry.grid(row=5, column=9)
        con5_ec_entry.grid_remove()
            

        
        con5_domicilio_entry=Entry(frame4, textvariable=con5_domicilio)
        con5_domicilio_entry.grid(row=5, column=10)
        con5_domicilio_entry.grid_remove()
        

        con5_sexo_entry=ttk.Combobox(frame4, textvariable=con5_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()) )
        con5_sexo_entry.grid(row=5, column=6)
        con5_sexo_entry.grid_remove()
        
        con5_estadocivil_entry=ttk.Combobox(frame4, state="roadonly", textvariable=con5_estadocivil, width=30)
        con5_estadocivil_entry.grid(row=5, column=11)
        con5_estadocivil_entry.grid_remove()
        con5_sexo_entry.bind("<<ComboboxSelected>>", escoger_gc5)
        con5_sexo_entry.current(0)
        

        boton_agregar_con5= ttk.Button(frame4, text="+", cursor="hand2", command=ingresarcomprador5)
        boton_agregar_con5.grid(row=4, column=13)
        boton_agregar_con5.configure(width=5)
        boton_agregar_con5.grid_remove()

        boton_quitar_con4= ttk.Button(frame4, text="-", cursor="hand2", command=eliminarcomprador4)
        boton_quitar_con4.grid(row=4, column=14)
        boton_quitar_con4.configure(width=5)
        boton_quitar_con4.grid_remove()

        buscarcedulac6_combo=ttk.Combobox(frame4, width=12, values=opcionesu1)
        buscarcedulac6_combo.grid(row=6, column=0)
        buscarcedulac6_combo.grid_remove()

        def actualizar_comboboxc6(event):
            # Obtener el valor seleccionado en el combobox
            seleccionar_valorc5 = buscarcedulac5_combo.get()
            # Consulta para obtener los valores de las columnas correspondientes al valor seleccionado
            cursor.execute("SELECT Cedula, Primer_Apellido, Segundo_Apellido, Primer_Nombre, Segundo_Nombre, Sexo , Domicilio, Estado_Civil FROM Ingresar_usuario1  WHERE Cedula=?", (seleccionar_valorc5,))
            result1 = cursor.fetchone()
            if result1:

              

                # Actualizar el valor del Entry para la column 
                cedulacomprador6_entry.delete(0, END)
                cedulacomprador6_entry.insert(0, result1[0])
                # Actualizar el valor del Entry para la columna 
                con6_primerApellido_entry.delete(0, END)
                con6_primerApellido_entry.insert(0, result1[1])

                con6_segundoApellido_entry.delete(0, END)
                con6_segundoApellido_entry.insert(0, result1[2])

                con6_primerNombre_entry.delete(0, END)
                con6_primerNombre_entry.insert(0, result1[3])

                con6_segundoNombre_entry.delete(0, END)
                con6_segundoNombre_entry.insert(0, result1[4])
                
                con6_sexo_entry.delete(0, END)
                con6_sexo_entry.insert(0, result1[5])

                con6_domicilio_entry.delete(0, END)
                con6_domicilio_entry.insert(0, result1[6])

                con6_estadocivil_entry.delete(0, END)
                con6_estadocivil_entry.insert(0, result1[7])

        buscarcedulac6_combo.bind("<<ComboboxSelected>>", actualizar_comboboxc6)


        cedulacomprador6_entry=Entry(frame4, textvariable=cedulacomprador6)
        cedulacomprador6_entry.grid(row=6, column=1)
        cedulacomprador6_entry.grid_remove()
        

        con6_primerApellido_entry=Entry(frame4, textvariable=con6_primerApellido)
        con6_primerApellido_entry.grid(row=6, column=2)
        con6_primerApellido_entry.grid_remove()

        con6_segundoApellido_entry=Entry(frame4, textvariable=con6_segundoApellido)
        con6_segundoApellido_entry.grid(row=6, column=3)
        con6_segundoApellido_entry.grid_remove()

        con6_primerNombre_entry=Entry(frame4, textvariable=con6_primerNombre)
        con6_primerNombre_entry.grid(row=6, column=4)
        con6_primerNombre_entry.grid_remove()

        con6_segundoNombre_entry=Entry(frame4, textvariable=con6_segundoNombre)
        con6_segundoNombre_entry.grid(row=6, column=5)
        con6_segundoNombre_entry.grid_remove()

        con6_fecha_nacimiento=Entry(frame4, width=1, textvariable=con6_fechanacimiento)
        con6_fecha_nacimiento.grid(row=6, column=7)
        con6_fecha_nacimiento.grid_remove()

        con6_rh_entry=Entry(frame4, width=1, textvariable=con6_rh)
        con6_rh_entry.grid(row=6, column=8)
        con6_rh_entry.grid_remove()

        con6_ec_entry=Entry(frame4, width=1)
        con6_ec_entry.grid(row=6, column=9)
        con6_ec_entry.grid_remove()
            

        
        con6_domicilio_entry=Entry(frame4, textvariable=con6_domicilio)
        con6_domicilio_entry.grid(row=6, column=10)
        con6_domicilio_entry.grid_remove()
        

        con6_sexo_entry=ttk.Combobox(frame4, textvariable=con6_sexo, state="roadonly", width=5, values=tuple(opcionesV1.keys()) )
        con6_sexo_entry.grid(row=6, column=6)
        con6_sexo_entry.grid_remove()
        
        con6_estadocivil_entry=ttk.Combobox(frame4, state="roadonly", textvariable=con6_estadocivil, width=30)
        con6_estadocivil_entry.grid(row=6, column=11)
        con6_estadocivil_entry.grid_remove()
        con6_sexo_entry.bind("<<ComboboxSelected>>", escoger_gc6)
        con6_sexo_entry.current(0)
        

        boton_agregar_con6= ttk.Button(frame4, text="+", cursor="hand2", command=ingresarcomprador6)
        boton_agregar_con6.grid(row=5, column=13)
        boton_agregar_con6.configure(width=5)
        boton_agregar_con6.grid_remove()



        boton_quitar_con5= ttk.Button(frame4, text="-", cursor="hand2", command=eliminarcomprador5)
        boton_quitar_con5.grid(row=5, column=14)
        boton_quitar_con5.configure(width=5)
        boton_quitar_con5.grid_remove()

        boton_quitar_con6=ttk.Button(frame4, text="-", cursor="hand2", command=eliminarcomprador6)
        boton_quitar_con6.grid(row=6, column=13)
        boton_quitar_con6.configure(width=5)
        boton_quitar_con6.grid_remove()

        
        boton_volver= ttk.Button(ventana5, text="Volver", command=ocultar_ventana_secundaria, cursor="hand2")
        boton_volver.place(relx=0.35, rely=0.95)
            
        cerrar_ventana5=ttk.Button(ventana5, text="Cerrar", command=cerrar_ventanaregistro, cursor="hand2").place(relx=0.65, rely=0.95)
        
        
                           
        labelfecha=Label(ventana5, text="Dia", bg="deep sky blue", font=("Arial Black", 11))
        labelfecha.place(relx=0.84, rely=0.87)
        entry_fecha = tk.Entry(ventana5, width=5)
        entry_fecha.insert(0, fecha_actual1)
        entry_fecha.place(relx=0.84, rely=0.9)

        labelfecha2=Label(ventana5, text="Mes", bg="deep sky blue", font=("Arial Black", 11))
        labelfecha2.place(relx=0.87, rely=0.87)
        entry_fecha2 = tk.Entry(ventana5, width=5)
        entry_fecha2.insert(0, fecha_actual2)
        entry_fecha2.place(relx=0.87, rely=0.9)

        labelfecha3=Label(ventana5, text="Año", bg="deep sky blue", font=("Arial Black", 11))
        labelfecha3.place(relx=0.90, rely=0.87)
        entry_fecha3 = tk.Entry(ventana5, width=5)
        entry_fecha3.insert(0, fecha_actual3)
        entry_fecha3.place(relx=0.90, rely=0.9)

        labelnescrituras=Label(ventana5, text="# Escritura", bg="deep sky blue", font=("Arial Black", 11))
        labelnescrituras.place(relx=0.64, rely=0.86)
        entry_nescrituras= tk.Entry(ventana5, width=5, textvariable=nescrituras)
        entry_nescrituras.insert(0, numeroescr)
        entry_nescrituras.place(relx=0.66, rely=0.9)

        labelnumerodepaginas=Label(ventana5, text="# Hojas Notarial", bg="deep sky blue", font=("Arial Black", 11))
        labelnumerodepaginas.place(relx=0.42, rely=0.86)
        entry_paginas=tk.Entry(ventana5, width=25, textvariable=paginas)
        entry_paginas.place(relx=0.42, rely=0.9)

        labelnradicado=Label(ventana5, text="# Radicado", bg="deep sky blue", font=("Arial Black", 11))
        labelnradicado.place(relx=0.55, rely=0.86)
        entry_nradicado= tk.Entry(ventana5, width=20, textvariable=radicado)
        entry_nradicado.place(relx=0.55, rely=0.9)

        def insertar_usuario():
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()   
            valor_cedula1= cedulavendedor_entry.get()
                       
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula1, ))
            data = mcursor.fetchone()
            if data is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( primerApellido1_entry.get().upper().strip(), segundoApellido1_entry.get().upper().strip(), primerNombre1_entry.get().upper().strip(), segundoNombre1_entry.get().upper().strip(), sexo1_entry.get().strip(), domicilio1_entry.get().strip(), estadocivil1_entry.get().strip(), cedulavendedor_entry.get() ))
                messagebox.showinfo("Usuario registrado", "El usuario se ha actualizado exitosamente")
                bd.commit()
                
            else:
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulavendedor_entry.get(), primerApellido1_entry.get().upper().strip(), segundoApellido1_entry.get().upper().strip(), primerNombre1_entry.get().upper().strip(), segundoNombre1_entry.get().upper().strip(), sexo1_entry.get().strip(), domicilio1_entry.get().strip(), estadocivil1_entry.get().strip() ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido agregado exitosamente")
                            
                    
                bd.commit()
                bd.close()
        button1=ttk.Button(frame3, width=8, text="Agregar", command=insertar_usuario)
        button1.grid(row=1, column=12)

        def insertar_usuario2():
            
            valor_cedula2= cedulavendedor2_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula2, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( Ven2_primerApellido_entry.get().upper().strip(), Ven2_segundoApellido_entry.get().upper().strip(), Ven2_primerNombre_entry.get().upper().strip(), Ven2_segundoNombre_entry.get().upper().strip(), Ven2_sexo_entry.get().strip(), Ven2_domicilio_entry.get().strip(), Ven2_estadocivil_entry.get().strip(), cedulavendedor2_entry.get() ))
                messagebox.showinfo("Usuario registrado", "El usuario se ha actualizado exitosamente")
                bd.commit()
                
            else:
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulavendedor2_entry.get().strip(), Ven2_primerApellido_entry.get().upper().strip(), Ven2_segundoApellido_entry.get().upper().strip(), Ven2_primerNombre_entry.get().upper().strip(), Ven2_segundoNombre_entry.get().upper().strip(), Ven2_sexo_entry.get().strip(), Ven2_domicilio_entry.get().strip(), Ven2_estadocivil_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()
        button2=ttk.Button(frame3, width=8, text="Agregar", command=insertar_usuario2 )
        button2.grid(row=2, column=12)
        button2.grid_remove()

        def insertar_usuario3():
            
            valor_cedula3= cedulavendedor3_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula3, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( Ven3_primerApellido_entry.get().upper().strip(), Ven3_segundoApellido_entry.get().upper().strip(), Ven3_primerNombre_entry.get().upper().strip(), Ven3_segundoNombre_entry.get().upper().strip(), Ven3_sexo_entry.get().strip(), Ven3_domicilio_entry.get().strip(), Ven3_estadocivil_entry.get().strip(), cedulavendedor3_entry.get() ))
                messagebox.showinfo("Usuario registrado", "El usuario se ha actualizado exitosamente")
                bd.commit()
            else:
               
                
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulavendedor3_entry.get().strip(),  Ven3_primerApellido_entry.get().upper().strip(), Ven3_segundoApellido_entry.get().upper().strip(), Ven3_primerNombre_entry.get().upper().strip(), Ven3_segundoNombre_entry.get().upper().strip(), Ven3_sexo_entry.get().strip(), Ven3_domicilio_entry.get().strip(), Ven3_estadocivil_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()
        button3=ttk.Button(frame3, width=8, text="Agregar", command=insertar_usuario3 )
        button3.grid(row=3, column=12)
        button3.grid_remove()

        def insertar_usuario4():
            
            valor_cedula4= cedulavendedor4_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula4, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( Ven4_primerApellido_entry.get().upper().strip(), Ven4_segundoApellido_entry.get().upper().strip(), Ven4_primerNombre_entry.get().upper().strip(), Ven4_segundoNombre_entry.get().upper().strip(), Ven4_sexo_entry.get().strip(), Ven4_domicilio_entry.get().strip(), Ven4_estadocivil_entry.get().strip(), cedulavendedor4_entry.get() ))
                messagebox.showinfo("Usuario registrado", "El usuario se ha actualizado exitosamente")
                bd.commit()
            else:
                                
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulavendedor4_entry.get().strip(),  Ven4_primerApellido_entry.get().upper().strip(), Ven4_segundoApellido_entry.get().upper().strip(), Ven4_primerNombre_entry.get().upper().strip(), Ven4_segundoNombre_entry.get().upper().strip(), Ven4_sexo_entry.get().strip(), Ven4_domicilio_entry.get().strip(), Ven4_estadocivil_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()
        button4=ttk.Button(frame3, width=8, text="Agregar", command=insertar_usuario4 )
        button4.grid(row=4, column=12)
        button4.grid_remove()

        def insertar_usuario5():
            
            valor_cedula5= cedulavendedor5_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula5, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( Ven5_primerApellido_entry.get().upper().strip(), Ven5_segundoApellido_entry.get().upper().strip(), Ven5_primerNombre_entry.get().upper().strip(), Ven5_segundoNombre_entry.get().upper().strip(), Ven5_sexo_entry.get().strip(), Ven5_domicilio_entry.get().strip(), Ven5_estadocivil_entry.get().strip(), cedulavendedor5_entry.get() ))
                messagebox.showinfo("Usuario registrado", "El usuario se ha actualizado exitosamente")
                bd.commit()
            else:
                                
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulavendedor5_entry.get().strip(),  Ven5_primerApellido_entry.get().upper().strip(), Ven5_segundoApellido_entry.get().upper().strip(), Ven5_primerNombre_entry.get().upper().strip(), Ven5_segundoNombre_entry.get().upper().strip(), Ven5_sexo_entry.get().strip(), Ven5_domicilio_entry.get().strip(), Ven5_estadocivil_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()
        button5=ttk.Button(frame3, width=8, text="Agregar", command=insertar_usuario5 )
        button5.grid(row=5, column=12)
        button5.grid_remove()

        def insertar_usuario6():
            
            valor_cedula6= cedulavendedor6_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula6, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( Ven6_primerApellido_entry.get().upper().strip(), Ven6_segundoApellido_entry.get().upper().strip(), Ven6_primerNombre_entry.get().upper().strip(), Ven6_segundoNombre_entry.get().upper().strip(), Ven6_sexo_entry.get().strip(), Ven6_domicilio_entry.get().strip(), Ven6_estadocivil_entry.get().strip(), cedulavendedor6_entry.get() ))
                messagebox.showinfo("Usuario registrado", "El usuario se ha actualizado exitosamente")
                bd.commit()
            else:
                               
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulavendedor6_entry.get().strip(),  Ven6_primerApellido_entry.get().upper().strip(), Ven6_segundoApellido_entry.get().upper().strip(), Ven6_primerNombre_entry.get().upper().strip(), Ven6_segundoNombre_entry.get().upper().strip(), Ven6_sexo_entry.get().strip(), Ven6_domicilio_entry.get().strip(), Ven6_estadocivil_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()
        button6=ttk.Button(frame3, width=8, text="Agregar", command=insertar_usuario6)
        button6.grid(row=6, column=12)
        button6.grid_remove()

        def insertar_usuario7():
                        
            valor_cedula7= cedulacomprador_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula7, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( primer_apellido2_entry.get().upper().strip(), segundo_apellido2_entry.get().upper().strip(), primer_nombre2_entry.get().upper().strip(), segundo_nombre2_entry.get().upper().strip(), sexo2_entry.get().strip(), domicilio2_entry.get().strip(), estadocivil2_entry.get().strip(), cedulacomprador_entry.get().strip() ))
                messagebox.showinfo("Usuario registrado", "El usuario se ha actualizado exitosamente")
                bd.commit()
            else:
                
               
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulacomprador_entry.get().strip(),  primer_apellido2_entry.get().upper().strip(), segundo_apellido2_entry.get().upper().strip(), primer_nombre2_entry.get().upper().strip(), segundo_nombre2_entry.get().upper().strip(), sexo2_entry.get().strip(), domicilio2_entry.get().strip(), estadocivil2_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()

        button7=ttk.Button(frame4, width=8, text="Agregar", command=insertar_usuario7)
        button7.grid(row=1, column=12)
              
       

        def insertar_usuario8():

                    
            valor_cedula8= cedulacomprador2_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula8, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( con2_primerApellido_entry.get().upper().strip(), con2_segundoApellido_entry.get().upper().strip(), con2_primerNombre_entry.get().upper().strip(), con2_segundoNombre_entry.get().upper().strip(), con2_sexo_entry.get().strip(), con2_domicilio_entry.get().strip(), con2_estadocivil_entry.get().strip(), cedulacomprador2_entry.get().strip() ))
                messagebox.showinfo("Usuario registrado", "El usuario se ha actualizado exitosamente")
                bd.commit()
            else:
                       
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulacomprador2_entry.get().strip(),  con2_primerApellido_entry.get().upper().strip(), con2_segundoApellido_entry.get().upper().strip(), con2_primerNombre_entry.get().upper().strip(), con2_segundoNombre_entry.get().upper().strip(), con2_sexo_entry.get().strip(), con2_domicilio_entry.get().strip(), con2_estadocivil_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "el usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()

        button8=ttk.Button(frame4, width=8, text="Agregar", command=insertar_usuario8)
        button8.grid(row=2, column=12)
        button8.grid_remove()

        def insertar_usuario9():

                    
            valor_cedula9= cedulacomprador3_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula9, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( con3_primerApellido_entry.get().upper().strip(), con3_segundoApellido_entry.get().upper().strip(), con3_primerNombre_entry.get().upper().strip(), con3_segundoNombre_entry.get().upper().strip(), con3_sexo_entry.get().strip(), con3_domicilio_entry.get().strip(), con3_estadocivil_entry.get().strip(), cedulacomprador3_entry.get().strip()))
                messagebox.showinfo("Usuario Registrado", "El usuario ha sido actualizado exitosamente")
                bd.commit()
            else:
                
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulacomprador3_entry.get().strip(),  con3_primerApellido_entry.get().upper().strip(), con3_segundoApellido_entry.get().upper().strip(), con3_primerNombre_entry.get().upper().strip(), con3_segundoNombre_entry.get().upper().strip(), con3_sexo_entry.get().strip(), con3_domicilio_entry.get().strip(), con3_estadocivil_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()
        button9=ttk.Button(frame4, width=8, text="Agregar", command=insertar_usuario9)
        button9.grid(row=3, column=12)
        button9.grid_remove()

        def insertar_usuario10():

                    
            valor_cedula10= cedulacomprador4_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula10, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( con4_primerApellido_entry.get().upper().strip(), con4_segundoApellido_entry.get().upper().strip(), con4_primerNombre_entry.get().upper().strip(), con4_segundoNombre_entry.get().upper().strip(), con4_sexo_entry.get().strip(), con4_domicilio_entry.get().strip(), con4_estadocivil_entry.get().strip(), cedulacomprador4_entry.get().strip()))
                messagebox.showinfo("Usuario registrado", "El usuario ha sido guardado exitosamente")
                bd.commit()
            else:
                               
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulacomprador4_entry.get().strip(),  con4_primerApellido_entry.get().upper().strip(), con4_segundoApellido_entry.get().upper().strip(), con4_primerNombre_entry.get().upper().strip(), con4_segundoNombre_entry.get().upper().strip(), con4_sexo_entry.get().strip(), con4_domicilio_entry.get().strip(), con4_estadocivil_entry.get().strip()  ))
                messagebox.showinfo("Usuario nuevo", "El usuario ha sido guardado exitosamente")
                bd.commit()
                bd.close()
        button10=ttk.Button(frame4, width=8, text="Agregar", command=insertar_usuario10)
        button10.grid(row=4, column=12)
        button10.grid_remove()

        def insertar_usuario11():

                    
            valor_cedula11= cedulacomprador5_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula11, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( con5_primerApellido_entry.get().upper().strip(), con5_segundoApellido_entry.get().upper().strip(), con5_primerNombre_entry.get().upper().strip(), con5_segundoNombre_entry.get().upper().strip(), con5_sexo_entry.get().strip(), con5_domicilio_entry.get().strip(), con5_estadocivil_entry.get().strip(), cedulacomprador5_entry.get().strip()))
                messagebox.showinfo("Usuario registrado", "El usuario ha sido guardado exitosamente")
            else:
                messagebox.showinfo("Usuario nuevo", "el usuario ha sido guardado exitosamente")
                bd = sqlite3.connect("login1.db")
                mcursor = bd.cursor()
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulacomprador5_entry.get().strip(),  con5_primerApellido_entry.get().upper().strip(), con5_segundoApellido_entry.get().upper().strip(), con5_primerNombre_entry.get().upper().strip(), con5_segundoNombre_entry.get().upper().strip(), con5_sexo_entry.get().strip(), con5_domicilio_entry.get().strip(), con5_estadocivil_entry.get().strip()  ))
                bd.commit()
                bd.close()

        button11=ttk.Button(frame4, width=8, text="Agregar", command=insertar_usuario11)
        button11.grid(row=5, column=12)
        button11.grid_remove()


        def insertar_usuario12():

                    
            valor_cedula12= cedulacomprador6_entry.get()
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            #mcursor.execute("CREATE TABLE Ingresar_usuario1 (Cedula TEXT PRIMARY KEY,  Primer_Apellido TEXT, Segundo_Apellido TEXT, Primer_Nombre TEXT, Segundo_nombre TEXT, Sexo TEXT, Domicilio TEXT, Estado_Civil TEXT)")
            mcursor.execute("SELECT * FROM Ingresar_usuario1 WHERE Cedula = ?", (valor_cedula12, ))
            if mcursor.fetchone() is not None:
                mcursor.execute("UPDATE Ingresar_usuario1 SET Primer_Apellido=?, Segundo_Apellido=?, Primer_Nombre=?, Segundo_nombre=?, Sexo=?, Domicilio=?, Estado_Civil=? WHERE Cedula=?", ( con6_primerApellido_entry.get().upper().strip(), con6_segundoApellido_entry.get().upper().strip(), con6_primerNombre_entry.get().upper().strip(), con6_segundoNombre_entry.get().upper().strip(), con6_sexo_entry.get().strip(), con6_domicilio_entry.get().strip(), con6_estadocivil_entry.get().strip(), cedulacomprador6_entry.get().strip()))
                messagebox.showinfo("Usuario registrado", "El usuario ha sido guardado exitosamente")
            else:
                
                mcursor.execute("INSERT INTO Ingresar_usuario1 VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (cedulacomprador6_entry.get().strip(),  con6_primerApellido_entry.get().upper().strip(), con6_segundoApellido_entry.get().upper().strip(), con6_primerNombre_entry.get().upper().strip(), con6_segundoNombre_entry.get().upper().strip(), con6_sexo_entry.get().strip(), con6_domicilio_entry.get().strip(), con6_estadocivil_entry.get().strip()  ))
                bd.commit()
                bd.close()

        button12=ttk.Button(frame4, width=8, text="Agregar", command=insertar_usuario12)
        button12.grid(row=6, column=12)
        button12.grid_remove()


        def insertardatos():
            global boton_insertar_datos
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()
            


        
            
            
            valorLetras= int(precioNuevo_entry.get())
            añoLetras= int (año_comb.get())
            escrLetras= int (escritura.get())
            precmil=float (int (precioNuevo_entry.get()))
            mesletras= int (entry_fecha2.get())
        
            año_letras= int(entry_fecha3.get())
            dia_letra= int(entry_fecha.get())

            mes_LETRAS = int (entry_fecha2.get())
            numero_escrituras= int (entry_nescrituras.get())
            valor_entry_chip= entry_chip.get()
            valor_entry_chip= "" if valor_entry_chip=="None" else valor_entry_chip

            #cursor.execute("CREATE TABLE datos_escrituras1 (ID INTEGER PRIMARY KEY AUTOINCREMENT, Matricula Text, Cedula_catastral TEXT, precio TEXT, precio1 Text, precioLetras TEXT, Ubicacion TEXT, Rural_urbano TEXT, Departamento TEXT, Municipio TEXT, Direccion Text, Modo_Adquirir TEXT, Escritura TEXT, EscrLetras TEXT,  Dia TEXT, Mes Text, Año TEXT, AñoLetras TEXT, Notaria TEXT, Municipio_notaria TEXT, Notario TEXT, cedulavendedor1 TEXT, primerApellidoVen1 TEXT, segundoApellidoVen1 TEXT, primerNombreVen1 TEXT, segundoNombre_Ven1 TEXT, sexoVen1 TEXT, fecha_nacimiento_Ven1 TEXT, RH_Ven1 TEXT, domicilio_Ven1 TEXT, estadocivil_ven1 TEXT, cedulavendedor2 TEXT, primerApellidoVen2 TEXT, segundoApellidoVen2 TEXT, primerNombreVen2 TEXT, segundoNombre_Ven2 TEXT, sexoVen2 TEXT, fecha_nacimiento_Ven2 TEXT, RH_Ven2 TEXT, domicilio_Ven2 TEXT, estadocivil_ven2 TEXT, cedulavendedor3 TEXT, primerApellidoVen3 TEXT, segundoApellidoVen3 TEXT, primerNombreVen3 TEXT, segundoNombre_Ven3 TEXT, sexoVen3 TEXT, fecha_nacimiento_Ven3 TEXT, RH_Ven3 TEXT, domicilio_Ven3 TEXT, estadocivil_ven3 TEXT, cedulavendedor4 TEXT, primerApellidoVen4 TEXT, segundoApellidoVen4 TEXT, primerNombreVen4 TEXT, segundoNombre_Ven4 TEXT, sexoVen4 TEXT, fecha_nacimiento_Ven4 TEXT, RH_Ven4 TEXT, domicilio_Ven4 TEXT, estadocivil_ven4 TEXT, cedulavendedor5 TEXT, primerApellidoVen5 TEXT, segundoApellidoVen5 TEXT, primerNombreVen5 TEXT, segundoNombre_Ven5 TEXT, sexoVen5 TEXT, fecha_nacimiento_Ven5 TEXT, RH_Ven5 TEXT, domicilio_Ven5 TEXT, estadocivil_ven5 TEXT, cedulavendedor6 TEXT, primerApellidoVen6 TEXT, segundoApellidoVen6 TEXT, primerNombreVen6 TEXT, segundoNombre_Ven6 TEXT, sexoVen6 TEXT, fecha_nacimiento_Ven6 TEXT, RH_Ven6 TEXT, domicilio_Ven6 TEXT, estadocivil_ven6 TEXT, cedula_comprador1 TEXT, primerApellido_comp1 TEXT, segundoApellido_comp1 TEXT, primerNombre_comp1 TEXT, segundoNombre_comp1 TEXT, sexo_comp1 TEXT, fecha_nacimiento_comp1 TEXT, RH_comp1 TEXT  ,domicilio_comp1 TEXT, estadocivil_comp1 TEXT, cedula_comprador2 TEXT, primerApellido_comp2 TEXT, segundoApellido_comp2 TEXT, primerNombre_comp2 TEXT, segundoNombre_comp2 TEXT, sexo_comp2 TEXT, fecha_nacimiento_comp2 TEXT, RH_comp2 TEXT  ,domicilio_comp2 TEXT, estadocivil_comp2 TEXT, cedula_comprador3 TEXT, primerApellido_comp3 TEXT, segundoApellido_comp3 TEXT, primerNombre_comp3 TEXT, segundoNombre_comp3 TEXT, sexo_comp3 TEXT, fecha_nacimiento_comp3 TEXT, RH_comp3 TEXT  ,domicilio_comp3 TEXT, estadocivil_comp3 TEXT, cedula_comprador4 TEXT, primerApellido_comp4 TEXT, segundoApellido_comp4 TEXT, primerNombre_comp4 TEXT, segundoNombre_comp4 TEXT, sexo_comp4 TEXT, fecha_nacimiento_comp4 TEXT, RH_comp4 TEXT, domicilio_comp4 TEXT, estadocivil_comp4 TEXT, cedula_comprador5 TEXT, primerApellido_comp5 TEXT, segundoApellido_comp5 TEXT, primerNombre_comp5 TEXT, segundoNombre_comp5 TEXT, sexo_comp5 TEXT, fecha_nacimiento_comp5 TEXT, RH_comp5 TEXT, domicilio_comp5 TEXT, estadocivil_comp5 TEXT, cedula_comprador6 TEXT, primerApellido_comp6 TEXT, segundoApellido_comp6 TEXT, primerNombre_comp6 TEXT, segundoNombre_comp6 TEXT, sexo_comp6 TEXT, fecha_nacimiento_comp6 TEXT, RH_comp6 TEXT, domicilio_comp6 TEXT, estadocivil_comp6 TEXT)")
            mcursor.execute("INSERT INTO datos_escrituras1  VALUES (NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", (matriculaNuevo_entry.get(), cedulaNuevo_entry.get(), ('{:,.2f}'.format(precmil).replace(',','~').replace('.',',').replace('~','.')), str(TextNumber(valorLetras)).strip(), ubicacion.get(), rural_urbano.get(),  departamento.get().upper(), municipio.get().upper(), direccionNueva_entry.get(), adquirir.get(), escritura.get(), str(TextNumber(escrLetras)).strip(), dia_comb.get(), mes_comb.get(), año_comb.get(), str(TextNumber(añoLetras)).strip(), notaria.get().upper(), municipioNota.get().upper(), notario.get().upper(), cedulavendedor_entry.get().strip(),  primerApellido1_entry.get().upper().strip(), segundoApellido1_entry.get().upper().strip(), primerNombre1_entry.get().upper().strip(), segundoNombre1_entry.get().upper().strip(), sexo1_entry.get().strip(), fecha_nacimiento1.get().strip(), rh1_entry.get().strip(), domicilio1_entry.get().strip(), estadocivil1_entry.get().strip(), cedulavendedor2_entry.get().strip(), Ven2_primerApellido_entry.get().upper().strip(), Ven2_segundoApellido_entry.get().upper().strip(), Ven2_primerNombre_entry.get().upper().strip(), Ven2_segundoNombre_entry.get().upper().strip(), Ven2_sexo_entry.get().strip(), Ven2_fecha_nacimiento.get().strip(), Ven2_rh_entry.get().strip(),  Ven2_domicilio_entry.get().strip(), Ven2_estadocivil_entry.get().strip(), cedulavendedor3_entry.get().strip(), Ven3_primerApellido_entry.get().upper().strip(), Ven3_segundoApellido_entry.get().upper().strip(),  Ven3_primerNombre_entry.get().upper().strip(), Ven3_segundoNombre_entry.get().upper().strip(),  Ven3_sexo_entry.get().strip(), Ven3_fecha_nacimiento.get().strip(), Ven3_rh_entry.get().strip(),   Ven3_domicilio_entry.get().strip(), Ven3_estadocivil_entry.get().strip(), cedulavendedor4_entry.get().strip(),  Ven4_primerApellido_entry.get().upper().strip(),  Ven4_segundoApellido_entry.get().upper().strip(), Ven4_primerNombre_entry.get().upper().strip(),  Ven4_segundoNombre_entry.get().upper().strip(), Ven4_sexo_entry.get().strip(), Ven4_fecha_nacimiento.get().strip(), Ven4_rh_entry.get().strip(), Ven4_domicilio_entry.get().strip(), Ven4_estadocivil_entry.get().strip(),  cedulavendedor5_entry.get().strip(), Ven5_primerApellido_entry.get().upper().strip(),  Ven5_segundoApellido_entry.get().upper().strip(), Ven5_primerNombre_entry.get().upper().strip(),  Ven5_segundoNombre_entry.get().upper().strip(), Ven5_sexo_entry.get().strip(), Ven5_fecha_nacimiento.get().strip(), Ven5_rh_entry.get().strip(),  Ven5_domicilio_entry.get().strip(), Ven5_estadocivil_entry.get().strip(), cedulavendedor6_entry.get().strip(), Ven6_primerApellido_entry.get().upper().strip(),   Ven6_segundoApellido_entry.get().upper().strip(), Ven6_primerNombre_entry.get().upper().strip(),  Ven6_segundoNombre_entry.get().upper().strip(), Ven6_sexo_entry.get().strip(),  Ven6_fecha_nacimiento.get().strip(), Ven6_rh_entry.get().strip(),  Ven6_domicilio_entry.get().strip(), Ven6_estadocivil_entry.get().strip(), cedulacomprador_entry.get().strip(), primer_apellido2_entry.get().upper().strip(), segundo_apellido2_entry.get().upper().strip(), primer_nombre2_entry.get().upper().strip(), segundo_nombre2_entry.get().upper().strip(), sexo2_entry.get().strip(), fecha_nacimiento2.get().strip(), rh2_entry.get().strip(), domicilio2_entry.get().strip(), estadocivil2_entry.get().strip(),  cedulacomprador2_entry.get().strip(), con2_primerApellido_entry.get().upper().strip(),  con2_segundoApellido_entry.get().upper().strip(), con2_primerNombre_entry.get().upper().strip(), con2_segundoNombre_entry.get().upper().strip(), con2_sexo_entry.get().strip(),  con2_fecha_nacimiento.get().strip(), con2_rh_entry.get().strip(), con2_domicilio_entry.get().strip(), con2_estadocivil_entry.get().strip(), cedulacomprador3_entry.get().strip(),  con3_primerApellido_entry.get().upper().strip(), con3_segundoApellido_entry.get().upper().strip(), con3_primerNombre_entry.get().upper().strip(), con3_segundoNombre_entry.get().upper().strip(), con3_sexo_entry.get().strip(), con3_fecha_nacimiento.get().strip(), con3_rh_entry.get().strip(), con3_domicilio_entry.get().strip(), con3_estadocivil_entry.get().strip(),  cedulacomprador4_entry.get().strip(), con4_primerApellido_entry.get().upper().strip(),  con4_segundoApellido_entry.get().upper().strip(), con4_primerNombre_entry.get().upper().strip(),  con4_segundoNombre_entry.get().upper().strip(), con4_sexo_entry.get().strip(), con4_fecha_nacimiento.get().strip(), con4_rh_entry.get().strip(), con4_domicilio_entry.get().strip(), con4_estadocivil_entry.get().strip(), cedulacomprador5_entry.get().strip(), con5_primerApellido_entry.get().upper().strip(), con5_segundoApellido_entry.get().upper().strip(), con5_primerNombre_entry.get().upper().strip(), con5_segundoNombre_entry.get().upper().strip(), con5_sexo_entry.get().strip(),  con5_fecha_nacimiento.get().strip(), con5_rh_entry.get().strip(), con5_domicilio_entry.get().strip(), con5_estadocivil_entry.get().strip(),  cedulacomprador6_entry.get().strip(), con6_primerApellido_entry.get().upper().strip(), con6_segundoApellido_entry.get().upper().strip(), con6_primerNombre_entry.get().upper().strip(), con6_segundoNombre_entry.get().upper().strip(), con6_sexo_entry.get().strip(),  con6_fecha_nacimiento.get().strip(), con6_rh_entry.get().strip(),  con6_domicilio_entry.get().strip(), con6_estadocivil_entry.get().strip(), precioNuevo_entry.get(), entry_fecha.get(), entry_fecha2.get(), entry_fecha3.get(), str(TextNumber(año_letras)).strip().upper(), str(TextNumber(dia_letra)).strip().upper(), str(calendar.month_name[mes_LETRAS]).strip().upper(), entry_nescrituras.get(), str(TextNumber(numero_escrituras)).strip().upper(), entry_nradicado.get(), linderomatricula.get().rstrip(""), pazysalvomatricula.get().rstrip(""), notariod.get(), cajaUR.get(),entry_paginas.get(), notaria_actual.get(), notario_notaria.get(), municipio_not.get(), username, valor_entry_chip )) 
            bd.commit()  

            messagebox.showinfo(message="Registro exitoso", title="Aviso")
            mostrar_tabla()
            limpiar_campos()

        boton_insertar_datos=ttk.Button(ventana5, text="Insertar", command=insertardatos, cursor="hand2")
        boton_insertar_datos.place(relx=0.45, rely=0.95)

        def actualizar_tabla():
            global boton_actualizar_datos
            bd=sqlite3.connect("login1.db")
            mcursor = bd.cursor()

            valorLetras= int (precioNuevo_entry.get())
            añoLetras= int (año_comb.get())
            escrLetras= int (escritura.get())
            precmil=float (valorLetras)
            año_letras= int (entry_fecha3.get())
            dia_letra= int (entry_fecha.get())
            mes_LETRAS = int (entry_fecha2.get())
            numero_escrituras= int (entry_nescrituras.get())
            valor_entry_chip= entry_chip.get()
            valor_entry_chip= "" if valor_entry_chip=="None" else valor_entry_chip 
            
            mcursor.execute("UPDATE datos_escrituras1 SET Matricula=?, Cedula_catastral=?, precio1=?, precioLetras=?, Ubicacion=?, Rural_urbano=?, Departamento=?, Municipio=?, Direccion=?, Modo_Adquirir=?, Escritura=?, EscrLetras=?,  Dia=?, Mes=?, Año=?, AñoLetras=? , Notaria=?, Municipio_notaria=?, Notario=?, cedulavendedor1=?, primerApellidoVen1=?, segundoApellidoVen1=?, primerNombreVen1=?, segundoNombre_Ven1=?, sexoVen1=?, fecha_nacimiento_Ven1=?, RH_Ven1=?, domicilio_Ven1=?, estadocivil_ven1=?, cedulavendedor2=?, primerApellidoVen2=?, segundoApellidoVen2=?, primerNombreVen2=?, segundoNombre_Ven2=?, sexoVen2=?, fecha_nacimiento_Ven2=?, RH_Ven2=?, domicilio_Ven2=?, estadocivil_ven2=?, cedulavendedor3=?, primerApellidoVen3=?, segundoApellidoVen3=?, primerNombreVen3=?, segundoNombre_Ven3=?, sexoVen3=?, fecha_nacimiento_Ven3=?, RH_Ven3=?, domicilio_Ven3=?, estadocivil_ven3=?, cedulavendedor4=?, primerApellidoVen4=?, segundoApellidoVen4=?, primerNombreVen4=?, segundoNombre_Ven4=?, sexoVen4=?, fecha_nacimiento_Ven4=?, RH_Ven4=?, domicilio_Ven4=?, estadocivil_ven4=?, cedulavendedor5=?, primerApellidoVen5=?, segundoApellidoVen5=?, primerNombreVen5=?, segundoNombre_Ven5=?, sexoVen5=?, fecha_nacimiento_Ven5=?, RH_Ven5=?, domicilio_Ven5=?, estadocivil_ven5=?, cedulavendedor6=?, primerApellidoVen6=?, segundoApellidoVen6=?, primerNombreVen6=?, segundoNombre_Ven6=?, sexoVen6=?, fecha_nacimiento_Ven6=?, RH_Ven6=?, domicilio_Ven6=?, estadocivil_ven6=?, cedula_comprador1=?, primerApellido_comp1=?, segundoApellido_comp1=?, primerNombre_comp1=?, segundoNombre_comp1=?, sexo_comp1=?, fecha_nacimiento_comp1=?, RH_comp1=?, domicilio_comp1=?, estadocivil_comp1=?, cedula_comprador2=?, primerApellido_comp2=?, segundoApellido_comp2=?, primerNombre_comp2=?, segundoNombre_comp2=?, sexo_comp2=?, fecha_nacimiento_comp2=?, RH_comp2=?, domicilio_comp2=?, estadocivil_comp2=?, cedula_comprador3=?, primerApellido_comp3=?, segundoApellido_comp3=?, primerNombre_comp3=?, segundoNombre_comp3=?, sexo_comp3=?, fecha_nacimiento_comp3=?, RH_comp3=?, domicilio_comp3=?, estadocivil_comp3=?, cedula_comprador4=?, primerApellido_comp4=?, segundoApellido_comp4=?, primerNombre_comp4=?, segundoNombre_comp4=?, sexo_comp4=?, fecha_nacimiento_comp4=?, RH_comp4=?, domicilio_comp4=?, estadocivil_comp4=?, cedula_comprador5=?, primerApellido_comp5=?, segundoApellido_comp5=?, primerNombre_comp5=?, segundoNombre_comp5=?, sexo_comp5=?, fecha_nacimiento_comp5=?, RH_comp5=?, domicilio_comp5=?, estadocivil_comp5=?, cedula_comprador6=?, primerApellido_comp6=?, segundoApellido_comp6=?, primerNombre_comp6=?, segundoNombre_comp6=?, sexo_comp6=?, fecha_nacimiento_comp6=?, RH_comp6=?, domicilio_comp6=?, estadocivil_comp6=?, precio=?, dia_actual=?, mes_actual=?, año_actual=?, año_letra=?, dia_letra=?, mes_letra=?, numeroescr=?, numeroescr_letras=?, Radicado=?, Linderos=?, pazysalvo=?, Declaracion=?, valor_urbano_rural=?, n_paginas=?, notaria_actual_1=?, notario_notaria_1=?, municipioA=?, username=?, chip_valor=? WHERE ID="+ID_registro.get(), ((matriculaNuevo_entry.get(), cedulaNuevo_entry.get(), ('{:,.2f}'.format(precmil).replace(',','~').replace('.',',').replace('~','.')), str(TextNumber(valorLetras)).strip(), ubicacion.get(), rural_urbano.get(),  departamento.get().upper(), municipio.get().upper(), direccionNueva_entry.get(), adquirir.get(), escritura.get(), str(TextNumber(escrLetras)).strip(), dia_comb.get(), mes_comb.get(), año_comb.get(), str(TextNumber(añoLetras)).strip(), notaria.get(), municipioNota.get(), notario.get(), cedulavendedor_entry.get().strip(), primerApellido1_entry.get().upper().strip(), segundoApellido1_entry.get().upper().strip(), primerNombre1_entry.get().upper().strip(), segundoNombre1_entry.get().upper().strip(), sexo1_entry.get().strip(), fecha_nacimiento1.get().strip(), rh1_entry.get().strip(), domicilio1_entry.get().strip(), estadocivil1_entry.get().strip(), cedulavendedor2_entry.get().strip(), Ven2_primerApellido_entry.get().upper().strip(), Ven2_segundoApellido_entry.get().upper().strip(), Ven2_primerNombre_entry.get().upper().strip(), Ven2_segundoNombre_entry.get().upper().strip(), Ven2_sexo_entry.get().strip(), Ven2_fecha_nacimiento.get().strip(), Ven2_rh_entry.get().strip(),  Ven2_domicilio_entry.get().strip(), Ven2_estadocivil_entry.get().strip(), cedulavendedor3_entry.get().strip(), Ven3_primerApellido_entry.get().upper().strip(), Ven3_segundoApellido_entry.get().upper().strip(),  Ven3_primerNombre_entry.get().upper().strip(), Ven3_segundoNombre_entry.get().upper().strip(),  Ven3_sexo_entry.get().strip(), Ven3_fecha_nacimiento.get().strip(), Ven3_rh_entry.get().strip(),   Ven3_domicilio_entry.get().strip(), Ven3_estadocivil_entry.get().strip(), cedulavendedor4_entry.get().strip(), Ven4_primerApellido_entry.get().upper().strip(),  Ven4_segundoApellido_entry.get().upper().strip(), Ven4_primerNombre_entry.get().upper().strip(),  Ven4_segundoNombre_entry.get().upper().strip(), Ven4_sexo_entry.get().strip(), Ven4_fecha_nacimiento.get().strip(), Ven4_rh_entry.get().strip(), Ven4_domicilio_entry.get().strip(), Ven4_estadocivil_entry.get().strip(),  cedulavendedor5_entry.get().strip(), Ven5_primerApellido_entry.get().upper().strip(),  Ven5_segundoApellido_entry.get().upper().strip(), Ven5_primerNombre_entry.get().upper().strip(),  Ven5_segundoNombre_entry.get().upper().strip(), Ven5_sexo_entry.get().strip(), Ven5_fecha_nacimiento.get().strip(), Ven5_rh_entry.get().strip(),  Ven5_domicilio_entry.get().strip(), Ven5_estadocivil_entry.get().strip(), cedulavendedor6_entry.get().strip(), Ven6_primerApellido_entry.get().upper().strip(),   Ven6_segundoApellido_entry.get().upper().strip(), Ven6_primerNombre_entry.get().upper().strip(),  Ven6_segundoNombre_entry.get().upper().strip(), Ven6_sexo_entry.get().strip(),  Ven6_fecha_nacimiento.get().strip(), Ven6_rh_entry.get().strip(),  Ven6_domicilio_entry.get().strip(), Ven6_estadocivil_entry.get().strip(), cedulacomprador_entry.get().strip(), primer_apellido2_entry.get().upper().strip(), segundo_apellido2_entry.get().upper().strip(), primer_nombre2_entry.get().upper().strip(), segundo_nombre2_entry.get().upper().strip(), sexo2_entry.get().strip(), fecha_nacimiento2.get().strip(), rh2_entry.get().strip(), domicilio2_entry.get().strip(), estadocivil2_entry.get().strip(),  cedulacomprador2_entry.get().strip(), con2_primerApellido_entry.get().upper().strip(),  con2_segundoApellido_entry.get().upper().strip(), con2_primerNombre_entry.get().upper().strip(), con2_segundoNombre_entry.get().upper().strip(), con2_sexo_entry.get().strip(),  con2_fecha_nacimiento.get().strip(), con2_rh_entry.get().strip(), con2_domicilio_entry.get().strip(), con2_estadocivil_entry.get().strip(), cedulacomprador3_entry.get().strip(), con3_primerApellido_entry.get().upper().strip(), con3_segundoApellido_entry.get().upper().strip(), con3_primerNombre_entry.get().upper().strip(), con3_segundoNombre_entry.get().upper().strip(), con3_sexo_entry.get().strip(), con3_fecha_nacimiento.get().strip(), con3_rh_entry.get().strip(), con3_domicilio_entry.get().strip(), con3_estadocivil_entry.get().strip(),  cedulacomprador4_entry.get().strip(), con4_primerApellido_entry.get().upper().strip(),  con4_segundoApellido_entry.get().upper().strip(), con4_primerNombre_entry.get().upper().strip(),  con4_segundoNombre_entry.get().upper().strip(), con4_sexo_entry.get().strip(), con4_fecha_nacimiento.get().strip(), con4_rh_entry.get().strip(), con4_domicilio_entry.get().strip(), con4_estadocivil_entry.get().strip(), cedulacomprador5_entry.get().strip(), con5_primerApellido_entry.get().upper().strip(), con5_segundoApellido_entry.get().upper().strip(), con5_primerNombre_entry.get().upper().strip(), con5_segundoNombre_entry.get().upper().strip(), con5_sexo_entry.get().strip(),  con5_fecha_nacimiento.get().strip(), con5_rh_entry.get().strip(), con5_domicilio_entry.get().strip(), con5_estadocivil_entry.get().strip(),  cedulacomprador6_entry.get().strip(), con6_primerApellido_entry.get().upper().strip(), con6_segundoApellido_entry.get().upper().strip(), con6_primerNombre_entry.get().upper().strip(), con6_segundoNombre_entry.get().upper().strip(), con6_sexo_entry.get().strip(),  con6_fecha_nacimiento.get().strip(), con6_rh_entry.get().strip(),  con6_domicilio_entry.get().strip(), con6_estadocivil_entry.get().strip(), precioNuevo.get(), entry_fecha.get(), entry_fecha2.get(), entry_fecha3.get(), str(TextNumber(año_letras)).strip().upper(), str(TextNumber(dia_letra)).strip().upper(), str(calendar.month_name[mes_LETRAS]).strip().upper(), entry_nescrituras.get(), str(TextNumber(numero_escrituras)).strip().upper(), entry_nradicado.get(), linderomatricula.get().strip(), pazysalvomatricula.get().strip(), notariod.get(), cajaUR.get(), entry_paginas.get(), notaria_actual.get(), notario_notaria.get(), municipio_not.get(), username, valor_entry_chip )))
            bd.commit()

            messagebox.showinfo(message="Actualizacion exitosa")
            mostrar_tabla()
            limpiar_campos()

        boton_actualizar_datos=ttk.Button(ventana5, text="Actualizar", command=actualizar_tabla, cursor="hand2")
        boton_actualizar_datos.place(relx=0.55, rely=0.95)
   
    barra1_menus1 = tk.Menu()
    menu1_archivo1 = tk.Menu(barra1_menus1, tearoff=False )
    menu1_archivo1.add_command(label="Cerrar sesion", command=cerrar_sesión)
    menu1_archivo1.add_command(label="Salir", command=salir_aplicacion)
    barra1_menus1.add_cascade(menu=menu1_archivo1, label="Inicio")

    menu1_archivo2 = tk.Menu(barra1_menus1, tearoff=False)
    menu1_archivo2.add_command(label="VENTA DE NUDA PROPIEDAD SIN RESERVA", command=automatizacion)
    menu1_archivo2.add_command(label="VENTA DE NUDA PROPIEDAD CON RESERVA", command=automatizacion2)
    menu1_archivo2.add_command(label="VENTA TOTAL DE TERRENO", command=automatizacion3)
    menu1_archivo2.add_command(label="VENTA DERECHO DE COUTA", command=automatizacion4)
    menu1_archivo2.add_command(label="VENTA DERECHO DE COUTA DE NUDA PROPIEDAD SIN RESERVA", command=automatizacion1)
    barra1_menus1.add_cascade(menu=menu1_archivo2, label="PROPIEDAD PLENA")

    menu1_archivo3 = tk.Menu(barra1_menus1, tearoff=False)
    menu1_archivo3.add_command(label="VENTA DE POSESIÓN MATERIAL", command=automatizacion5)
    menu1_archivo3.add_command(label="VENTA NUDA PROPIEDAD EN POSESIÓN", command=automatizacion6)
    menu1_archivo3.add_command(label="VENTA DERECHO DE COUTA EN POSESIÓN MATERIAL", command=automatizacion7)
    menu1_archivo3.add_command(label="VENTA NUDA PROPIEDAD DERECHO DE COUTA POSESIÓN", command=automatizacion8)
    barra1_menus1.add_cascade(menu=menu1_archivo3, label="POSESIÓN")

    menu1_archivo4 = tk.Menu(barra1_menus1, tearoff=False)
    menu1_archivo4.add_command(label="VENTA DERECHOS HERENCIALES VINCULADOS ADQUIRIDOS POR COMPRA", command=automatizacion9)
    menu1_archivo4.add_command(label="VENTA DERECHOS HERENCIALES VINCULADOS COMO HEREDEROS", command=automatizacion10)
    barra1_menus1.add_cascade(menu=menu1_archivo4, label="DERECHOS HERENCIALES Y/O GANANCIALES")
    ventana6.config(menu=barra1_menus1)


    menu1_archivo5 = tk.Menu(barra1_menus1, tearoff=False)
    menu1_archivo5.add_command(label="VENTA_E_HIPOTECA_CORPORACION_SOCIAL_DE_CUNDINAMARCA", command=automatizacion11)
    menu1_archivo5.add_command(label="VENTA_E_HIPOTECA_BCSC", command=automatizacion12)
    menu1_archivo5.add_command(label="VENTA_E_HIPOTECA_BANCO_DE_BOGOTA", command=automatizacion13)
    menu1_archivo5.add_command(label="VENTA_E_HIPOTECA_BANCO_AGRARIO", command=automatizacion14)
    menu1_archivo5.add_command(label="RESCILIACION_ DE_ESCRITURA", command=automatizacion15)
    menu1_archivo5.add_command(label="PODER_GENERA", command=automatizacion16)
    menu1_archivo5.add_command(label="DECLARACION_DE_UNION_MARITAL_Y_SOCIEDAD_PATRIMONIA", command=automatizacion17)
    menu1_archivo5.add_command(label="DECLARACION_DE_CONSTRUCCION_PARTICULARES", command=automatizacion18)
    menu1_archivo5.add_command(label="DECLARACION_DE_CONSTRUCCION_DE_COLSUBSIDIO", command=automatizacion19)
    menu1_archivo5.add_command(label="DACION_EN_PAGO", command=automatizacion20)
    menu1_archivo5.add_command(label="CONSTITUCION_HIPOTECA_BANCO_DE_BOGOTA", command=automatizacion21)
    menu1_archivo5.add_command(label="CONSTITUCION_HIPOTECA_BANCO_AGRARIOS", command=automatizacion22)
    menu1_archivo5.add_command(label="CONSTITUCION_HIPOTECA_ABIERTA_PERSONA_NATURAL", command=automatizacion23)
    menu1_archivo5.add_command(label="COMPRAVENTA_CON_PACTO_DE_RETROVENTA", command=automatizacion24)
    menu1_archivo5.add_command(label="CANCELACION_DE_USUFRUCTO", command=automatizacion25)
    menu1_archivo5.add_command(label="CANCELACION_AFECTACION_A_VIVIENDA_FAMILIAR_AFECTO_UNOS", command=automatizacion26)
    menu1_archivo5.add_command(label="CANCELACION_AFECTACION_A_VIVIENDA_FAMILIAR_AFECTARON_AMBOS", command=automatizacion27)
    menu1_archivo5.add_command(label="AUTORIZACION_SALIDA_DEL_PAIS", command=automatizacion28)
    menu1_archivo5.add_command(label="AFECTACION _A_ VIVIENDA FAMILIAR_PROPIEDAD_DE_UNO", command=automatizacion29)
    menu1_archivo5.add_command(label="AFECTACION _A_ VIVIENDA FAMILIAR_PROPIEDAD_DE_DOS", command=automatizacion30)
    menu1_archivo5.add_command(label="ACTUALIZACION_DE_AREA", command=automatizacion31)
    menu1_archivo5.add_command(label="ACTUALIZACION_CEDULA_CATASTRAL", command=automatizacion32)
    menu1_archivo5.add_command(label="ACLARACION", command=automatizacion33)
    
    
    barra1_menus1.add_cascade(menu=menu1_archivo5, label="ultimos formatos")
    ventana6.config(menu=barra1_menus1)


   
    boton_abrir = ttk.Button(ventana6, text="Registrar/Editar", command=registrar_escritura, cursor="hand2")
    boton_abrir.place(relx=0.1, rely=0.8)

    boton_notario=ttk.Button(ventana6, text="Registrar Notario y/o Notaría", command=ingresar_notario, cursor="hand2")
    boton_notario.place(relx=0.3, rely=0.8)

    boton_lindero=ttk.Button(ventana6, text="Agregar lindero y paz y salvo", command=linderos, cursor="hand2")
    boton_lindero.place(relx=0.5, rely=0.8)
      

    boton_mostrar = ttk.Button(ventana6, text="Mostrar", command=mostrar_tabla, cursor="hand2")
    boton_mostrar.place(relx=0.7, rely=0.8)

  

    boton_eliminar = ttk.Button(ventana6, text="Eliminar", state="disabled", command=borrar_tabla, cursor="hand2")
    boton_eliminar.place(relx=0.9, rely=0.8)

    img_adelante = PhotoImage(file="adelante.png")
    img_atras = PhotoImage(file="atras.png") 

    boton_ventana_secundaria = Button(ventana6, image=img_adelante, state="disabled", command=mostrar_ventana_secundaria, cursor="hand2")
    boton_ventana_secundaria.place(relx=0.1, rely=0.87)

   
    tabla = ttk.Treeview(ventana6, height=10,  columns=[f"#{n}" for n in range(1, 6)])
    tabla.place(x=15, y=45, width=1340)

    tabla.column("#0", anchor=CENTER, width=10)
    tabla.column("#1", anchor=CENTER, width=100)
    tabla.column("#2", anchor=CENTER, width=100)
    tabla.column("#3", anchor=CENTER, width=100)
    tabla.column("#4", anchor=CENTER, width=100)
    tabla.column("#5", anchor=CENTER, width=100)
   


    style = ttk.Style()
    style.configure("Treeview.Heading", font=("ARIAL", 10, "bold"))

    # Configura los encabezados de las columnas
    tabla.heading("#0", text="ID", anchor=CENTER)
    tabla.heading("#1", text="MATRICULA", anchor=CENTER)
    tabla.heading("#2", text="CD CATASTRAL", anchor=CENTER)
    tabla.heading("#3", text="PRECIO", anchor=CENTER)
    tabla.heading("#4", text="PRECIO LETRAS", anchor=CENTER)
    tabla.heading("#5", text="UBICACIÓN", anchor=CENTER)


    def limpiarb():
        # Limpiar el texto del Entry
        entrybuscar.delete(0, "end")
        # Dar foco al Entry
        entrybuscar.focus()
    

    
    def buscar():
        # Obtener el texto de búsqueda
        search_text = entrybuscar.get()
        # Obtener todos los ítems del Treeview
        items = tabla.get_children()
        # Obtener el ítem seleccionado actualmente
        selected_item = tabla.selection()
        # Inicializar el índice de inicio de búsqueda
        start_index = 0
        # Verificar si hay un ítem seleccionado
        if selected_item:
            # Obtener el índice del ítem seleccionado
            selected_index = items.index(selected_item[0])
            # Establecer el índice de inicio de búsqueda en el siguiente ítem
            start_index = selected_index + 1
        # Crear una bandera para indicar si se encontró un ítem coincidente
        found = False
        # Recorrer los ítems desde el índice de inicio de búsqueda hasta el final
        for item in items[start_index:]:
            # Obtener los valores del ítem
            values = tabla.item(item, "values")
            # Verificar si el texto de búsqueda está en los valores del ítem
            if search_text in values:
                # Seleccionar el ítem encontrado
                tabla.selection_set(item)
                # Asegurarse de que el ítem sea visible
                tabla.see(item)
                # Establecer la bandera en verdadero
                found = True
                break
        # Verificar si no se encontró un ítem coincidente
        if not found:
            # Recorrer los ítems desde el principio hasta el índice de inicio de búsqueda
            for item in items[:start_index]:
                # Obtener los valores del ítem
                values = tabla.item(item, "values")
                # Verificar si el texto de búsqueda está en los valores del ítem
                if search_text in values:
                    # Seleccionar el ítem encontrado
                    tabla.selection_set(item)
                    # Asegurarse de que el ítem sea visible
                    tabla.see(item)
                    break

    entrybuscar = tk.Entry(ventana6)
    entrybuscar.focus()
    entrybuscar.place(relx=0.3, rely=0.9)
    button_buscar = tk.Button(ventana6, text="Buscar", command=buscar, cursor="hand2")
    button_buscar.place(relx=0.41, rely=0.9)


    button_limpiar = tk.Button(ventana6, text="Limpiar busqueda", command=limpiarb, cursor="hand2")
    button_limpiar.place(relx=0.45, rely=0.9)


    tabla.bind("<ButtonRelease-1>", seleccionar_tabla)
    
    
    ventana6.wait_window(ventana6)

def linderos():
    global ventanalinderos
    global cuadrolinderos
    global tabla4
    global cuadropazysalvo
    global cuadromatricula
    global cuadrodebusqueda
    
    ventanalinderos = Toplevel(ventana6)
    ventanalinderos.title("Linderos, paz y salvo")
    ventanalinderos.geometry("1200x600")
    ventanalinderos.iconbitmap("Logo.ico")
    ventanalinderos.focus_set()
    ventanalinderos.grab_set()
    ventanalinderos.transient(master=ventana6)



    Label(ventanalinderos, text="").pack()

    tabla4 = ttk.Treeview(ventanalinderos, height=1,   columns=[f"#{n}" for n in range(1, 3)])
    tabla4.pack()
    tabla4.heading("#0", text = "Matricula", anchor = CENTER)
    tabla4.heading("#1", text = "Lindero", anchor = CENTER)
    tabla4.heading("#2", text = "paz y salvo", anchor = CENTER)

    tabla4.column("#0", width=100)
    tabla4.column("#1", width=300)
    tabla4.column("#2", width=300)
    

        
    tabla4.bind("<<TreeviewSelect>>", seleccionar_tabla4)

    Label(ventanalinderos, text="").pack()
    Label(ventanalinderos, text="Matricula", font=("ARIAL BLACK", 11)).pack()
    cuadromatricula= ttk.Entry(ventanalinderos)
    cuadromatricula.pack()
    
    titulo1=Label(ventanalinderos, text="Lindero", font=("ARIAL BLACK", 11))
    titulo1.place(relx=0.2, rely=0.28)
    titulo2=Label(ventanalinderos, text="paz y salvo", font=("ARIAL BLACK", 11))
    titulo2.place(relx=0.7, rely=0.28)
    cuadrolinderos= tk.Text(ventanalinderos)
    cuadrolinderos.pack(side=tk.LEFT)
    cuadropazysalvo= tk.Text(ventanalinderos)
    cuadropazysalvo.pack(side=tk.LEFT) 

    ttk.Button(ventanalinderos, text="Insertar Linderos y paz y salvo", cursor="hand2", command=insertar_linderos).place(relx=0.1, rely=0.95)
    
    ttk.Button(ventanalinderos, text="Actualizar lindero y paz y salvo", cursor="hand2", command=actualizar_linderos).place(relx=0.3, rely=0.95)
    
    ttk.Button(ventanalinderos, text="Eliminar registro", cursor="hand2", command=eliminar_linderos).place(relx=0.5, rely=0.95)
    load_data()

    def cuadrodebusquedaM():
            # Obtener el texto de búsqueda
            search_text = cuadrodebusqueda.get()
            # Obtener todos los ítems del Treeview
            items = tabla4.get_children()
            # Obtener el ítem seleccionado actualmente
            selected_item = tabla4.selection()
            # Inicializar el índice de inicio de búsqueda
            start_index = 0
            # Verificar si hay un ítem seleccionado
            if selected_item:
                # Obtener el índice del ítem seleccionado
                selected_index = items.index(selected_item[0])
                # Establecer el índice de inicio de búsqueda en el siguiente ítem
                start_index = selected_index + 1
            # Crear una bandera para indicar si se encontró un ítem coincidente
            found = False
            # Recorrer los ítems desde el índice de inicio de búsqueda hasta el final
            for item in items[start_index:]:
                # Obtener los valores del ítem
                values = tabla4.item(item, "text")
                # Verificar si el texto de búsqueda está en los valores del ítem
                if search_text in values:
                    # Seleccionar el ítem encontrado
                    tabla4.selection_set(item)
                    # Asegurarse de que el ítem sea visible
                    tabla4.see(item)
                    # Establecer la bandera en verdadero
                    found = True
                    break
            # Verificar si no se encontró un ítem coincidente
            if not found:
                # Recorrer los ítems desde el principio hasta el índice de inicio de búsqueda
                for item in items[:start_index]:
                    # Obtener los valores del ítem
                    values = tabla4.item(item, "text")
                    # Verificar si el texto de búsqueda está en los valores del ítem
                    if search_text in values:
                        # Seleccionar el ítem encontrado
                        tabla4.selection_set(item)
                        # Asegurarse de que el ítem sea visible
                        tabla4.see(item)
                        break

    cuadrodebusqueda = tk.Entry(ventanalinderos)
    cuadrodebusqueda.focus()
    cuadrodebusqueda.place(relx=0.7, rely=0.95)
    button_buscarL = tk.Button(ventanalinderos, text="Buscar", command=cuadrodebusquedaM, cursor="hand2")
    button_buscarL.place(relx=0.8, rely=0.95)

    

def ingresar_notario():    
    global ventana7
    ventana7 = Toplevel(ventana6)
    ventana7.title("Notario")
    ventana7.geometry("800x600")
    ventana7.iconbitmap("Logo.ico")
    ventana7.focus_set()
    ventana7.grab_set()
    ventana7.transient(master=ventana6)

    global Id_notario
    global reginotario
    global resolucion_decreto
    global notariaACTUAL
    global notario_notaria1
    global municipio_not1
    

    Id_notario=StringVar()
    reginotario=StringVar()
    resolucion_decreto=StringVar()
    notariaACTUAL=StringVar()
    notario_notaria1=StringVar()
    municipio_not1=StringVar()
    

    global tabla3
    global nombre_notario_entry
    global resolucion_entry
    global notaria_actual_entry
    global notario_notaria_comb
    global municipio_comb



    

    

    Label(ventana7).pack()
    Label(ventana7, text="Nombre Notario(a)", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).pack()
    nombre_notario_entry=Entry(ventana7, textvariable=reginotario, width=34, font=("calibri", 14))
    nombre_notario_entry.focus()
    nombre_notario_entry.pack(ipady=3)
    Label(ventana7).pack()
    

    Label(ventana7, text="Cargo y Resolución ó Decreto", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).pack()
    resolucion_entry=Entry(ventana7, textvariable=resolucion_decreto, width=34, font=("calibri", 14))
    resolucion_entry.pack(ipady=3)
    Label(ventana7).pack()
    Label(ventana7, text="Número de Notaría", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).pack()
    notaria_actual_entry=Entry(ventana7, textvariable=notariaACTUAL, width=34, font=("calibri", 14))
    notaria_actual_entry.pack(ipady=3)
    Label(ventana7).pack()

    opciones1 = {

            "":(),
            "Amazonas": ("Leticia", "El Encanto", "La Chorrera", "La Pedrera","La Victoria",  "Miriti-Parana", 
            "Puerto alegria", "Puerto Arica", "Puerto Nariño", "Puero santander","Tarapaca"), 

            "Antioquia": ("Medellin", "Abejorral", "Abriaquí", "Alejandría", "Amagá", "Amalfi", "Andes", "Angelópolis", "Angostura", "Anorí",
            "Anza", "Apartadó", "Arboletes", "Argelia", "Armenia", "Barbosa", "Bello", "Belmira", "Betania", "Betulia",
            "Briceño", "Buriticá", "Cáceres", "Caicedo", "Caldas", "Campamento", "Cañasgordas", "Caracolí", "Caramanta", "Carepa",
            "Carolina", "Caucasia", "Chigorodó", "Cisneros", "Ciudad Bolívar", "Cocorná", "Concepción", "Concordia", "	Copacabana", "Dabeiba",
            "Don Matías", "Ebéjico", "El Bagre", "El Carmen de Viboral", "El Peñol", "El Retiro", "El Santuario", "Entrerrios", "Envigado", "Fredonia",
            "Frontino", "Giraldo", "Girardota", "Gómez Plata", "Granada", "Guadalupe", "Guarne", "Guatapé", "Heliconia", "Hispania",
            "Itagui", "Ituango", "Jardín", "Jericó", "La Ceja", "La Estrella", "La Pintada", "La Unión", "Liborina", "Maceo",
            "Marinilla", "Montebello", "Murindó", "Mutatá", "Nariño", "Nechí", "Necoclí", "Olaya", "Peque", "Pueblorrico",
            "Puerto Berrío", "Puerto Nare", "Puerto Triunfo", "Remedios", "Rionegro", "Sabanalarga", "Sabaneta", "Salgar", "San Andrés de Cuerquía", "San Carlos",
            "San Francisco", "San Jerónimo", "San José de La Montaña", "San Juan de Urabá", "San Luis", "San Pedro", "San Pedro de Uraba", "San Rafael", "San Roque", "San Vicente",
            "Santa Bárbara", "	Santa Rosa de Osos", "Santafé de Antioquia", "Santo Domingo", "Segovia", "Sonson", "Sopetrán", "9	Támesis", "Tarazá", "Tarso",
            "Titiribí", "Toledo", "Turbo", "Uramita", "Urrao", "Valdivia", "Valparaíso", "Vegachí", "Venecia", "Vigía del Fuerte",
            "Yalí", "	Yarumal", "Yolombó", "Yondó", "Zaragoza"),

            "Arauca":("Arauca", "Arauquita", "Cravo Norte", "Fortul", "Puerto Rodón", "Saravena", "Tame"),

            "Atlantico":("Barranquilla", "Baranoa", "Campo de la Cruz", "Candelaria", "Galapa", "Juan de Acosta", "Luruaco", "Malambo", "Manatí", "Palmar de Varela",
            "Piojó", "Polonuevo", "Ponedera", "Puerto Colombia", "Repelón", "Sabanagrande", "Sabanalarga", "Santa Lucía", "Santo Tomás", "Soledad",
            "Suán", "Tubará", "Usiacurí"),

            
            "Bolivar": ("Cartagena", "Achí", "Altos del Rosario", "Arenal",
            "Arjona", "Arroyohondo", "Barranco de Loba", "Brazuelo de Papayal", "Calamar", "Cantagallo", "El Carmen de Bolívar", "El Carmen de Bolívar", "Cicuco", "Clemencia",
            "Córdoba", "El Guamo", "El Peñón", "Hatillo de Loba", "Magangué", "Mahates", "Margarita", "María La Baja", "Montecristo", "Morales",
            "Norosí", "Pinillos", "Regidor", "Río Viejo", "San Cristóbal", "San Estanislao", "San Fernando", "San Jacinto", "San Jacinto del Cauca", "San Juan Nepomuceno",
            "San Martín de Loba", "San Pablo", "Santa Catalina", "Santa Cruz de Mompox", "Santa Rosa", "Santa Rosa del Sur", "Simití", "Soplaviento", "Talaigua Nuevo", "Tiquisio",
            "Turbaco", "Turbaná", "Villanueva", "Zambrano"),

            "Boyaca":("Tunja", "Almeida", "Aquitania", "Arcabuco", "Belén", "Berbeo", "Betéitiva", "Boavita", "Boyacá", "Briceño",
            "Buenavista", "Busbanzá", "Caldas", "Campohermoso", "Cerinza", "Chinavita", "Chiquinquirá", "Chíquiza", "Chiscas", "Chita",
            "Chitaraque", "Chivatá", "Chivor", "Ciénega", "Cómbita", "Coper", "Corrales", "Covarachía", "Cubará", "Cucaita",
            "Cuítiva", "Duitama", "El Cocuy", "El Espino", "Firavitoba", "Floresta", "Gachantivá", "Gameza", "Garagoa", "Guacamayas",
            "Guateque", "	Guayatá", "Güicán", "Iza", "Jenesano", "Jericó", "La Capilla", "La Uvita", "La Victoria", "Labranzagrande",
            "Macanal", "Maripí", "Miraflores", "Mongua", "Monguí", "Moniquirá", "Motavita", "Muzo", "Nobsa", "Nuevo Colón",
            "Oicatá", "Otanche", "Pachavita", "Páez", "Paipa", "Pajarito", "Panqueba", "Pauna", "Paya","Paz de Río", "Pesca", "Pisba", 
            "Puerto Boyacá", "Quípama", "Ramiriquí", "	Ráquira", "Rondón", "Saboyá", "Sáchica", "Samacá", "San Eduardo", "San José de Pare", 
            "San Luis de Gaceno", "San Mateo", "San Miguel de Sema", "San Pablo de Borbur", "Santa María", "Santa Rosa de Viterbo", "Santa Sofía",
            "Santana", "Sativanorte", "Sativasur", "Siachoque", "Soatá", "Socha", "Socotá", "Sogamoso", "Somondoco", "Sora","Soracá", "Sotaquirá", 
            "Susacón", "Sutamarchán", "Sutatenza", "Tasco", "Tenza", "Tibaná", "Tibasosa", "Tinjacá","Tipacoque", "Toca", "Togüí", "Tópaga", "Tota",
            "Tununguá", "Turmequé", "Tuta", "Tutazá", "Umbita","Ventaquemada", "Villa de Leyva", "Viracachá", "Zetaquira"),

            "Caldas": ("Manizales", "Aguadas", "Anserma", "Aranzazu", "Belalcazar", "Chinchiná", "Filadelfia", "La Dorada", "La Merced", "Manzanares",
            "Marmato", "Marquetalia", "Marulanda", "Neira", "Norcasia", "Pacora", "Palestina", "Pensilvania", "Riosucio", "Risaralda",
            "Salamina", "Samana", "San Jose", "Supía", "Victoria", "Villamaría", "Viterbo"),

            "Caqueta": ("Florencia", "Albania", " Cartagena del Chairá", "Curillo", "El Doncello", "El Paujil", "La Montañita", "Morelia", 
            "Puerto Milán", "Puerto Rico", "San José del Fragua", "San Vicente del Caguán", "Solano", "Solita", "Valparaíso"),

            "Casanare": ("Yopal", "Aguazul", "Chámeza", "Hato Corozal", "La Salina", "Maní", "Monterrey", "Nunchía", "Orocué", "Paz de Ariporo", 
            "Pore", "Recetor", "Sabanalarga", "Sácama", "San Luis de Palenque", "Támara", "Tauramena", "Trinidad", "Villanueva"),

            "Cauca": ("Popayán", "Almaguer", "Argelia", "Balboa", "Bolívar", "Buenos Aires", "Cajibio", "Caldono", "Caloto", "Corinto", "El Tambo", 
            "Florencia", "Guapi", "Inza", "Jambaló", "La Sierra", "La Vega", "López", "Mercaderes", "Miranda", "Morales", "Padilla", 
            "Páez", "Patia (El Bordo)", "Piamonte", "Piendamo",  "Puerto Tejada", "Purace", "Rosas", "San Sebastián", "Santa Rosa",
                "Santander de Quilichao", "Silvia", "Sotara", "Suárez", "Sucre", "Timbío", "Timbiquí", "Toribio", "Totoro", "Villa Rica"),

            "Cesar": ("Valledupar", "Aguachica", "Codazzi", "Astrea", "Becerril", "Bosconia", "Chimichagua", "Curumaní",
            "El Copey", "El Paso", "Gamarra", "González", "La Gloria", "La Jagua Ibirico", "Manaure Balcón Del Cesar", "Pailitas", "Pelaya", 
            "Pueblo Bello","Río De Oro", "Robles (La Paz)", "San Alberto", "San Diego", "San Martín", "Tamalameque"),

            "Choco": ("Quibdó", " Acandi", "Alto Baudo (pie de pato)","Atrato", "Bagado", "Bahia Solano (mutis)", "Bajo Baudo (pizarro)", 
            "Bojaya (bellavista)", "Canton de San Pablo", "Carmen del Darien", "Certegui", "Condoto", "El Carmen", "Istmina", "Jurado", "Litoral del san juan",
            "Lloro","Medio Atrato", "Medio Baudo (boca de pepe)", "Medio San Juan", "Novita", "Nuqui", "Rio iro", "Rio Quito", "Riosucio", 
            "San Jose del Palmar","Sipi", "Tado", "Unguia", "Unión Panamericana"),

            "Córdoba": ("Montería", "Ayapel", "Buenavista", "Canalete", "Cereté", "Chima", "Chinú", "Cienaga De Oro", "Cotorra", "La Apartada", "Lorica",
            "Los Córdobas", "Momil", "Montelíbano",  "Moñitos", "Planeta Rica", "Pueblo Nuevo", "Puerto Escondido", "Puerto Libertador", 
            "Purísima","Sahagún", "San Andrés de Sotavento", "San Antero", "San Bernardo del Viento", "San Carlos", "San Pelayo", "Tierralta", 
            "Valencia"),

            "Cundinamarca": ("Bogotá_D.C","Agua de Dios", "Albán", "Anapoima", "Anolaima", "Apulo", "Arbeláez", "Beltrán", "Bituima", "Bojacá", "Cabrera",
            "Cachipay", "Cajicá", "Caparrapí", "Cáqueza", "Carmen de Carupa", "Chaguaní", "Chía", "Chipaque", "Choachí", 
            "Chocontá","Cogua", "Cota", "Cucunubá", "El Colegio", "El Peñón", "El Rosal", "Facatativá", "Fómeque", "Fosca", "Funza",
            "Fúquene", "Fusagasugá", "Gachalá", "Gachancipá", "Gachetá", "Gama", "Girardot", "Granada", "Guachetá", "Guaduas", "Guasca",
            "Guataquí","Guatavita", "Guayabal de Síquima", "Guayabetal", "Gutiérrez", "Jerusalén", "Junín", "La Calera", "La Mesa", "La Palma",
            "La Peña", "La Vega", "Lenguazaque", "Machetá", "Madrid", "Manta", "Medina", "Mosquera", "Nariño", "Nemocón", "Nilo", "Nimaima",
            "Nocaima", "Pacho", "Paime", "Pandi", "Paratebueno", "Pasca", "Puerto Salgar", "Pulí", "Quebradanegra", "Quetame", "Quipile",
            "Ricaurte","San Antonio del Tequendama", "San Bernardo", "San Cayetano", "San Francisco", "San Juan de Rioseco", "Sasaima", 
            "Sesquilé", "Sibaté", "Silvania", "Simijaca", "Soacha", "Sopó", "Subachoque", "Suesca", "Supatá", "Susa", "Sutatausa", "Tabio",
            "Tausa", "Tena", "Tenjo", "Tibacuy", "Tibirita", "Tocaima", "Tocancipá", "Topaipí", "Ubalá", "Ubaque", "Ubaté", "Une", "Útica",
            "Venecia", "Vergara", "Vianí", "Villagómez", "Villapinzón", "Villeta", "Viotá", "Yacopí", "Zipacón", "Zipaquirá"),

            "Guainía": ("Inírida", "Barrancominas", "Cacahual", "La Guadalupe", "Mapiripana", "Morichal Nuevo", "Pana Pana", 
            "Puerto Colombia", "San Felipe"),

            "Guaviare": ("San Jose del Guaviare", "Calamar", "El Retorno", "Miraflorez"),

            "Huila": ("Neiva", "Acevedo", "Aipe", "Algeciras", "Altamira", "Baraya", "Campoalegre", "Colombia", "Elías", "El Agrado","Garzón",
            "Gigante", "Guadalupe", "Hobo", "Íquira", "Isnos", "La Argentina", "La Plata", "Nátaga", "Oporapa", "Paicol", "Palermo", "Palestina",
            "Pital", "Pitalito", "Rivera", "Saladoblanco", "Santa María", "San Agustín", "Suaza", "Tarqui", "Tello", "Teruel", "Tesalia", "Timaná",
            "Villavieja", "Yaguará"), 

            "La Guajira": ("Riohacha","Albania", "Barrancas", "Dibulla", "Distracción", "El Molino", "Fonseca", "Hatonuevo", "La Jagua del Pilar", "Maicao", 
            "Manaure", "San Juan del Cesar", "Uribia", "Urumita", "Villanueva"),

            "Magdalena": ("Santa Martha", "Algarrobo", "Aracataca", "Ariguaní", "Cerro de San Antonio", "Chibolo", "Ciénaga", "Concordia", "El Banco",
            "El Piñon","El Retén", "Fundación", "Guamal", "Nueva Granada", "Pedraza", "Pijino del Carmen", "Pivijai", "Plato", "Pueblo Viejo", 
            "Remolino", "Sabanas de San Ángel", "Salamina", "San Sebastián de Buenavista", "Santa Ana", "Santa Bárbara de Pinto", "San Zenón", 
            "Sitionuevo", "Tenerife", "Zapayán", "Zona Bananera"),

            "Meta": ("Villavicencio", "Acacías", "Barranca de Upía", "Cabuyaro", "Castilla La Nueva", "	Cubarral", "Cumaral", "El Calvario", "El Castillo",
            "El Dorado","Fuente de Oro", "Granada", "Guamal", "La Macarena", "Lejanías", "Mapiripán", "Mesetas", "Puerto Concordia", "Puerto Gaitán", 
            "Puerto Lleras","Puerto López", "Puerto Rico", "Restrepo", "San Carlos de Guaroa", "San Juan de Arama", "San Juanito", "San Martín", "Uribe", "Vista Hermosa"),

            "Nariño": ("Pasto", "Alban", "Aldaña", "Ancuya", "Arboleda", "Barbacoas", "Belen", "Buesaco", "Chachagui", "Colon(genova)","Consaca", 
            "Contadero", "Cordoba", "Cuaspud", "Cumbal", "Cumbitara", "El Charco", "El Peñol", "El Rosario", "El tablón", "El Tambo", "Funes", "Guachucal",
            "Guaitarilla", "Gualmatan", "Iles", "Imues", "Ipiales", "La Cruz", "La florida", "La llanada","La Tola", "La Unión", "Leiva", "Leiva", "Linares",
            "Los Andes", "Magui", "Mallama", "Mosquera", "Nariño", "Olaya Herrera", "Ospina", "Pizarro", "Policarpa", "Potosi", "Providencia", "Puerres", 
            "Pupiales", "Ricaurte", "Roberto Payan", "Samaniego", "San Bernardo", "San Lorenzo", "San Pablo", "Nariño", "San Pedro de Cartago", "Ospina", 
            "Sandona", "Santa Barbara", "Santacruz", "Sapuyes", "Taminango", "Tangua", "Tumaco", "Tuquerres", "Yacuanquer"),

            "Norte de Santander": ("Cucuta", "Ábrego", "Arboledas", "Bochalema", "Bochalema", "Cáchira", "Cácota", "Chinácota", "Chitagá", "",
            "Convención", "Cucutilla", "Durania", "El Carmen", "El Tarra", "El Zulia", "Gramalote", "Hacarí", "Herrán", "La Esperanza", "La Playa de Belén",
            "Labateca", "Los Patios", "	Lourdes", "Mutiscua", "Ocaña", "Pamplona", "Pamplonita", "Puerto Santander", "Ragonvalia", "Salazar de Las Palmas",
            "San Calixto", "San Cayetano", "Santiago", "Santo Domingo de Silos", "Sardinata", "Teorama", "Tibú", "Toledo", "Villa Caro", "Villa del Rosario"),

            "Putumayo": ("Mocoa", "Colón", "Orito", "Puerto Asís", "Puerto Caicedo", "Puerto Guzmán", "Puerto Leguízamo", "San Francisco", "San Miguel", 
            "Santiago", "Sibundoy", "Valle del Guamuez", "Villagarzón"),

            "Quindio": ("Armenia", "Buenavista", "Calarcá", "Circasia", "Córdoba", "Filandia", "Génova", "La Tebaida", "Montenegro", "Pijao", "Quimbaya","Salento"),

            "Risaralda": ("Pereira", "Apía", "Balboa", "Belén de Umbría", "Dosquebradas", "Guática", "La Celia", "La Virginia", "Marsella", "Mistrató",
            "Pueblo Rico", "Quinchía", "Santa Rosa de Cabal", "Santuario"),

            "San Andres": ("San_Andres"),

            "Santander": ("Bucaramanga", "Aguada", "Albania", "Aratoca", "Barbosa", "Barichara", "Barrancabermeja", "Betulia", "Bolívar", "Cabrera",
            " California", "Capitanejo", "Carcasí", "Cepitá", "Cerrito", "Charalá", "Charta", "Chima", "Chipatá", "Cimitarra", "Concepción", "Confines",
            "Contratación", "Coromoro", "Curití", "El Carmen de Chucurí", "El Guacamayo", "El Peñón", "El Playón", "Encino", "Enciso","Florián", "Floridablanca",
            "Galán", "Gámbita", "Girón", "Guaca", "Guadalupe", "Guapotá", "Guavatá", "Güepsa", "Hato","Jesús María", "Jordán", "La Belleza", "La Paz", "Landázuri",
            "Lebrija", "Los Santos", "Macaravita", "Málaga", "Matanza", "Mogotes", "Molagavita", " Ocamonte", "Oiba", "Onzaga", "Palmar", "Palmas del Socorro",
            "Páramo", "Piedecuesta", "Pinchote", "Puente Nacional", "Puerto Parra", "Puerto Wilches", "Rionegro", "Sabana de Torres", "San Andrés", "San Benito", 
            "San Gil", "San Joaquín", "San José de Miranda", "San Miguel", "San Vicente de Chucurí", "Santa Bárbara", "Santa Helena del Opón", "Simacota",
            "Socorro", "Suaita", "Sucre", "Suratá", "Tona", "Valle de San José", "Vélez", "Vetas", "Villanueva", "Zapatoca"),

            "Sucre": ("Sincelejo", "Buenavista", "Caimito", "Chalán", "Colosó", "Corozal", "Coveñas", "El Roble", "Galeras", "Guaranda", "La Unión", 
            "Los Palmitos", "Majagual", "Morroa", "Ovejas", "Palmito", "Sampués", "San Benito Abad", "San Juan de Betulia", "San Marcos", "San Onofre", 
            "San Pedro", "Santiago de Tolú", "Sincé", "Sucre", "Tolúviejo"),

            "Tolima":("Ibagué", "Alpujarra", "Alvarado", "Ambalema", "Anzoátegui", "Armero guayabal", "Ataco", "cunday", "Cajamarca", "Carmen de Apicalá", "Casabianca",
            "Chaparral", "Coello", "Coyaima", "Dolores", "Espinal", "Falan", "Flandes", "Fresno ", "Guamo", "Herveo", "Honda", "Icononzo", "Lérida ", "Líbano",
            "Mariquita", "Melgar", "Murillo", "Natagaima", "Ortega", "Palocabildo", "Piedras", "Planadas", "Prado", "Purificación", "Rioblanco", "Roncesvalles",
            "Rovira", "Saldaña", "San Antonio", "San Luis", "Santa Isabel", "Suárez", "Valle de San Juan", "Venadillo", "Villahermosa", "Villarrica"),

            "Valle del Cauca": ("Cali", "Alcalá", "Andalucía", "Ansermanuevo", "Argelia", "Bolívar", "Buenaventura", "Buga", "Bugalagrande", "Caicedonia",
            "Calima - El Darién", "Candelaria", "Cartago", "Dagua", "El Águila", "El Cairo", "El Cerrito", "El Dovio", "Florida", "Ginebra", "Guacarí", "Jamundí",
            "La Cumbre", "La Unión", "La Victoria", "Obando", "Palmira", "Pradera", "Restrepo", "Riofrío", "Roldanillo", "San Pedro", "Sevilla", "Toro", "Trujillo", 
            "Tuluá", "Ulloa", "Versalles", "Vijes", "Yotoco", "Yumbo", "Zarzal"),

            "Vaupés": ("Mitú", "Caruru", "Pacoa", "Taraira", "Papunaua", "Yavaraté"),

            "Vichada": ("Puerto Carreño", "Cumaribo", "La Primavera", "Santa Rosalía"),
            }

    def lugar_seleccion1(event):

            
        municipio_comb.set("")
        municipio_comb.config(values=opciones1[notario_notaria_comb.get()])

    Label(ventana7, text="Departamento", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).pack()
    notario_notaria_comb=ttk.Combobox(ventana7, textvariable=notario_notaria1, width=53, state="roadonly", values=tuple(opciones1.keys()))
    notario_notaria_comb.pack(ipady=3)
    notario_notaria_comb.current(0)

    Label(ventana7, text="Municipio", bg="deep sky blue", fg="white", width=34, height=1, font=("calibri", 14)).pack()
    municipio_comb=ttk.Combobox(ventana7, textvariable=municipio_not1, width=53, state="roadonly")
    municipio_comb.pack(ipady=3)
    notario_notaria_comb.bind("<<ComboboxSelected>>", lugar_seleccion1)
    notario_notaria_comb.current(0)
    

    


    
    def insertar_tabla3():
        bd=sqlite3.connect("login1.db")
        mcursor = bd.cursor()

        
        mcursor.execute("INSERT INTO registro_notario VALUES (NULL, ?, ?, ?, ?, ?)", (nombre_notario_entry.get().upper(), resolucion_entry.get().upper(), notaria_actual_entry.get(), notario_notaria_comb.get().upper(), municipio_comb.get().upper()))
        bd.commit()
        messagebox.showinfo(message="Registro exitoso", title="Aviso")

        nombre_notario_entry.delete(0, END)
        resolucion_entry.delete(0, END)
        mostrar_tabla3()
        

    ttk.Button(ventana7, text="Registrar", cursor="hand2", command=insertar_tabla3).place(relx=0.2, rely=0.95)

    ttk.Button(ventana7, text="Mostrar", cursor="hand2", command=mostrar_tabla3).place(relx=0.4, rely=0.95)
    Label(ventana7).pack()
    ttk.Button(ventana7, text="Actualizar", cursor="hand2", command=actualizar_tabla3).place(relx=0.6, rely=0.95)
    Label(ventana7).pack()
    ttk.Button(ventana7, text="Eliminar", cursor="hand2", command=borrar_tabla3).place(relx=0.8, rely=0.95)
    
    tabla3 = ttk.Treeview(ventana7, height=5,  columns=[f"#{n}" for n in range(1, 5)])
    tabla3.pack()
    tabla3.heading("#0", text = "ID", anchor = CENTER)
    tabla3.heading("#1", text = "Nombre Notario", anchor = CENTER)
    tabla3.heading("#2", text = "Resolución ó Decreto", anchor = CENTER)
    tabla3.heading("#3", text = "Departamento", anchor = CENTER)
    tabla3.heading("#4", text = "Municipio", anchor = CENTER)
    tabla3.column("#0", width=30)

    tabla3.bind("<ButtonRelease-1>", seleccionar_tabla3)
    ventana7.wait_window(ventana7)

  
   



 
def mostrar_ventana_secundaria():
    ventana5.deiconify()

def ocultar_ventana_secundaria():
    ventana6.deiconify()
    
  

def cerrar_ventanaregistro():
    ventana5.destroy()
    boton_abrir["state"]="normal"
    boton_notario["state"]="normal"
    boton_lindero["state"]="normal"    
    boton_ventana_secundaria["state"]="disable"
    
    
   
def mostrar_tabla():
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()
    registros=tabla.get_children()
    for elemento in registros:
        tabla.delete(elemento)

    
    mcursor.execute("SELECT * FROM datos_escrituras1")
    for row in mcursor:
        tabla.insert("", 0, text=row[0], values=(row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8], row[9], row[10], row[11], row[12], row[13], row[14], row[15], row[16], row[17], row[18], row[19], row[20], row[21], row[22], row[23], row[24], row[25], row[26], row[27], row[28], row[29], row[30], row[31], row[32], row[33], row[34], row[35], row[36], row[37], row[38], row[39], row[40], row[41], row[42], row[43], row[44], row[45], row[46], row[47], row[48], row[49], row[50], row[51], row[52], row[53], row[54], row[55], row[56], row[57], row[58], row[59], row[60], row[61], row[62], row[63], row[64], row[65], row[66], row[67], row[68], row[69], row[70], row[71], row[72], row[73], row[74], row[75], row[76], row[77], row[78], row[79], row[80], row[81], row[82], row[83], row[84], row[85], row[86], row[87], row[88], row[89], row[90], row[91], row[92], row[93], row[94], row[95], row[96], row[97], row[98], row[99], row[100], row[101], row[102], row[103], row[104], row[105], row[106], row[107], row[108], row[109], row[110], row[111], row[112], row[113], row[114], row[115], row[116], row[117], row[118], row[119], row[120], row[121], row[122], row[123], row[124], row[125], row[126], row[127], row[128], row[129], row[130], row[131], row[132], row[133], row[134], row[135], row[136], row[137], row[138], row[139], row[140], row[141], row[142], row[143], row[144], row[145], row[146], row[147], row[148], row[149], row[150], row[151], row[152], row[153], row[154], row[155], row[156], row[157], row[158], row[159], row[160], row[161], row[162], row[163], row[164], row[165], row[166], row[167], row[168], row[169], row[170], row[171], row[172], row[173], row[174], row[175], row[176], row[177], row[178], row[179], row[180], row[181], row[182], row[183], row[184], row[185], row[186], row[187], row[188], row[189], row[190], row[191], row[192], row[193], row[194], row[195], row[196], row[197], row[198], row[199], row[200], row[201], row[202], row[203], row[204], row[205], row[206], row[207], row[208], row[209], row[210], row[211], row[212], row[213], row[214], row[215], row[216], row[217], row[218], row[219], row[220], row[221], row[222], row[223], row[224], row[225], row[226], row[227], row[228], row[229], row[230], row[231], row[232], row[233], row[234], row[235], row[236], row[237], row[238], row[239], row[240], row[241], row[242], row[243], row[244], row[245], row[246], row[247], row[248], row[249], row[250], row[251], row[252], row[253], row[254], row[255], row[256], row[257], row[258], row[259], row[260], row[261], row[262], row[263], row[264], row[265], row[266], row[267], row[268], row[269], row[270], row[271], row[272], row[273], row[274], row[275], row[276], row[277], row[278], row[279], row[280], row[281], row[282], row[283], row[284], row[285], row[286], row[287], row[288], row[289], row[290], row[291], row[292], row[293], row[294], row[295], row[296], row[297], row[298], row[299], row[300], row[301], row[302], row[303], row[304], row[305], row[306], row[307], row[308], row[309], row[310], row[311], row[312], row[313], row[314], row[315], row[316], row[317], row[318], row[319], row[320], row[321], row[322], row[323], row[324], row[325], row[326], row[327], row[328], row[329], row[330], row[331], row[332], row[333], row[334], row[335], row[336], row[337], row[338], row[339], row[340], row[341], row[342], row[343], row[344], row[345], row[346], row[347], row[348], row[349], row[350], row[351], row[352], row[353], row[354], row[355], row[356], row[357], row[358], row[359], row[360], row[361], row[362], row[363], row[364], row[365], row[366], row[367], row[368], row[369], row[370], row[371], row[372], row[373], row[374], row[375], row[376], row[377], row[378], row[379], row[380], row[381], row[382], row[383], row[384], row[385], row[386], row[387], row[388], row[389], row[390], row[391], row[392], row[393], row[394], row[395], row[396], row[397], row[398], row[399], row[400], row[401], row[402], row[403], row[404], row[405], row[406], row[407], row[408], row[409], row[410], row[411], row[412], row[413], row[414], row[415], row[416], row[417] ))
        
def borrar_tabla():
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()
    try:
        if messagebox.askyesno(message="Desea eliminar el registro?", title="Advertencia"):
            mcursor.execute("DELETE FROM datos_escrituras1 WHERE ID="+ID_registro.get())
            bd.commit()
    except:
        messagebox.showinfo(message="Seleccione un registro")
        pass
    mostrar_tabla()


def limpiar_campos():


    buscarcedulav1_combo.delete(0, END)
    buscarcedulav2_combo.delete(0, END)
    buscarcedulav3_combo.delete(0, END)
    buscarcedulav4_combo.delete(0, END)
    buscarcedulav5_combo.delete(0, END)
    buscarcedulav6_combo.delete(0, END)
    buscarcedulac1_combo.delete(0, END)
    buscarcedulac2_combo.delete(0, END)
    buscarcedulac3_combo.delete(0, END)
    buscarcedulac4_combo.delete(0, END)
    buscarcedulac5_combo.delete(0, END)
    buscarcedulac6_combo.delete(0, END)
    ec_entry.delete(0, END)
    Ven2_ec_entry.delete(0, END)
    Ven3_ec_entry.delete(0, END)
    Ven4_ec_entry.delete(0, END)
    Ven5_ec_entry.delete(0, END)
    Ven6_ec_entry.delete(0, END)

    ecc_entry.delete(0, END)
    con2_ec_entry.delete(0, END)
    con3_ec_entry.delete(0, END)
    con4_ec_entry.delete(0, END)
    con5_ec_entry.delete(0, END)
    con6_ec_entry.delete(0, END)


    buscarcedulav2_combo.grid_remove()
    buscarcedulav3_combo.grid_remove()
    buscarcedulav4_combo.grid_remove()
    buscarcedulav5_combo.grid_remove()
    buscarcedulav6_combo.grid_remove()

    buscarcedulac2_combo.grid_remove()
    buscarcedulac3_combo.grid_remove()
    buscarcedulac4_combo.grid_remove()
    buscarcedulac5_combo.grid_remove()
    buscarcedulac6_combo.grid_remove()

    Ven2_ec_entry.grid_remove()
    Ven3_ec_entry.grid_remove()
    Ven4_ec_entry.grid_remove()
    Ven5_ec_entry.grid_remove()
    Ven6_ec_entry.grid_remove()

    con2_ec_entry.grid_remove()
    con3_ec_entry.grid_remove()
    con4_ec_entry.grid_remove()
    con5_ec_entry.grid_remove()
    con6_ec_entry.grid_remove()

    button1.grid()
    button2.grid_remove()
    button3.grid_remove()
    button4.grid_remove()
    button5.grid_remove()
    button6.grid_remove()

    button7.grid()
    button8.grid_remove()
    button9.grid_remove()
    button10.grid_remove()
    button11.grid_remove()
    button12.grid_remove()

    matriculaNuevo_entry.delete(0, END)
    cedulaNuevo_entry.delete(0, END)
    precioNuevo_entry.delete(0, END)
    ubicacion.delete(0, END)
    rural_urbano.delete(0, END)
    departamento.set("")
    municipio.set("")
    direccionNueva_entry.delete(0, END)
    adquirir.delete(0, END)
    escritura.delete(0, END)
    dia_comb.set("")
    mes_comb.delete(0, END)
    año_comb.delete(0, END)
    notaria.delete(0, END)
    municipioNota.delete(0, END)
    notario.delete(0, END)
    notariod.delete(0, END)
    notaria_actual.delete(0, END)
    notario_notaria.delete(0, END)





    cedulavendedor_entry.delete(0, END)
    primerApellido1_entry.delete(0, END)
    segundoApellido1_entry.delete(0, END)
    primerNombre1_entry.delete(0, END)
    segundoNombre1_entry.delete(0,END)
    sexo1_entry.delete(0, END)
    domicilio1_entry.delete(0, END)
    estadocivil1_entry.delete(0, END)
    rh1_entry.delete(0, END)
    fecha_nacimiento1.delete(0, END)

    cedulavendedor2_entry.delete(0, END)
    Ven2_primerApellido_entry.delete(0, END)
    Ven2_segundoApellido_entry.delete(0, END)
    Ven2_primerNombre_entry.delete(0, END)
    Ven2_segundoNombre_entry.delete(0, END)
    Ven2_sexo_entry.delete(0, END)
    Ven2_fecha_nacimiento.delete(0, END) 
    Ven2_rh_entry.delete(0, END)
    Ven2_domicilio_entry.delete(0, END)
    Ven2_estadocivil_entry.delete(0, END)

    cedulavendedor3_entry.delete(0, END)
    Ven3_primerApellido_entry.delete(0, END)
    Ven3_segundoApellido_entry.delete(0, END)
    Ven3_primerNombre_entry.delete(0, END)
    Ven3_segundoNombre_entry.delete(0, END)
    Ven3_sexo_entry.delete(0, END)
    Ven3_fecha_nacimiento.delete(0, END) 
    Ven3_rh_entry.delete(0, END)
    Ven3_domicilio_entry.delete(0, END)
    Ven3_estadocivil_entry.delete(0, END)

    cedulavendedor4_entry.delete(0, END)
    Ven4_primerApellido_entry.delete(0, END)
    Ven4_segundoApellido_entry.delete(0, END)
    Ven4_primerNombre_entry.delete(0, END)
    Ven4_segundoNombre_entry.delete(0, END)
    Ven4_sexo_entry.delete(0, END)
    Ven4_fecha_nacimiento.delete(0, END) 
    Ven4_rh_entry.delete(0, END)
    Ven4_domicilio_entry.delete(0, END)
    Ven4_estadocivil_entry.delete(0, END)

    cedulavendedor5_entry.delete(0, END)
    Ven5_primerApellido_entry.delete(0, END)
    Ven5_segundoApellido_entry.delete(0, END)
    Ven5_primerNombre_entry.delete(0, END)
    Ven5_segundoNombre_entry.delete(0, END)
    Ven5_sexo_entry.delete(0, END)
    Ven5_fecha_nacimiento.delete(0, END) 
    Ven5_rh_entry.delete(0, END)
    Ven5_domicilio_entry.delete(0, END)
    Ven5_estadocivil_entry.delete(0, END)

    cedulavendedor6_entry.delete(0, END)
    Ven6_primerApellido_entry.delete(0, END)
    Ven6_segundoApellido_entry.delete(0, END)
    Ven6_primerNombre_entry.delete(0, END)
    Ven6_segundoNombre_entry.delete(0, END)
    Ven6_sexo_entry.delete(0, END)
    Ven6_fecha_nacimiento.delete(0, END) 
    Ven6_rh_entry.delete(0, END)
    Ven6_domicilio_entry.delete(0, END)
    Ven6_estadocivil_entry.delete(0, END)




    cedulacomprador_entry.delete(0, END)
    primer_apellido2_entry.delete(0, END)
    segundo_apellido2_entry.delete(0, END)
    primer_nombre2_entry.delete(0, END)
    segundo_nombre2_entry.delete(0,END)
    sexo2_entry.delete(0, END)
    domicilio2_entry.delete(0, END)
    estadocivil2_entry.delete(0, END)
    rh2_entry.delete(0, END)
    fecha_nacimiento2.delete(0, END)

    cedulacomprador2_entry.delete(0, END) 
    con2_primerApellido_entry.delete(0, END)
    con2_segundoApellido_entry.delete(0, END)
    con2_primerNombre_entry.delete(0, END)
    con2_segundoNombre_entry.delete(0, END)
    con2_sexo_entry.delete(0, END)
    con2_fecha_nacimiento.delete(0, END)
    con2_rh_entry.delete(0, END)
    con2_domicilio_entry.delete(0, END)
    con2_estadocivil_entry.delete(0, END)

    cedulacomprador3_entry.delete(0, END) 
    con3_primerApellido_entry.delete(0, END)
    con3_segundoApellido_entry.delete(0, END)
    con3_primerNombre_entry.delete(0, END)
    con3_segundoNombre_entry.delete(0, END)
    con3_sexo_entry.delete(0, END)
    con3_fecha_nacimiento.delete(0, END)
    con3_rh_entry.delete(0, END)
    con3_domicilio_entry.delete(0, END)
    con3_estadocivil_entry.delete(0, END)

    cedulacomprador4_entry.delete(0, END) 
    con4_primerApellido_entry.delete(0, END)
    con4_segundoApellido_entry.delete(0, END)
    con4_primerNombre_entry.delete(0, END)
    con4_segundoNombre_entry.delete(0, END)
    con4_sexo_entry.delete(0, END)
    con4_fecha_nacimiento.delete(0, END)
    con4_rh_entry.delete(0, END)
    con4_domicilio_entry.delete(0, END)
    con4_estadocivil_entry.delete(0, END)

    cedulacomprador5_entry.delete(0, END) 
    con5_primerApellido_entry.delete(0, END)
    con5_segundoApellido_entry.delete(0, END)
    con5_primerNombre_entry.delete(0, END)
    con5_segundoNombre_entry.delete(0, END)
    con5_sexo_entry.delete(0, END)
    con5_fecha_nacimiento.delete(0, END)
    con5_rh_entry.delete(0, END)
    con5_domicilio_entry.delete(0, END)
    con5_estadocivil_entry.delete(0, END)

    cedulacomprador6_entry.delete(0, END) 
    con6_primerApellido_entry.delete(0, END)
    con6_segundoApellido_entry.delete(0, END)
    con6_primerNombre_entry.delete(0, END)
    con6_segundoNombre_entry.delete(0, END)
    con6_sexo_entry.delete(0, END)
    con6_fecha_nacimiento.delete(0, END)
    con6_rh_entry.delete(0, END)
    con6_domicilio_entry.delete(0, END)
    con6_estadocivil_entry.delete(0, END)

    entry_nradicado.delete(0, END)
    entry_nescrituras.delete(0, END)

    cajaUR.delete(0, END)
    entry_paginas.delete(0, END)
    linderomatricula.delete(0, END)
    pazysalvomatricula.delete(0, END)
    municipio_not.delete(0, END)
    entry_chip.delete(0, END)

    label_chip.grid_remove()
    entry_chip.grid_remove()



    cedulavendedor2_entry.grid_remove()
    Ven2_primerApellido_entry.grid_remove()
    Ven2_segundoApellido_entry.grid_remove()
    Ven2_primerNombre_entry.grid_remove()
    Ven2_segundoNombre_entry.grid_remove()
    Ven2_sexo_entry.grid_remove()
    Ven2_fecha_nacimiento.grid_remove() 
    Ven2_rh_entry.grid_remove()
    Ven2_domicilio_entry.grid_remove()
    Ven2_estadocivil_entry.grid_remove()

    cedulavendedor3_entry.grid_remove()
    Ven3_primerApellido_entry.grid_remove()
    Ven3_segundoApellido_entry.grid_remove()
    Ven3_primerNombre_entry.grid_remove()
    Ven3_segundoNombre_entry.grid_remove()
    Ven3_sexo_entry.grid_remove()
    Ven3_fecha_nacimiento.grid_remove()
    Ven3_rh_entry.grid_remove()
    Ven3_domicilio_entry.grid_remove()
    Ven3_estadocivil_entry.grid_remove()

    cedulavendedor4_entry.grid_remove()
    Ven4_primerApellido_entry.grid_remove()
    Ven4_segundoApellido_entry.grid_remove()
    Ven4_primerNombre_entry.grid_remove()
    Ven4_segundoNombre_entry.grid_remove()
    Ven4_sexo_entry.grid_remove()
    Ven4_fecha_nacimiento.grid_remove()
    Ven4_rh_entry.grid_remove()
    Ven4_domicilio_entry.grid_remove()
    Ven4_estadocivil_entry.grid_remove()

    cedulavendedor5_entry.grid_remove()
    Ven5_primerApellido_entry.grid_remove()
    Ven5_segundoApellido_entry.grid_remove()
    Ven5_primerNombre_entry.grid_remove()
    Ven5_segundoNombre_entry.grid_remove()
    Ven5_sexo_entry.grid_remove()
    Ven5_fecha_nacimiento.grid_remove() 
    Ven5_rh_entry.grid_remove()
    Ven5_domicilio_entry.grid_remove()
    Ven5_estadocivil_entry.grid_remove()

    cedulavendedor6_entry.grid_remove()
    Ven6_primerApellido_entry.grid_remove()
    Ven6_segundoApellido_entry.grid_remove()
    Ven6_primerNombre_entry.grid_remove()
    Ven6_segundoNombre_entry.grid_remove()
    Ven6_sexo_entry.grid_remove()
    Ven6_fecha_nacimiento.grid_remove() 
    Ven6_rh_entry.grid_remove()
    Ven6_domicilio_entry.grid_remove()
    Ven6_estadocivil_entry.grid_remove()





    cedulacomprador2_entry.grid_remove()
    con2_primerApellido_entry.grid_remove()
    con2_segundoApellido_entry.grid_remove()
    con2_primerNombre_entry.grid_remove()
    con2_segundoNombre_entry.grid_remove()
    con2_sexo_entry.grid_remove()
    con2_fecha_nacimiento.grid_remove()
    con2_rh_entry.grid_remove()
    con2_domicilio_entry.grid_remove()
    con2_estadocivil_entry.grid_remove()

    cedulacomprador3_entry.grid_remove() 
    con3_primerApellido_entry.grid_remove()
    con3_segundoApellido_entry.grid_remove()
    con3_primerNombre_entry.grid_remove()
    con3_segundoNombre_entry.grid_remove()
    con3_sexo_entry.grid_remove()
    con3_fecha_nacimiento.grid_remove()
    con3_rh_entry.grid_remove()
    con3_domicilio_entry.grid_remove()
    con3_estadocivil_entry.grid_remove()

    cedulacomprador4_entry.grid_remove() 
    con4_primerApellido_entry.grid_remove()
    con4_segundoApellido_entry.grid_remove()
    con4_primerNombre_entry.grid_remove()
    con4_segundoNombre_entry.grid_remove()
    con4_sexo_entry.grid_remove()
    con4_fecha_nacimiento.grid_remove()
    con4_rh_entry.grid_remove()
    con4_domicilio_entry.grid_remove()
    con4_estadocivil_entry.grid_remove()

    cedulacomprador5_entry.grid_remove() 
    con5_primerApellido_entry.grid_remove()
    con5_segundoApellido_entry.grid_remove()
    con5_primerNombre_entry.grid_remove()
    con5_segundoNombre_entry.grid_remove()
    con5_sexo_entry.grid_remove()
    con5_fecha_nacimiento.grid_remove()
    con5_rh_entry.grid_remove()
    con5_domicilio_entry.grid_remove()
    con5_estadocivil_entry.grid_remove()

    cedulacomprador6_entry.grid_remove() 
    con6_primerApellido_entry.grid_remove()
    con6_segundoApellido_entry.grid_remove()
    con6_primerNombre_entry.grid_remove()
    con6_segundoNombre_entry.grid_remove()
    con6_sexo_entry.grid_remove()
    con6_fecha_nacimiento.grid_remove()
    con6_rh_entry.grid_remove()
    con6_domicilio_entry.grid_remove()
    con6_estadocivil_entry.grid_remove()

    boton_agregar_ven2.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()

    boton_agregar_ven4.grid_remove()
    boton_quitar_ven3.grid_remove()

    boton_agregar_ven5.grid_remove()
    boton_quitar_ven4.grid_remove()

    boton_agregar_ven6.grid_remove()
    boton_quitar_ven5.grid_remove()

    boton_quitar_ven6.grid_remove()

    boton_agregar_con2.grid()
    boton_agregar_con3.grid_remove()
    boton_quitar_con2.grid_remove()

    boton_agregar_con4.grid_remove()
    boton_quitar_con3.grid_remove()

    boton_agregar_con5.grid_remove()
    boton_quitar_con4.grid_remove()

    boton_agregar_con6.grid_remove()
    boton_quitar_con5.grid_remove()

    boton_quitar_con6.grid_remove()

def seleccionar_tabla4(event):
    item4 = tabla4.selection()[0]
    cuadromatricula.delete(0, "end")
    cuadromatricula.insert("end", tabla4.item(item4, "text"))
    cuadrolinderos.delete("1.0", "end")
    cuadrolinderos.insert("end", tabla4.item(item4, "values")[0])
    cuadropazysalvo.delete("1.0", "end")
    cuadropazysalvo.insert("end", tabla4.item(item4, "values")[1])
    
    

def insertar_linderos():
    
    valor_matricula = cuadromatricula.get()
    conn = sqlite3.connect("login1.db")
    c = conn.cursor()
    c.execute("SELECT * FROM linderos2 WHERE Matricula_lindero = ?", (valor_matricula, ))
    if c.fetchone() is not None:
        messagebox.showerror("ERROR", "Ya existe la matricula")
    else:
        messagebox.showinfo("Lindero, paz y salvo", "Lindero, paz y salvo guardado por matricula, exitosamente")
        conn = sqlite3.connect("login1.db")
        c = conn.cursor()
        c.execute("INSERT INTO linderos2 VALUES (?, ?, ?)", (cuadromatricula.get(), cuadrolinderos.get("1.0", "end").strip(), cuadropazysalvo.get("1.0", "end").strip() ))
        conn.commit()
        conn.close()
        load_data()
        cuadromatricula.delete(0, END)
        cuadrolinderos.delete("1.0", "end")
        cuadropazysalvo.delete("1.0", "end")

def actualizar_linderos():
    
    valor_matricula= cuadromatricula.get()
    new_name = cuadrolinderos.get("1.0", "end").strip()
    paz_salvo = cuadropazysalvo.get("1.0", "end").strip()
    cuadromatricula.config(state= "disabled")
    
    conn = sqlite3.connect("login1.db")
    c = conn.cursor()
    c.execute("UPDATE linderos2 SET lindero = ?, paz_salvo = ? WHERE  Matricula_lindero = ?", (new_name, paz_salvo, valor_matricula))
    conn.commit()
    conn.close()
    load_data()
    
    

def eliminar_linderos():
    item4 = tabla4.selection()[0]
    item_id = tabla4.item(item4, "text")
   

    conn = sqlite3.connect("login1.db")
    c = conn.cursor()
    c.execute("DELETE FROM linderos2 WHERE Matricula_lindero = ?", (item_id,))
    conn.commit()
    conn.close()
    load_data()
    cuadromatricula.delete(0, END)
    cuadrolinderos.delete("1.0", "end")
    cuadropazysalvo.delete("1.0", "end")
    

def load_data():
    tabla4.delete(*tabla4.get_children())
    conn = sqlite3.connect("login1.db")
    c = conn.cursor()
    c.execute("SELECT * FROM linderos2")
    rows = c.fetchall()
    for row in rows:
        tabla4.insert("", "end", text=row[0], values=(row[1], row[2],))
    conn.close()







def mostrar_tabla3():
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()
    registros=tabla3.get_children()
    for elemento in registros:
        tabla3.delete(elemento)

    
    mcursor.execute("SELECT * FROM registro_notario")
    for row in mcursor:
        tabla3.insert("", 0, text=row[0], values=(row[1], row[2], row[3], row[4], row[5] ))

def borrar_tabla3():
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()
    try:
        if messagebox.askyesno(message="Desea eliminar el registro?", title="Advertencia"):
            mcursor.execute("DELETE FROM registro_notario WHERE ID="+Id_notario.get())
            bd.commit()
    except:
        messagebox.showinfo(message="Seleccione un registro")
        pass
    mostrar_tabla3()

def actualizar_tabla3():
    bd=sqlite3.connect("login1.db")
    mcursor = bd.cursor()

       
    mcursor.execute("UPDATE registro_notario SET Nombre_Notario=?, Declaración=?, notaria_actual=?, notario_notaria=?, municipio1=?  WHERE ID="+Id_notario.get(), (nombre_notario_entry.get().upper(), resolucion_entry.get().upper(), notaria_actual_entry.get().upper(), notario_notaria_comb.get().upper(), municipio_comb.get().upper()))
    bd.commit()

    messagebox.showinfo(message="Actualizacion exitosa")
    mostrar_tabla3()

def seleccionar_tabla3(event):
    item3= tabla3.identify('item', event.x, event.y)
    Id_notario.set(tabla3.item(item3, "text"))
    reginotario.set(tabla3.item(item3, "values")[0])
    resolucion_decreto.set(tabla3.item(item3, "values")[1])
    notariaACTUAL.set(tabla3.item(item3, "values")[2])
    notario_notaria1.set(tabla3.item(item3, "values")[3])
    municipio_not1.set(tabla3.item(item3, "values")[4])
    
    
   

def seleccionar_tabla(event):
    item1= tabla.identify('item', event.x, event.y)
    ID_registro.set(tabla.item(item1, "text"))
    matriculaNuevo.set(tabla.item(item1, "values")[0])
    cedulaNuevo.set(tabla.item(item1, "values")[1])
    precioNuevo.set(tabla.item(item1, "values")[139])
    
    print(len(tabla.item(item1, "values")))

   
    ubicacion1.set(tabla.item(item1, "values")[4])
    rural_urbano1.set(tabla.item(item1, "values")[5])
    
   

    departamento1.set(tabla.item(item1, "values")[6])         
    municipio1.set(tabla.item(item1, "values")[7]) 
    
                                                                 
    direccionNueva1.set(tabla.item(item1, "values")[8])
       
    Modo_adquirir1.set(tabla.item(item1, "values")[9])
    escritura1.set(tabla.item(item1, "values")[10])
    dia1.set(tabla.item(item1, "values")[12])
    mes1.set(tabla.item(item1, "values")[13])
    año1.set(tabla.item(item1, "values")[14])
    notaria1.set(tabla.item(item1, "values")[16])
    Notaria_municipio1.set(tabla.item(item1, "values")[17])
    notario1.set(tabla.item(item1, "values")[18])
    notariod1.set(tabla.item(item1, "values")[348])
         
    

    ven1_cedulavendedor.set(tabla.item(item1, "values")[19])
    ven1_primerApellido.set(tabla.item(item1, "values")[20])
    ven1_segundoApellido.set(tabla.item(item1, "values")[21])
    ven1_primerNombre.set(tabla.item(item1, "values")[22])
    ven1_segundoNombre.set(tabla.item(item1, "values")[23])
    ven1_sexo.set(tabla.item(item1, "values")[24])
    ven1_fechadenacimiento.set(tabla.item(item1, "values")[25])
    ven1_rh.set(tabla.item(item1, "values")[26])
    ven1_domicilio.set(tabla.item(item1, "values")[27])
    ven1_estadocivil.set(tabla.item(item1, "values")[28])

    cedulavendedor2.set(tabla.item(item1, "values")[29])
    Ven2_primerApellido.set(tabla.item(item1, "values")[30])
    Ven2_segundoApellido.set(tabla.item(item1, "values")[31])
    Ven2_primerNombre.set(tabla.item(item1, "values")[32])
    Ven2_segundoNombre.set(tabla.item(item1, "values")[33])
    Ven2_sexo.set(tabla.item(item1, "values")[34])
    Ven2_fechadenacimiento.set(tabla.item(item1, "values")[35])
    Ven2_rh.set(tabla.item(item1, "values")[36])
    Ven2_domicilio.set(tabla.item(item1, "values")[37])
    Ven2_estadocivil.set(tabla.item(item1, "values")[38])


    cedulavendedor3.set(tabla.item(item1, "values")[39])
    Ven3_primerApellido.set(tabla.item(item1, "values")[40])
    Ven3_segundoApellido.set(tabla.item(item1, "values")[41])
    Ven3_primerNombre.set(tabla.item(item1, "values")[42])
    Ven3_segundoNombre.set(tabla.item(item1, "values")[43])
    Ven3_sexo.set(tabla.item(item1, "values")[44])
    ven3_fechadenacimiento.set(tabla.item(item1, "values")[45])
    Ven3_rh.set(tabla.item(item1, "values")[46])
    Ven3_domicilio.set(tabla.item(item1, "values")[47])
    Ven3_estadocivil.set(tabla.item(item1, "values")[48])


    cedulavendedor4.set(tabla.item(item1, "values")[49])
    Ven4_primerApellido.set(tabla.item(item1, "values")[50])
    Ven4_segundoApellido.set(tabla.item(item1, "values")[51])
    Ven4_primerNombre.set(tabla.item(item1, "values")[52])
    Ven4_segundoNombre.set(tabla.item(item1, "values")[53])
    Ven4_sexo.set(tabla.item(item1, "values")[54])
    Ven4_fechanacimiento.set(tabla.item(item1, "values")[55])
    Ven4_rh.set(tabla.item(item1, "values")[56])
    Ven4_domicilio.set(tabla.item(item1, "values")[57])
    Ven4_estadocivil.set(tabla.item(item1, "values")[58])

    cedulavendedor5.set(tabla.item(item1, "values")[59])
    Ven5_primerApellido.set(tabla.item(item1, "values")[60])
    Ven5_segundoApellido.set(tabla.item(item1, "values")[61])
    Ven5_primerNombre.set(tabla.item(item1, "values")[62])
    Ven5_segundoNombre.set(tabla.item(item1, "values")[63])
    Ven5_sexo.set(tabla.item(item1, "values")[64])
    Ven5_fechanacimiento.set(tabla.item(item1, "values")[65])
    Ven5_rh.set(tabla.item(item1, "values")[66])
    Ven5_domicilio.set(tabla.item(item1, "values")[67])
    Ven5_estadocivil.set(tabla.item(item1, "values")[68])

    cedulavendedor6.set(tabla.item(item1, "values")[69])
    Ven6_primerApellido.set(tabla.item(item1, "values")[70])
    Ven6_segundoApellido.set(tabla.item(item1, "values")[71])
    Ven6_primerNombre.set(tabla.item(item1, "values")[72])
    Ven6_segundoNombre.set(tabla.item(item1, "values")[73])
    Ven6_sexo.set(tabla.item(item1, "values")[74])
    Ven6_fechanacimiento.set(tabla.item(item1, "values")[75])
    Ven6_rh.set(tabla.item(item1, "values")[76])
    Ven6_domicilio.set(tabla.item(item1, "values")[77])
    Ven6_estadocivil.set(tabla.item(item1, "values")[78])



    cedulacomprador1.set(tabla.item(item1, "values")[79])
    con1_primer_apellido.set(tabla.item(item1, "values")[80])
    con1_segundo_apellido.set(tabla.item(item1, "values")[81])
    con1_primer_nombre.set(tabla.item(item1, "values")[82])
    con1_segundo_nombre.set(tabla.item(item1, "values")[83])
    con1_sexo.set(tabla.item(item1, "values")[84])
    con1_fechanacimiento.set(tabla.item(item1, "values")[85])
    con1_rh.set(tabla.item(item1, "values")[86])
    con1_domicilio.set(tabla.item(item1, "values")[87])
    con1_estadocivil.set(tabla.item(item1, "values")[88])

    cedulacomprador2.set(tabla.item(item1, "values")[89])
    con2_primerApellido.set(tabla.item(item1, "values")[90])
    con2_segundoApellido.set(tabla.item(item1, "values")[91])
    con2_primerNombre.set(tabla.item(item1, "values")[92])
    con2_segundoNombre.set(tabla.item(item1, "values")[93])
    con2_sexo.set(tabla.item(item1, "values")[94])
    con2_fechanacimiento.set(tabla.item(item1, "values")[95])
    con2_rh.set(tabla.item(item1, "values")[96])
    con2_domicilio.set(tabla.item(item1, "values")[97])
    con2_estadocivil.set(tabla.item(item1, "values")[98])

    cedulacomprador3.set(tabla.item(item1, "values")[99])
    con3_primerApellido.set(tabla.item(item1, "values")[100])
    con3_segundoApellido.set(tabla.item(item1, "values")[101])
    con3_primerNombre.set(tabla.item(item1, "values")[102])
    con3_segundoNombre.set(tabla.item(item1, "values")[103])
    con3_sexo.set(tabla.item(item1, "values")[104])
    con3_fechanacimiento.set(tabla.item(item1, "values")[105])
    con3_rh.set(tabla.item(item1, "values")[106])
    con3_domicilio.set(tabla.item(item1, "values")[107])
    con3_estadocivil.set(tabla.item(item1, "values")[108])

    cedulacomprador4.set(tabla.item(item1, "values")[109])
    con4_primerApellido.set(tabla.item(item1, "values")[110])
    con4_segundoApellido.set(tabla.item(item1, "values")[111])
    con4_primerNombre.set(tabla.item(item1, "values")[112])
    con4_segundoNombre.set(tabla.item(item1, "values")[113])
    con4_sexo.set(tabla.item(item1, "values")[114])
    con4_fechanacimiento.set(tabla.item(item1, "values")[115])
    con4_rh.set(tabla.item(item1, "values")[116])
    con4_domicilio.set(tabla.item(item1, "values")[117])
    con4_estadocivil.set(tabla.item(item1, "values")[118])

    cedulacomprador5.set(tabla.item(item1, "values")[119])
    con5_primerApellido.set(tabla.item(item1, "values")[120])
    con5_segundoApellido.set(tabla.item(item1, "values")[121])
    con5_primerNombre.set(tabla.item(item1, "values")[122])
    con5_segundoNombre.set(tabla.item(item1, "values")[123])
    con5_sexo.set(tabla.item(item1, "values")[124])
    con5_fechanacimiento.set(tabla.item(item1, "values")[125])
    con5_rh.set(tabla.item(item1, "values")[126])
    con5_domicilio.set(tabla.item(item1, "values")[127])
    con5_estadocivil.set(tabla.item(item1, "values")[128])

    cedulacomprador6.set(tabla.item(item1, "values")[129])
    con6_primerApellido.set(tabla.item(item1, "values")[130])
    con6_segundoApellido.set(tabla.item(item1, "values")[131])
    con6_primerNombre.set(tabla.item(item1, "values")[132])
    con6_segundoNombre.set(tabla.item(item1, "values")[133])
    con6_sexo.set(tabla.item(item1, "values")[134])
    con6_fechanacimiento.set(tabla.item(item1, "values")[135])
    con6_rh.set(tabla.item(item1, "values")[136])
    con6_domicilio.set(tabla.item(item1, "values")[137])
    con6_estadocivil.set(tabla.item(item1, "values")[138])
    radicado.set(tabla.item(item1, "values")[345])
    nescrituras.set(tabla.item(item1, "values")[257])

    linderomatricula1.set (tabla.item(item1, "values")[346])
    

    pazysalvomatricula1.set(tabla.item(item1, "values")[347])
    

    notariad1.set(tabla.item(item1, "values")[348])
    cajaUR1.set(tabla.item(item1, "values")[390])
    paginas.set(tabla.item(item1, "values")[391])
    notaria_actual1.set(tabla.item(item1, "values")[396])
    notario_notaria1.set(tabla.item(item1, "values")[397])
    municipio_not1.set(tabla.item(item1, "values")[398])

    entry_chip1.set(tabla.item(item1, "values")[416])

    seleccionar_widget()
    
    
   

    boton_eliminar["state"]="normal"

def ingresarvendedor1a():
    buscarcedulav1_combo.grid()
    cedulavendedor_entry.grid()
    primerApellido1_entry.grid()
    segundoApellido1_entry.grid()
    primerNombre1_entry.grid()
    segundoNombre1_entry.grid()
    sexo1_entry.grid()
    fecha_nacimiento1.grid()
    ec_entry.grid()
    rh1_entry.grid()
    domicilio1_entry.grid()
    estadocivil1_entry.grid()
    boton_agregar_ven2.grid()
    button1.grid() 

    buscarcedulav2_combo.grid_remove()
    cedulavendedor2_entry.grid_remove()
    Ven2_primerApellido_entry.grid_remove()
    Ven2_segundoApellido_entry.grid_remove()
    Ven2_primerNombre_entry.grid_remove()
    Ven2_segundoNombre_entry.grid_remove()
    Ven2_sexo_entry.grid_remove()
    Ven2_fecha_nacimiento.grid_remove()
    Ven2_rh_entry.grid_remove()
    Ven2_ec_entry.grid_remove()
    Ven2_domicilio_entry.grid_remove()
    Ven2_estadocivil_entry.grid_remove()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button2.grid_remove()
    
    buscarcedulav3_combo.grid_remove()
    cedulavendedor3_entry.grid_remove()
    Ven3_primerApellido_entry.grid_remove()
    Ven3_segundoApellido_entry.grid_remove()
    Ven3_primerNombre_entry.grid_remove()
    Ven3_segundoNombre_entry.grid_remove()
    Ven3_sexo_entry.grid_remove()
    Ven3_fecha_nacimiento.grid_remove()
    Ven3_rh_entry.grid_remove()
    Ven3_ec_entry.grid_remove()
    Ven3_domicilio_entry.grid_remove()
    Ven3_estadocivil_entry.grid_remove()
    boton_agregar_ven4.grid_remove()
    boton_quitar_ven3.grid_remove()
    button3.grid_remove()

    buscarcedulav4_combo.grid_remove()
    cedulavendedor4_entry.grid_remove()
    Ven4_primerApellido_entry.grid_remove()
    Ven4_segundoApellido_entry.grid_remove()
    Ven4_primerNombre_entry.grid_remove()
    Ven4_segundoNombre_entry.grid_remove()
    Ven4_sexo_entry.grid_remove()
    Ven4_fecha_nacimiento.grid_remove()
    Ven4_rh_entry.grid_remove()
    Ven4_ec_entry.grid_remove()
    Ven4_domicilio_entry.grid_remove()
    Ven4_estadocivil_entry.grid_remove()
    boton_agregar_ven5.grid_remove()
    boton_quitar_ven4.grid_remove()
    button4.grid_remove()

    buscarcedulav5_combo.grid_remove()
    cedulavendedor5_entry.grid_remove()
    Ven5_primerApellido_entry.grid_remove()
    Ven5_segundoApellido_entry.grid_remove()
    Ven5_primerNombre_entry.grid_remove()
    Ven5_segundoNombre_entry.grid_remove()
    Ven5_sexo_entry.grid_remove()
    Ven5_fecha_nacimiento.grid_remove()
    Ven5_rh_entry.grid_remove()
    Ven5_ec_entry.grid_remove()
    Ven5_domicilio_entry.grid_remove()
    Ven5_estadocivil_entry.grid_remove()
    boton_agregar_ven6.grid_remove()
    boton_quitar_ven5.grid_remove()
    button5.grid_remove()

    buscarcedulav6_combo.grid_remove()
    cedulavendedor6_entry.grid_remove()
    Ven6_primerApellido_entry.grid_remove()
    Ven6_segundoApellido_entry.grid_remove()
    Ven6_primerNombre_entry.grid_remove()
    Ven6_segundoNombre_entry.grid_remove()
    Ven6_sexo_entry.grid_remove()
    Ven6_fecha_nacimiento.grid_remove()
    Ven6_rh_entry.grid_remove()
    Ven6_ec_entry.grid_remove()
    Ven6_domicilio_entry.grid_remove()
    Ven6_estadocivil_entry.grid_remove()
    boton_quitar_ven6.grid_remove()
    button6.grid_remove()


            
            
    
def ingresarvendedor2a():
    buscarcedulav1_combo.grid()
    cedulavendedor_entry.grid()
    primerApellido1_entry.grid()
    segundoApellido1_entry.grid()
    primerNombre1_entry.grid()
    segundoNombre1_entry.grid()
    sexo1_entry.grid()
    fecha_nacimiento1.grid()
    rh1_entry.grid()
    domicilio1_entry.grid()
    estadocivil1_entry.grid()
    boton_agregar_ven2.grid_remove()
    button1.grid_remove()
               
    buscarcedulav2_combo.grid()        
    cedulavendedor2_entry.grid()
    Ven2_primerApellido_entry.grid()
    Ven2_segundoApellido_entry.grid()
    Ven2_primerNombre_entry.grid()
    Ven2_segundoNombre_entry.grid()
    Ven2_sexo_entry.grid()
    Ven2_fecha_nacimiento.grid()
    Ven2_rh_entry.grid()
    Ven2_ec_entry.grid()
    Ven2_domicilio_entry.grid()
    Ven2_estadocivil_entry.grid()
    boton_agregar_ven3.grid()
    boton_quitar_ven2.grid()
    button2.grid()
    
    buscarcedulav3_combo.grid_remove()
    cedulavendedor3_entry.grid_remove()
    Ven3_primerApellido_entry.grid_remove()
    Ven3_segundoApellido_entry.grid_remove()
    Ven3_primerNombre_entry.grid_remove()
    Ven3_segundoNombre_entry.grid_remove()
    Ven3_sexo_entry.grid_remove()
    Ven3_fecha_nacimiento.grid_remove()
    Ven3_rh_entry.grid_remove()
    Ven3_domicilio_entry.grid_remove()
    Ven3_estadocivil_entry.grid_remove()
    boton_agregar_ven4.grid_remove()
    boton_quitar_ven3.grid_remove()
    button3.grid_remove()
    Ven3_ec_entry.grid_remove()

    buscarcedulav4_combo.grid_remove()
    cedulavendedor4_entry.grid_remove()
    Ven4_primerApellido_entry.grid_remove()
    Ven4_segundoApellido_entry.grid_remove()
    Ven4_primerNombre_entry.grid_remove()
    Ven4_segundoNombre_entry.grid_remove()
    Ven4_sexo_entry.grid_remove()
    Ven4_fecha_nacimiento.grid_remove()
    Ven4_rh_entry.grid_remove()
    Ven4_domicilio_entry.grid_remove()
    Ven4_estadocivil_entry.grid_remove()
    boton_agregar_ven5.grid_remove()
    boton_quitar_ven4.grid_remove()
    button4.grid_remove()
    Ven4_ec_entry.grid_remove()

    buscarcedulav5_combo.grid_remove()
    cedulavendedor5_entry.grid_remove()
    Ven5_primerApellido_entry.grid_remove()
    Ven5_segundoApellido_entry.grid_remove()
    Ven5_primerNombre_entry.grid_remove()
    Ven5_segundoNombre_entry.grid_remove()
    Ven5_sexo_entry.grid_remove()
    Ven5_fecha_nacimiento.grid_remove()
    Ven5_rh_entry.grid_remove()
    Ven5_domicilio_entry.grid_remove()
    Ven5_estadocivil_entry.grid_remove()
    boton_agregar_ven6.grid_remove()
    boton_quitar_ven5.grid_remove()
    button5.grid_remove()
    Ven5_ec_entry.grid_remove()

    buscarcedulav6_combo.grid_remove()
    cedulavendedor6_entry.grid_remove()
    Ven6_primerApellido_entry.grid_remove()
    Ven6_segundoApellido_entry.grid_remove()
    Ven6_primerNombre_entry.grid_remove()
    Ven6_segundoNombre_entry.grid_remove()
    Ven6_sexo_entry.grid_remove()
    Ven6_fecha_nacimiento.grid_remove()
    Ven6_rh_entry.grid_remove()
    Ven6_domicilio_entry.grid_remove()
    Ven6_estadocivil_entry.grid_remove()
    boton_quitar_ven6.grid_remove()
    button6.grid_remove()
    Ven6_ec_entry.grid_remove()



   
   

def ingresarvendedor3a():
    buscarcedulav1_combo.grid()
    cedulavendedor_entry.grid()
    primerApellido1_entry.grid()
    segundoApellido1_entry.grid()
    primerNombre1_entry.grid()
    segundoNombre1_entry.grid()
    sexo1_entry.grid()
    fecha_nacimiento1.grid()
    rh1_entry.grid()
    domicilio1_entry.grid()
    estadocivil1_entry.grid()
    boton_agregar_ven2.grid_remove()
    button1.grid_remove()
    ec_entry.grid()
    

    buscarcedulav2_combo.grid()        
    cedulavendedor2_entry.grid()
    Ven2_primerApellido_entry.grid()
    Ven2_segundoApellido_entry.grid()
    Ven2_primerNombre_entry.grid()
    Ven2_segundoNombre_entry.grid()
    Ven2_sexo_entry.grid()
    Ven2_fecha_nacimiento.grid()
    Ven2_rh_entry.grid()
    Ven2_domicilio_entry.grid()
    Ven2_estadocivil_entry.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button2.grid_remove()
    Ven2_ec_entry.grid()

    buscarcedulav3_combo.grid()        
    cedulavendedor3_entry.grid()
    Ven3_primerApellido_entry.grid()
    Ven3_segundoApellido_entry.grid()
    Ven3_primerNombre_entry.grid()
    Ven3_segundoNombre_entry.grid()
    Ven3_sexo_entry.grid()
    Ven3_fecha_nacimiento.grid()
    Ven3_rh_entry.grid()
    Ven3_domicilio_entry.grid()
    Ven3_estadocivil_entry.grid()
    boton_agregar_ven4.grid()
    boton_quitar_ven3.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button3.grid()
    Ven3_ec_entry.grid()

    buscarcedulav4_combo.grid_remove()
    cedulavendedor4_entry.grid_remove()
    Ven4_primerApellido_entry.grid_remove()
    Ven4_segundoApellido_entry.grid_remove()
    Ven4_primerNombre_entry.grid_remove()
    Ven4_segundoNombre_entry.grid_remove()
    Ven4_sexo_entry.grid_remove()
    Ven4_fecha_nacimiento.grid_remove()
    Ven4_rh_entry.grid_remove()
    Ven4_domicilio_entry.grid_remove()
    Ven4_estadocivil_entry.grid_remove()
    boton_agregar_ven5.grid_remove()
    boton_quitar_ven4.grid_remove()
    button4.grid_remove()
    Ven4_ec_entry.grid_remove()

    buscarcedulav5_combo.grid_remove()
    cedulavendedor5_entry.grid_remove()
    Ven5_primerApellido_entry.grid_remove()
    Ven5_segundoApellido_entry.grid_remove()
    Ven5_primerNombre_entry.grid_remove()
    Ven5_segundoNombre_entry.grid_remove()
    Ven5_sexo_entry.grid_remove()
    Ven5_fecha_nacimiento.grid_remove()
    Ven5_rh_entry.grid_remove()
    Ven5_domicilio_entry.grid_remove()
    Ven5_estadocivil_entry.grid_remove()
    boton_agregar_ven6.grid_remove()
    boton_quitar_ven5.grid_remove()
    button5.grid_remove()
    Ven5_ec_entry.grid_remove()

    buscarcedulav6_combo.grid_remove()
    cedulavendedor6_entry.grid_remove()
    Ven6_primerApellido_entry.grid_remove()
    Ven6_segundoApellido_entry.grid_remove()
    Ven6_primerNombre_entry.grid_remove()
    Ven6_segundoNombre_entry.grid_remove()
    Ven6_sexo_entry.grid_remove()
    Ven6_fecha_nacimiento.grid_remove()
    Ven6_rh_entry.grid_remove()
    Ven6_domicilio_entry.grid_remove()
    Ven6_estadocivil_entry.grid_remove()
    boton_quitar_ven6.grid_remove()
    button6.grid_remove()
    Ven6_ec_entry.grid_remove()

    
   
def ingresarvendedor4a():
    buscarcedulav1_combo.grid()
    cedulavendedor_entry.grid()
    primerApellido1_entry.grid()
    segundoApellido1_entry.grid()
    primerNombre1_entry.grid()
    segundoNombre1_entry.grid()
    sexo1_entry.grid()
    fecha_nacimiento1.grid()
    rh1_entry.grid()
    domicilio1_entry.grid()
    estadocivil1_entry.grid()
    boton_agregar_ven2.grid_remove()
    button1.grid_remove()
    ec_entry.grid()
    

    buscarcedulav2_combo.grid()        
    cedulavendedor2_entry.grid()
    Ven2_primerApellido_entry.grid()
    Ven2_segundoApellido_entry.grid()
    Ven2_primerNombre_entry.grid()
    Ven2_segundoNombre_entry.grid()
    Ven2_sexo_entry.grid()
    Ven2_fecha_nacimiento.grid()
    Ven2_rh_entry.grid()
    Ven2_domicilio_entry.grid()
    Ven2_estadocivil_entry.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button2.grid_remove()
    Ven2_ec_entry.grid()

    buscarcedulav3_combo.grid()       
    cedulavendedor3_entry.grid()
    Ven3_primerApellido_entry.grid()
    Ven3_segundoApellido_entry.grid()
    Ven3_primerNombre_entry.grid()
    Ven3_segundoNombre_entry.grid()
    Ven3_sexo_entry.grid()
    Ven3_fecha_nacimiento.grid()
    Ven3_rh_entry.grid()
    Ven3_domicilio_entry.grid()
    Ven3_estadocivil_entry.grid()
    boton_agregar_ven4.grid()
    boton_quitar_ven3.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button3.grid_remove()
    Ven3_ec_entry.grid()

    buscarcedulav4_combo.grid()
    cedulavendedor4_entry.grid()
    Ven4_primerApellido_entry.grid()
    Ven4_segundoApellido_entry.grid()
    Ven4_primerNombre_entry.grid()
    Ven4_segundoNombre_entry.grid()
    Ven4_sexo_entry.grid()
    Ven4_fecha_nacimiento.grid()
    Ven4_rh_entry.grid()
    Ven4_domicilio_entry.grid()
    Ven4_estadocivil_entry.grid()
    boton_agregar_ven5.grid()
    boton_quitar_ven4.grid()
    boton_agregar_ven4.grid_remove()
    boton_quitar_ven3.grid_remove()
    button4.grid()
    Ven4_ec_entry.grid()

    buscarcedulav5_combo.grid_remove()
    cedulavendedor5_entry.grid_remove()
    Ven5_primerApellido_entry.grid_remove()
    Ven5_segundoApellido_entry.grid_remove()
    Ven5_primerNombre_entry.grid_remove()
    Ven5_segundoNombre_entry.grid_remove()
    Ven5_sexo_entry.grid_remove()
    Ven5_fecha_nacimiento.grid_remove()
    Ven5_rh_entry.grid_remove()
    Ven5_domicilio_entry.grid_remove()
    Ven5_estadocivil_entry.grid_remove()
    boton_agregar_ven6.grid_remove()
    boton_quitar_ven5.grid_remove()
    button5.grid_remove()
    Ven5_ec_entry.grid_remove()

    buscarcedulav6_combo.grid_remove()
    cedulavendedor6_entry.grid_remove()
    Ven6_primerApellido_entry.grid_remove()
    Ven6_segundoApellido_entry.grid_remove()
    Ven6_primerNombre_entry.grid_remove()
    Ven6_segundoNombre_entry.grid_remove()
    Ven6_sexo_entry.grid_remove()
    Ven6_fecha_nacimiento.grid_remove()
    Ven6_rh_entry.grid_remove()
    Ven6_domicilio_entry.grid_remove()
    Ven6_estadocivil_entry.grid_remove()
    boton_quitar_ven6.grid_remove()
    button6.grid_remove()
    Ven6_ec_entry.grid_remove()

    

    
def ingresarvendedor5a():

    buscarcedulav1_combo.grid()
    cedulavendedor_entry.grid()
    primerApellido1_entry.grid()
    segundoApellido1_entry.grid()
    primerNombre1_entry.grid()
    segundoNombre1_entry.grid()
    sexo1_entry.grid()
    fecha_nacimiento1.grid()
    rh1_entry.grid()
    domicilio1_entry.grid()
    estadocivil1_entry.grid()
    boton_agregar_ven2.grid_remove()
    button1.grid_remove()
    ec_entry.grid()

    buscarcedulav2_combo.grid()        
    cedulavendedor2_entry.grid()
    Ven2_primerApellido_entry.grid()
    Ven2_segundoApellido_entry.grid()
    Ven2_primerNombre_entry.grid()
    Ven2_segundoNombre_entry.grid()
    Ven2_sexo_entry.grid()
    Ven2_fecha_nacimiento.grid()
    Ven2_rh_entry.grid()
    Ven2_domicilio_entry.grid()
    Ven2_estadocivil_entry.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button2.grid_remove()
    Ven2_ec_entry.grid()

    buscarcedulav3_combo.grid()      
    cedulavendedor3_entry.grid()
    Ven3_primerApellido_entry.grid()
    Ven3_segundoApellido_entry.grid()
    Ven3_primerNombre_entry.grid()
    Ven3_segundoNombre_entry.grid()
    Ven3_sexo_entry.grid()
    Ven3_fecha_nacimiento.grid()
    Ven3_rh_entry.grid()
    Ven3_domicilio_entry.grid()
    Ven3_estadocivil_entry.grid()
    boton_agregar_ven4.grid()
    boton_quitar_ven3.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button3.grid_remove()
    Ven3_ec_entry.grid()

    buscarcedulav4_combo.grid()
    cedulavendedor4_entry.grid()
    Ven4_primerApellido_entry.grid()
    Ven4_segundoApellido_entry.grid()
    Ven4_primerNombre_entry.grid()
    Ven4_segundoNombre_entry.grid()
    Ven4_sexo_entry.grid()
    Ven4_fecha_nacimiento.grid()
    Ven4_rh_entry.grid()
    Ven4_domicilio_entry.grid()
    Ven4_estadocivil_entry.grid()
    boton_agregar_ven4.grid_remove()
    boton_quitar_ven3.grid_remove()
    button4.grid_remove()
    Ven4_ec_entry.grid()

    buscarcedulav5_combo.grid()
    cedulavendedor5_entry.grid()
    Ven5_primerApellido_entry.grid()
    Ven5_segundoApellido_entry.grid()
    Ven5_primerNombre_entry.grid()
    Ven5_segundoNombre_entry.grid()
    Ven5_sexo_entry.grid()
    Ven5_fecha_nacimiento.grid()
    Ven5_rh_entry.grid()
    Ven5_domicilio_entry.grid()
    Ven5_estadocivil_entry.grid()
    boton_agregar_ven6.grid()
    boton_quitar_ven5.grid()
    boton_agregar_ven5.grid_remove()
    boton_quitar_ven4.grid_remove()
    button5.grid()
    Ven5_ec_entry.grid()

    buscarcedulav6_combo.grid_remove()
    cedulavendedor6_entry.grid_remove()
    Ven6_primerApellido_entry.grid_remove()
    Ven6_segundoApellido_entry.grid_remove()
    Ven6_primerNombre_entry.grid_remove()
    Ven6_segundoNombre_entry.grid_remove()
    Ven6_sexo_entry.grid_remove()
    Ven6_fecha_nacimiento.grid_remove()
    Ven6_rh_entry.grid_remove()
    Ven6_domicilio_entry.grid_remove()
    Ven6_estadocivil_entry.grid_remove()
    boton_quitar_ven6.grid_remove()
    Ven6_ec_entry.grid_remove()

    
def ingresarvendedor6a():

    buscarcedulav1_combo.grid()
    cedulavendedor_entry.grid()
    primerApellido1_entry.grid()
    segundoApellido1_entry.grid()
    primerNombre1_entry.grid()
    segundoNombre1_entry.grid()
    sexo1_entry.grid()
    fecha_nacimiento1.grid()
    rh1_entry.grid()
    domicilio1_entry.grid()
    estadocivil1_entry.grid()
    boton_agregar_ven2.grid_remove()
    button1.grid_remove()
    ec_entry.grid()
               
    buscarcedulav2_combo.grid()        
    cedulavendedor2_entry.grid()
    Ven2_primerApellido_entry.grid()
    Ven2_segundoApellido_entry.grid()
    Ven2_primerNombre_entry.grid()
    Ven2_segundoNombre_entry.grid()
    Ven2_sexo_entry.grid()
    Ven2_fecha_nacimiento.grid()
    Ven2_rh_entry.grid()
    Ven2_domicilio_entry.grid()
    Ven2_estadocivil_entry.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button2.grid_remove()
    Ven2_ec_entry.grid()

    buscarcedulav3_combo.grid()        
    cedulavendedor3_entry.grid()
    Ven3_primerApellido_entry.grid()
    Ven3_segundoApellido_entry.grid()
    Ven3_primerNombre_entry.grid()
    Ven3_segundoNombre_entry.grid()
    Ven3_sexo_entry.grid()
    Ven3_fecha_nacimiento.grid()
    Ven3_rh_entry.grid()
    Ven3_domicilio_entry.grid()
    Ven3_estadocivil_entry.grid()
    boton_agregar_ven4.grid()
    boton_quitar_ven3.grid()
    boton_agregar_ven3.grid_remove()
    boton_quitar_ven2.grid_remove()
    button3.grid_remove()
    Ven3_ec_entry.grid()

    buscarcedulav4_combo.grid()
    cedulavendedor4_entry.grid()
    Ven4_primerApellido_entry.grid()
    Ven4_segundoApellido_entry.grid()
    Ven4_primerNombre_entry.grid()
    Ven4_segundoNombre_entry.grid()
    Ven4_sexo_entry.grid()
    Ven4_fecha_nacimiento.grid()
    Ven4_rh_entry.grid()
    Ven4_domicilio_entry.grid()
    Ven4_estadocivil_entry.grid()
    boton_agregar_ven4.grid_remove()
    boton_quitar_ven3.grid_remove()
    button4.grid_remove()
    Ven4_ec_entry.grid()

    buscarcedulav5_combo.grid()
    cedulavendedor5_entry.grid()
    Ven5_primerApellido_entry.grid()
    Ven5_segundoApellido_entry.grid()
    Ven5_primerNombre_entry.grid()
    Ven5_segundoNombre_entry.grid()
    Ven5_sexo_entry.grid()
    Ven5_fecha_nacimiento.grid()
    Ven5_rh_entry.grid()
    Ven5_domicilio_entry.grid()
    Ven5_estadocivil_entry.grid()
    boton_agregar_ven5.grid_remove()
    boton_quitar_ven4.grid_remove()
    button5.grid_remove()
    Ven5_ec_entry.grid()

    buscarcedulav6_combo.grid()
    cedulavendedor6_entry.grid()
    Ven6_primerApellido_entry.grid()
    Ven6_segundoApellido_entry.grid()
    Ven6_primerNombre_entry.grid()
    Ven6_segundoNombre_entry.grid()
    Ven6_sexo_entry.grid()
    Ven6_fecha_nacimiento.grid()
    Ven6_rh_entry.grid()
    Ven6_domicilio_entry.grid()
    Ven6_estadocivil_entry.grid()
    boton_quitar_ven6.grid()
    boton_agregar_ven6.grid_remove()
    boton_quitar_ven5.grid_remove()
    button6.grid()
    Ven6_ec_entry.grid()

def ingresarcomprador1a():

    buscarcedulac1_combo.grid()
    button7.grid()
    cedulacomprador_entry.grid()
    primer_apellido2_entry.grid()
    segundo_apellido2_entry.grid()
    primer_nombre2_entry.grid()
    segundo_nombre2_entry.grid()
    sexo2_entry.grid()
    fecha_nacimiento2.grid()
    rh2_entry.grid()
    domicilio2_entry.grid()
    estadocivil2_entry.grid() 
    boton_agregar_con2.grid()
    ecc_entry.grid()

    buscarcedulac2_combo.grid_remove()
    button8.grid_remove()
    cedulacomprador2_entry.grid_remove()
    con2_primerApellido_entry.grid_remove()
    con2_segundoApellido_entry.grid_remove()
    con2_primerNombre_entry.grid_remove()
    con2_segundoNombre_entry.grid_remove()
    con2_sexo_entry.grid_remove()
    con2_fecha_nacimiento.grid_remove()
    con2_rh_entry.grid_remove()
    con2_domicilio_entry.grid_remove()
    con2_estadocivil_entry.grid_remove()
    boton_agregar_con3.grid_remove()
    boton_quitar_con2.grid_remove()
    con2_ec_entry.grid_remove()

    buscarcedulac3_combo.grid_remove()
    button9.grid_remove()
    cedulacomprador3_entry.grid_remove()
    con3_primerApellido_entry.grid_remove()
    con3_segundoApellido_entry.grid_remove()
    con3_primerNombre_entry.grid_remove()
    con3_segundoNombre_entry.grid_remove()
    con3_sexo_entry.grid_remove()
    con3_fecha_nacimiento.grid_remove()
    con3_rh_entry.grid_remove()
    con3_domicilio_entry.grid_remove()
    con3_estadocivil_entry.grid_remove()
    boton_agregar_con4.grid_remove()
    boton_quitar_con3.grid_remove()
    con3_ec_entry.grid_remove()

    buscarcedulac4_combo.grid_remove()
    button10.grid_remove()
    cedulacomprador4_entry.grid_remove()
    con4_primerApellido_entry.grid_remove()
    con4_segundoApellido_entry.grid_remove()
    con4_primerNombre_entry.grid_remove()
    con4_segundoNombre_entry.grid_remove()
    con4_sexo_entry.grid_remove()
    con4_fecha_nacimiento.grid_remove()
    con4_rh_entry.grid_remove()
    con4_domicilio_entry.grid_remove()
    con4_estadocivil_entry.grid_remove()
    boton_agregar_con5.grid_remove()
    boton_quitar_con4.grid_remove()
    con4_ec_entry.grid_remove()

    buscarcedulac5_combo.grid_remove()
    button11.grid_remove()
    cedulacomprador5_entry.grid_remove()
    con5_primerApellido_entry.grid_remove()
    con5_segundoApellido_entry.grid_remove()
    con5_primerNombre_entry.grid_remove()
    con5_segundoNombre_entry.grid_remove()
    con5_sexo_entry.grid_remove()
    con5_fecha_nacimiento.grid_remove()
    con5_rh_entry.grid_remove()
    con5_domicilio_entry.grid_remove()
    con5_estadocivil_entry.grid_remove()
    boton_agregar_con6.grid_remove()
    boton_quitar_con5.grid_remove()
    con5_ec_entry.grid_remove()

    buscarcedulac6_combo.grid_remove()
    button12.grid_remove()
    cedulacomprador6_entry.grid_remove()
    con6_primerApellido_entry.grid_remove()
    con6_segundoApellido_entry.grid_remove()
    con6_primerNombre_entry.grid_remove()
    con6_segundoNombre_entry.grid_remove()
    con6_sexo_entry.grid_remove()
    con6_fecha_nacimiento.grid_remove()
    con6_rh_entry.grid_remove()
    con6_domicilio_entry.grid_remove()
    con6_estadocivil_entry.grid_remove()
    boton_quitar_con6.grid_remove()
    con6_ec_entry.grid_remove()


    
    
def ingresarcomprador2a():
    buscarcedulac1_combo.grid()
    button7.grid_remove()
    cedulacomprador_entry.grid()
    primer_apellido2_entry.grid()
    segundo_apellido2_entry.grid()
    primer_nombre2_entry.grid()
    segundo_nombre2_entry.grid()
    sexo2_entry.grid()
    fecha_nacimiento2.grid()
    rh2_entry.grid()
    domicilio2_entry.grid()
    estadocivil2_entry.grid()
    ecc_entry.grid()

    buscarcedulac2_combo.grid()
    button8.grid()        
    cedulacomprador2_entry.grid()
    con2_primerApellido_entry.grid()
    con2_segundoApellido_entry.grid()
    con2_primerNombre_entry.grid()
    con2_segundoNombre_entry.grid()
    con2_sexo_entry.grid()
    con2_fecha_nacimiento.grid()
    con2_rh_entry.grid()
    con2_domicilio_entry.grid()
    con2_estadocivil_entry.grid()
    boton_agregar_con2.grid_remove()
    boton_agregar_con3.grid()
    boton_quitar_con2.grid()
    con2_ec_entry.grid()

    buscarcedulac3_combo.grid_remove()
    button9.grid_remove()
    cedulacomprador3_entry.grid_remove()
    con3_primerApellido_entry.grid_remove()
    con3_segundoApellido_entry.grid_remove()
    con3_primerNombre_entry.grid_remove()
    con3_segundoNombre_entry.grid_remove()
    con3_sexo_entry.grid_remove()
    con3_fecha_nacimiento.grid_remove()
    con3_rh_entry.grid_remove()
    con3_domicilio_entry.grid_remove()
    con3_estadocivil_entry.grid_remove()
    boton_agregar_con4.grid_remove()
    boton_quitar_con3.grid_remove()
    con3_ec_entry.grid_remove()


    buscarcedulac4_combo.grid_remove()
    button10.grid_remove()
    cedulacomprador4_entry.grid_remove()
    con4_primerApellido_entry.grid_remove()
    con4_segundoApellido_entry.grid_remove()
    con4_primerNombre_entry.grid_remove()
    con4_segundoNombre_entry.grid_remove()
    con4_sexo_entry.grid_remove()
    con4_fecha_nacimiento.grid_remove()
    con4_rh_entry.grid_remove()
    con4_domicilio_entry.grid_remove()
    con4_estadocivil_entry.grid_remove()
    boton_agregar_con5.grid_remove()
    boton_quitar_con4.grid_remove()
    con4_ec_entry.grid_remove()

    buscarcedulac5_combo.grid_remove()
    button11.grid_remove()
    cedulacomprador5_entry.grid_remove()
    con5_primerApellido_entry.grid_remove()
    con5_segundoApellido_entry.grid_remove()
    con5_primerNombre_entry.grid_remove()
    con5_segundoNombre_entry.grid_remove()
    con5_sexo_entry.grid_remove()
    con5_fecha_nacimiento.grid_remove()
    con5_rh_entry.grid_remove()
    con5_domicilio_entry.grid_remove()
    con5_estadocivil_entry.grid_remove()
    boton_agregar_con6.grid_remove()
    boton_quitar_con5.grid_remove()
    con5_ec_entry.grid_remove()

    buscarcedulac6_combo.grid_remove()
    button12.grid_remove()
    cedulacomprador6_entry.grid_remove()
    con6_primerApellido_entry.grid_remove()
    con6_segundoApellido_entry.grid_remove()
    con6_primerNombre_entry.grid_remove()
    con6_segundoNombre_entry.grid_remove()
    con6_sexo_entry.grid_remove()
    con6_fecha_nacimiento.grid_remove()
    con6_rh_entry.grid_remove()
    con6_domicilio_entry.grid_remove()
    con6_estadocivil_entry.grid_remove()
    boton_quitar_con6.grid_remove()
    con6_ec_entry.grid_remove()

def ingresarcomprador3a():
    buscarcedulac1_combo.grid()
    button7.grid_remove()
    cedulacomprador_entry.grid()
    primer_apellido2_entry.grid()
    segundo_apellido2_entry.grid()
    primer_nombre2_entry.grid()
    segundo_nombre2_entry.grid()
    sexo2_entry.grid()
    fecha_nacimiento2.grid()
    rh2_entry.grid()
    domicilio2_entry.grid()
    estadocivil2_entry.grid()
    ecc_entry.grid() 

    buscarcedulac2_combo.grid()
    button8.grid_remove()
    cedulacomprador2_entry.grid()
    con2_primerApellido_entry.grid()
    con2_segundoApellido_entry.grid()
    con2_primerNombre_entry.grid()
    con2_segundoNombre_entry.grid()
    con2_sexo_entry.grid()
    con2_fecha_nacimiento.grid()
    con2_rh_entry.grid()
    con2_domicilio_entry.grid()
    con2_estadocivil_entry.grid()
    boton_agregar_con2.grid_remove()
    con2_ec_entry.grid()

    buscarcedulac3_combo.grid()
    button9.grid()    
    cedulacomprador3_entry.grid()
    con3_primerApellido_entry.grid()
    con3_segundoApellido_entry.grid()
    con3_primerNombre_entry.grid()
    con3_segundoNombre_entry.grid()
    con3_sexo_entry.grid()
    con3_fecha_nacimiento.grid()
    con3_rh_entry.grid()
    con3_domicilio_entry.grid()
    con3_estadocivil_entry.grid()
    boton_agregar_con4.grid()
    boton_quitar_con3.grid()
    boton_agregar_con3.grid_remove()
    boton_quitar_con2.grid_remove()
    con3_ec_entry.grid()    

    buscarcedulac4_combo.grid_remove()
    button10.grid_remove()
    cedulacomprador4_entry.grid_remove()
    con4_primerApellido_entry.grid_remove()
    con4_segundoApellido_entry.grid_remove()
    con4_primerNombre_entry.grid_remove()
    con4_segundoNombre_entry.grid_remove()
    con4_sexo_entry.grid_remove()
    con4_fecha_nacimiento.grid_remove()
    con4_rh_entry.grid_remove()
    con4_domicilio_entry.grid_remove()
    con4_estadocivil_entry.grid_remove()
    boton_agregar_con5.grid_remove()
    boton_quitar_con4.grid_remove()
    con4_ec_entry.grid_remove()

    buscarcedulac5_combo.grid_remove()
    button11.grid_remove()
    cedulacomprador5_entry.grid_remove()
    con5_primerApellido_entry.grid_remove()
    con5_segundoApellido_entry.grid_remove()
    con5_primerNombre_entry.grid_remove()
    con5_segundoNombre_entry.grid_remove()
    con5_sexo_entry.grid_remove()
    con5_fecha_nacimiento.grid_remove()
    con5_rh_entry.grid_remove()
    con5_domicilio_entry.grid_remove()
    con5_estadocivil_entry.grid_remove()
    boton_agregar_con6.grid_remove()
    boton_quitar_con5.grid_remove()
    con5_ec_entry.grid_remove()


    buscarcedulac6_combo.grid_remove()
    button12.grid_remove()
    cedulacomprador6_entry.grid_remove()
    con6_primerApellido_entry.grid_remove()
    con6_segundoApellido_entry.grid_remove()
    con6_primerNombre_entry.grid_remove()
    con6_segundoNombre_entry.grid_remove()
    con6_sexo_entry.grid_remove()
    con6_fecha_nacimiento.grid_remove()
    con6_rh_entry.grid_remove()
    con6_domicilio_entry.grid_remove()
    con6_estadocivil_entry.grid_remove()
    boton_quitar_con6.grid_remove()
    con6_ec_entry.grid_remove()


    

def ingresarcomprador4a():
    buscarcedulac1_combo.grid()
    button7.grid_remove()
    cedulacomprador_entry.grid()
    primer_apellido2_entry.grid()
    segundo_apellido2_entry.grid()
    primer_nombre2_entry.grid()
    segundo_nombre2_entry.grid()
    sexo2_entry.grid()
    fecha_nacimiento2.grid()
    rh2_entry.grid()
    domicilio2_entry.grid()
    estadocivil2_entry.grid()
    ecc_entry.grid()


    buscarcedulac2_combo.grid()
    button8.grid_remove()        
    cedulacomprador2_entry.grid()
    con2_primerApellido_entry.grid()
    con2_segundoApellido_entry.grid()
    con2_primerNombre_entry.grid()
    con2_segundoNombre_entry.grid()
    con2_sexo_entry.grid()
    con2_fecha_nacimiento.grid()
    con2_rh_entry.grid()
    con2_domicilio_entry.grid()
    con2_estadocivil_entry.grid()
    boton_agregar_con2.grid_remove()
    con2_ec_entry.grid()


    buscarcedulac3_combo.grid()
    button9.grid_remove()    
    cedulacomprador3_entry.grid()
    con3_primerApellido_entry.grid()
    con3_segundoApellido_entry.grid()
    con3_primerNombre_entry.grid()
    con3_segundoNombre_entry.grid()
    con3_sexo_entry.grid()
    con3_fecha_nacimiento.grid()
    con3_rh_entry.grid()
    con3_domicilio_entry.grid()
    con3_estadocivil_entry.grid()
    boton_agregar_con3.grid_remove()
    boton_quitar_con2.grid_remove()
    con3_ec_entry.grid()


    buscarcedulac4_combo.grid()
    button10.grid()
    cedulacomprador4_entry.grid()
    con4_primerApellido_entry.grid()
    con4_segundoApellido_entry.grid()
    con4_primerNombre_entry.grid()
    con4_segundoNombre_entry.grid()
    con4_sexo_entry.grid()
    con4_fecha_nacimiento.grid()
    con4_rh_entry.grid()
    con4_domicilio_entry.grid()
    con4_estadocivil_entry.grid()
    boton_agregar_con5.grid()
    boton_quitar_con4.grid()
    boton_agregar_con4.grid_remove()
    boton_quitar_con3.grid_remove()
    con4_ec_entry.grid()

    buscarcedulac5_combo.grid_remove()
    button11.grid_remove()
    cedulacomprador5_entry.grid_remove()
    con5_primerApellido_entry.grid_remove()
    con5_segundoApellido_entry.grid_remove()
    con5_primerNombre_entry.grid_remove()
    con5_segundoNombre_entry.grid_remove()
    con5_sexo_entry.grid_remove()
    con5_fecha_nacimiento.grid_remove()
    con5_rh_entry.grid_remove()
    con5_domicilio_entry.grid_remove()
    con5_estadocivil_entry.grid_remove()
    boton_agregar_con6.grid_remove()
    boton_quitar_con5.grid_remove()
    con5_ec_entry.grid_remove()


    buscarcedulac6_combo.grid_remove()
    button12.grid_remove()
    cedulacomprador6_entry.grid_remove()
    con6_primerApellido_entry.grid_remove()
    con6_segundoApellido_entry.grid_remove()
    con6_primerNombre_entry.grid_remove()
    con6_segundoNombre_entry.grid_remove()
    con6_sexo_entry.grid_remove()
    con6_fecha_nacimiento.grid_remove()
    con6_rh_entry.grid_remove()
    con6_domicilio_entry.grid_remove()
    con6_estadocivil_entry.grid_remove()
    boton_quitar_con6.grid_remove()
    con6_ec_entry.grid_remove()



    

def ingresarcomprador5a():

    buscarcedulac1_combo.grid()
    button7.grid_remove()
    cedulacomprador_entry.grid()
    primer_apellido2_entry.grid()
    segundo_apellido2_entry.grid()
    primer_nombre2_entry.grid()
    segundo_nombre2_entry.grid()
    sexo2_entry.grid()
    fecha_nacimiento2.grid()
    rh2_entry.grid()
    domicilio2_entry.grid()
    estadocivil2_entry.grid()
    ecc_entry.grid()


    buscarcedulac2_combo.grid()
    button8.grid_remove()        
    cedulacomprador2_entry.grid()
    con2_primerApellido_entry.grid()
    con2_segundoApellido_entry.grid()
    con2_primerNombre_entry.grid()
    con2_segundoNombre_entry.grid()
    con2_sexo_entry.grid()
    con2_fecha_nacimiento.grid()
    con2_rh_entry.grid()
    con2_domicilio_entry.grid()
    con2_estadocivil_entry.grid()
    boton_agregar_con2.grid_remove()
    con2_ec_entry.grid()


    buscarcedulac3_combo.grid()
    button9.grid_remove()    
    cedulacomprador3_entry.grid()
    con3_primerApellido_entry.grid()
    con3_segundoApellido_entry.grid()
    con3_primerNombre_entry.grid()
    con3_segundoNombre_entry.grid()
    con3_sexo_entry.grid()
    con3_fecha_nacimiento.grid()
    con3_rh_entry.grid()
    con3_domicilio_entry.grid()
    con3_estadocivil_entry.grid()
    boton_agregar_con3.grid_remove()
    boton_quitar_con2.grid_remove()
    con3_ec_entry.grid()


    buscarcedulac4_combo.grid()
    button10.grid_remove()
    cedulacomprador4_entry.grid()
    con4_primerApellido_entry.grid()
    con4_segundoApellido_entry.grid()
    con4_primerNombre_entry.grid()
    con4_segundoNombre_entry.grid()
    con4_sexo_entry.grid()
    con4_fecha_nacimiento.grid()
    con4_rh_entry.grid()
    con4_domicilio_entry.grid()
    con4_estadocivil_entry.grid()
    boton_agregar_con5.grid()
    boton_quitar_con4.grid()
    boton_agregar_con4.grid_remove()
    boton_quitar_con3.grid_remove()
    con4_ec_entry.grid()



    buscarcedulac5_combo.grid()
    button11.grid()
    cedulacomprador5_entry.grid()
    con5_primerApellido_entry.grid()
    con5_segundoApellido_entry.grid()
    con5_primerNombre_entry.grid()
    con5_segundoNombre_entry.grid()
    con5_sexo_entry.grid()
    con5_fecha_nacimiento.grid()
    con5_rh_entry.grid()
    con5_domicilio_entry.grid()
    con5_estadocivil_entry.grid()
    boton_agregar_con6.grid()
    boton_quitar_con5.grid()
    boton_agregar_con5.grid_remove()
    boton_quitar_con4.grid_remove()
    con5_ec_entry.grid()


    buscarcedulac6_combo.grid_remove()
    button12.grid_remove()
    cedulacomprador6_entry.grid_remove()
    con6_primerApellido_entry.grid_remove()
    con6_segundoApellido_entry.grid_remove()
    con6_primerNombre_entry.grid_remove()
    con6_segundoNombre_entry.grid_remove()
    con6_sexo_entry.grid_remove()
    con6_fecha_nacimiento.grid_remove()
    con6_rh_entry.grid_remove()
    con6_domicilio_entry.grid_remove()
    con6_estadocivil_entry.grid_remove()
    boton_quitar_con6.grid_remove()
    con6_ec_entry.grid_remove()


    

def ingresarcomprador6a():
    buscarcedulac1_combo.grid()
    button7.grid_remove()
    cedulacomprador_entry.grid()
    primer_apellido2_entry.grid()
    segundo_apellido2_entry.grid()
    primer_nombre2_entry.grid()
    segundo_nombre2_entry.grid()
    sexo2_entry.grid()
    fecha_nacimiento2.grid()
    rh2_entry.grid()
    domicilio2_entry.grid()
    estadocivil2_entry.grid()
    ecc_entry.grid()


    buscarcedulac2_combo.grid()
    button8.grid_remove()        
    cedulacomprador2_entry.grid()
    con2_primerApellido_entry.grid()
    con2_segundoApellido_entry.grid()
    con2_primerNombre_entry.grid()
    con2_segundoNombre_entry.grid()
    con2_sexo_entry.grid()
    con2_fecha_nacimiento.grid()
    con2_rh_entry.grid()
    con2_domicilio_entry.grid()
    con2_estadocivil_entry.grid()
    boton_agregar_con2.grid_remove()
    con2_ec_entry.grid()

    buscarcedulac3_combo.grid()
    button9.grid_remove()    
    cedulacomprador3_entry.grid()
    con3_primerApellido_entry.grid()
    con3_segundoApellido_entry.grid()
    con3_primerNombre_entry.grid()
    con3_segundoNombre_entry.grid()
    con3_sexo_entry.grid()
    con3_fecha_nacimiento.grid()
    con3_rh_entry.grid()
    con3_domicilio_entry.grid()
    con3_estadocivil_entry.grid()
    boton_agregar_con3.grid_remove()
    boton_quitar_con2.grid_remove()
    con3_ec_entry.grid()


    buscarcedulac4_combo.grid()
    button10.grid_remove()
    cedulacomprador4_entry.grid()
    con4_primerApellido_entry.grid()
    con4_segundoApellido_entry.grid()
    con4_primerNombre_entry.grid()
    con4_segundoNombre_entry.grid()
    con4_sexo_entry.grid()
    con4_fecha_nacimiento.grid()
    con4_rh_entry.grid()
    con4_domicilio_entry.grid()
    con4_estadocivil_entry.grid()
    con4_ec_entry.grid()


    boton_quitar_con4.grid()
    boton_agregar_con4.grid_remove()
    boton_quitar_con3.grid_remove()
    
    buscarcedulac5_combo.grid()
    button11.grid_remove()
    cedulacomprador5_entry.grid()
    con5_primerApellido_entry.grid()
    con5_segundoApellido_entry.grid()
    con5_primerNombre_entry.grid()
    con5_segundoNombre_entry.grid()
    con5_sexo_entry.grid()
    con5_fecha_nacimiento.grid()
    con5_rh_entry.grid()
    con5_domicilio_entry.grid()
    con5_estadocivil_entry.grid()
    boton_agregar_con5.grid_remove()
    boton_quitar_con4.grid_remove()
    con5_ec_entry.grid()


    buscarcedulac6_combo.grid()
    button12.grid()
    cedulacomprador6_entry.grid()
    con6_primerApellido_entry.grid()
    con6_segundoApellido_entry.grid()
    con6_primerNombre_entry.grid()
    con6_segundoNombre_entry.grid()
    con6_sexo_entry.grid()
    con6_fecha_nacimiento.grid()
    con6_rh_entry.grid()
    con6_domicilio_entry.grid()
    con6_estadocivil_entry.grid()
    boton_quitar_con6.grid()
    boton_agregar_con6.grid_remove()
    boton_agregar_con5.grid_remove()
    boton_quitar_con5.grid_remove()
    con6_ec_entry.grid()    

   

def seleccionar_widget():
    chip_entrada=entry_chip.get()
    entryv1_valor=cedulavendedor_entry.get() 
    entryv2_valor=cedulavendedor2_entry.get() 
    entryv3_valor=cedulavendedor3_entry.get()
    entryv4_valor=cedulavendedor4_entry.get()
    entryv5_valor=cedulavendedor5_entry.get()
    entryv6_valor=cedulavendedor6_entry.get()

    entryc1_valor=cedulacomprador_entry.get()
    entryc2_valor=cedulacomprador2_entry.get() 
    entryc3_valor=cedulacomprador3_entry.get()
    entryc4_valor=cedulacomprador4_entry.get()
    entryc5_valor=cedulacomprador5_entry.get()
    entryc6_valor=cedulacomprador6_entry.get()

    if chip_entrada !="":
        entry_chip.grid()
        label_chip.grid()
    
    else:
        entry_chip.grid_remove()
        label_chip.grid_remove()
    
    if entryv1_valor !="" and entryv2_valor =="" and entryv3_valor =="" and entryv4_valor =="" and entryv5_valor =="" and entryv6_valor =="":
        ingresarvendedor1a()

    
    
    if entryv1_valor !="" and entryv2_valor !="" and entryv3_valor =="" and entryv4_valor =="" and entryv5_valor =="" and entryv6_valor =="":
        ingresarvendedor2a()
    
    
    
    
    if entryv1_valor !="" and entryv2_valor !="" and entryv3_valor !="" and entryv4_valor =="" and entryv5_valor =="" and entryv6_valor =="":
        ingresarvendedor3a()
    
    
    
    if entryv1_valor !="" and entryv2_valor !="" and entryv3_valor !="" and entryv4_valor !="" and entryv5_valor =="" and entryv6_valor =="":
        ingresarvendedor4a()
    
    
    
    if entryv1_valor !="" and entryv2_valor !="" and entryv3_valor !="" and entryv4_valor !="" and entryv5_valor !="" and entryv6_valor =="":
        ingresarvendedor5a()
    
    
    
    if entryv1_valor !="" and entryv2_valor !="" and entryv3_valor !="" and entryv4_valor !="" and entryv5_valor !="" and entryv6_valor !="":
        ingresarvendedor6a()
    
    if entryc1_valor !="" and entryc2_valor =="" and entryc3_valor =="" and entryc4_valor =="" and entryc5_valor =="" and entryc6_valor =="":
        ingresarcomprador1a()

    if entryc1_valor !="" and entryc2_valor !="" and entryc3_valor =="" and entryc4_valor =="" and entryc5_valor =="" and entryc6_valor =="":
        ingresarcomprador2a()
    
    if entryc1_valor !="" and entryc2_valor !="" and entryc3_valor !="" and entryc4_valor =="" and entryc5_valor =="" and entryc6_valor =="":
        ingresarcomprador3a()
       
   
    
    if entryc1_valor !="" and entryc2_valor !="" and entryc3_valor !="" and entryc4_valor !="" and entryc5_valor =="" and entryc6_valor =="":
        ingresarcomprador4a()
    
    
    
    if entryc1_valor !="" and entryc2_valor !="" and entryc3_valor !="" and entryc4_valor !="" and entryc5_valor !="" and entryc6_valor =="":
        ingresarcomprador5a()
    
    
    
    if entryc1_valor !="" and entryc2_valor !="" and entryc3_valor !="" and entryc4_valor !="" and entryc5_valor !="" and entryc6_valor !="":
        ingresarcomprador6a()
    
     


def automatizacion():
    doc = DocxTemplate("VENTANUDAPROPIEDADSINRESERVA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],
            'pcc' : fila["puntocomaCC"],




            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],  

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],
            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],
                      
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escrituras.docx")
        documento = Document(f"archivos/escrituras.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [34, 49, 51, 53, 55, 56, 57, 58, 61] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta)

def automatizacion1():
    doc = DocxTemplate("VENTADERECHODECUOTAENNUDAPROPIEDADSINRESERVA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"], 

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"], 
            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],          
                 

                 
        } 
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escrituras.docx")
        documento = Document(f"archivos/escrituras.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [34, 53, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta)


def automatizacion2():
    doc = DocxTemplate("VENTANUDAPROPIEDADCONRESERVA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/Venta_nuda_propiedad_sin_reserva.docx")
        documento = Document(f"archivos/Venta_nuda_propiedad_sin_reserva.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [34, 53, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion3():
    doc = DocxTemplate("VENTAPLENAPROPIEDAD.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],

            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion4():
    doc = DocxTemplate("VENTADERECHODECUOTAPLENAPROPIEDAD.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion5():
    doc = DocxTemplate("VENTAPOSESION.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion6():
    doc = DocxTemplate("VENTANUDAPROPIEDADENPOSESION.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion7():
    doc = DocxTemplate("VENTADERECHODECUOTAENPOSESION.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 


def automatizacion8():
    doc = DocxTemplate("VENTADERECHODECUOTANUDAPROPIEDADENPOSESION.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],

            


            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 


def automatizacion9():
    doc = DocxTemplate("VENTADERECHOSHERENCIALESVINCULADOSADQUIRIDOSPORCOMPRA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 


def automatizacion10():
    doc = DocxTemplate("VENTADERECHOSHERENCIALESVINCULADOSCOMOHEREDERO.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion11():
    doc = DocxTemplate("VENTA_E_HIPOTECA_CORPORACION_SOCIAL_DE_CUNDINAMARCA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

                
            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion12():
    doc = DocxTemplate("VENTA_E_HIPOTECA_BCSC.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],
            





            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion13():
    doc = DocxTemplate("VENTA_E_HIPOTECA_BANCO_DE_BOGOTA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion14():
    doc = DocxTemplate("VENTA_E_HIPOTECA_BANCO_AGRARIO.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],
            
            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],



            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion15():
    doc = DocxTemplate("RESCILIACION_ DE_ESCRITURA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],


            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion16():
    doc = DocxTemplate("PODER_GENERAL.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],

            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion17():
    doc = DocxTemplate("DECLARACION_DE_UNION_MARITAL_Y_SOCIEDAD_PATRIMONIAL.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion18():
    doc = DocxTemplate("DECLARACION_DE_CONSTRUCCION_PARTICULARES.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion19():
    doc = DocxTemplate("DECLARACION_DE_CONSTRUCCION_DE_COLSUBSIDIO.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion20():
    doc = DocxTemplate("DACION_EN_PAGO.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion21():
    doc = DocxTemplate("CONSTITUCION_HIPOTECA_BANCO_DE_BOGOTA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion22():
    doc = DocxTemplate("CONSTITUCION_HIPOTECA_BANCO_AGRARIO.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion23():
    doc = DocxTemplate("CONSTITUCION_HIPOTECA_ABIERTA_PERSONA_NATURAL.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion24():
    doc = DocxTemplate("COMPRAVENTA_CON_PACTO_DE_RETROVENTA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],



            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion25():
    doc = DocxTemplate("CANCELACION_DE_USUFRUCTO.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion26():
    doc = DocxTemplate("CANCELACION_AFECTACION_A_VIVIENDA_FAMILIAR_AFECTO_UNO.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion27():
    doc = DocxTemplate("CANCELACION_AFECTACION_A_VIVIENDA_FAMILIAR_AFECTARON_AMBOS.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion28():
    doc = DocxTemplate("AUTORIZACION_SALIDA_DEL_PAIS.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion29():
    doc = DocxTemplate("AFECTACION _A_ VIVIENDA FAMILIAR_PROPIEDAD_DE_UNO.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion30():
    doc = DocxTemplate("AFECTACION _A_ VIVIENDA FAMILIAR_PROPIEDAD_DE_DOS.docx")
    
    

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion31():
    doc = DocxTemplate("ACTUALIZACION_DE_AREA.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion32():
    doc = DocxTemplate("ACTUALIZACION_CEDULA_CATASTRAL.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 

def automatizacion33():
    doc = DocxTemplate("ACLARACION.docx")
    
   
      

    
    bd=sqlite3.connect("login1.db")
    df=pd.read_sql("SELECT * FROM datos_escrituras1 WHERE ID="+ID_registro.get(), bd)

    for index,  fila in df.iterrows():
        datos = { 
            'numero_escritura' : fila["Escritura"] ,
            'letra_escritura' : fila["EscrLetras"],
            'matricula' : fila["Matricula"], 
            'cd_catastral' : fila["Cedula_catastral"],
            'ubicacion' : fila["Ubicacion"],
            'municipio' : fila["Municipio"], 
            'departamento' : fila["Departamento"],
            'rural_urbano' : fila["Rural_urbano"],
            'precio' : fila["precio1"],
            'precioletras' : fila["precioLetras"],
            'direccion' : fila["Direccion"],
            'notaria' : fila["Notaria"],
            'munota' : fila["Municipio_notaria"],
            'notario' : fila["Notario"],
            'modadquirir' : fila["Modo_Adquirir"],


            'dia' : fila["Dia"],
            'mes' : fila["Mes"],
            'año' : fila["Año"],
            'añoletras' : fila["AñoLetras"],

            'cedula_ven1' : fila["cedulavendedor1"],
            'pnv1' : fila["primerNombreVen1"],
            'snv1' : fila["segundoNombre_Ven1"],
            'pav1' : fila["primerApellidoVen1"],
            'sav1' : fila["segundoApellidoVen1"],
            'sexov1' : fila["sexoVen1"],
            'domv1' : fila["domicilio_Ven1"],
            'estv1' : fila["estadocivil_ven1"],

            'cedula_ven2' : fila["cedulavendedor2"],
            'pnv2' : fila["primerNombreVen2"],
            'snv2' : fila["segundoNombre_Ven2"],
            'pav2' : fila["primerApellidoVen2"],
            'sav2' : fila["segundoApellidoVen2"],
            'sexov2' : fila["sexoVen2"],
            'domv2' : fila["domicilio_Ven2"],
            'estv2' : fila["estadocivil_ven2"],

            'cedula_ven3' : fila["cedulavendedor3"],
            'pnv3' : fila["primerNombreVen3"],
            'snv3' : fila["segundoNombre_Ven3"],
            'pav3' : fila["primerApellidoVen3"],
            'sav3' : fila["segundoApellidoVen3"],
            'sexov3' : fila["sexoVen3"],
            'domv3' : fila["domicilio_Ven3"],
            'estv3' : fila["estadocivil_ven3"],

            'cedula_ven4' : fila["cedulavendedor4"],
            'pnv4' : fila["primerNombreVen4"],
            'snv4' : fila["segundoNombre_Ven4"],
            'sav4' : fila["segundoApellidoVen4"],
            'sexov4' : fila["sexoVen4"],
            'domv4' : fila["domicilio_Ven4"],
            'estv4' : fila["estadocivil_ven4"],

            'cedula_ven5' : fila["cedulavendedor5"],
            'pnv5' : fila["primerNombreVen5"],
            'snv5' : fila["segundoNombre_Ven5"],
            'pav5' : fila["primerApellidoVen5"],
            'sav5' : fila["segundoApellidoVen5"],
            'sexov5' : fila["sexoVen5"],
            'domv5' : fila["domicilio_Ven5"],
            'estv5' : fila["estadocivil_ven5"],

            'cedula_ven6' : fila["cedulavendedor6"],
            'pnv6' : fila["primerNombreVen6"],
            'snv6' : fila["segundoNombre_Ven6"],
            'pav6' : fila["primerApellidoVen6"],
            'sav6' : fila["segundoApellidoVen6"],
            'sexov6' : fila["sexoVen6"],
            'domv6' : fila["domicilio_Ven6"],
            'estv6' : fila["estadocivil_ven6"],

            'cedula_comp1' : fila["cedula_comprador1"],
            'pnc1' : fila["primerNombre_comp1"],
            'snc1' : fila["segundoNombre_comp1"],
            'pac1' : fila["primerApellido_comp1"],
            'sac1' : fila["segundoApellido_comp1"] ,
            'sexoc1' : fila["sexo_comp1"],
            'domc1' : fila["domicilio_comp1"],
            'estc1' : fila["estadocivil_comp1"], 

            'cedula_comp2' : fila["cedula_comprador2"],
            'pnc2' : fila["primerNombre_comp2"],
            'snc2' : fila["segundoNombre_comp2"],
            'pac2' : fila["primerApellido_comp2"],
            'sac2' : fila["segundoApellido_comp2"],
            'sexoc2' : fila["sexo_comp2"],
            'domc2' : fila["domicilio_comp2"],
            'estc2' : fila["estadocivil_comp2"],

            'cedula_comp3' : fila["cedula_comprador3"],
            'pnc3' : fila["primerNombre_comp3"],
            'snc3' : fila["segundoNombre_comp3"],
            'pac3' : fila["primerApellido_comp3"],
            'sac3' : fila["segundoApellido_comp3"],
            'sexoc3' : fila["sexo_comp3"],
            'domc3' : fila["domicilio_comp3"],
            'estc3' : fila["estadocivil_comp3"],

            'cedula_comp4' : fila["cedula_comprador4"],
            'pnc4' : fila["primerNombre_comp4"],
            'snc4' : fila["segundoNombre_comp4"],
            'pac4' : fila["primerApellido_comp4"],
            'sac4' : fila["segundoApellido_comp4"],
            'sexoc4' : fila["sexo_comp4"],
            'domc4' : fila["domicilio_comp4"],
            'estc4' : fila["estadocivil_comp4"],

            'cedula_comp5' : fila["cedula_comprador5"],
            'pnc5' : fila["primerNombre_comp5"],
            'snc5' : fila["segundoNombre_comp5"],
            'pac5' : fila["primerApellido_comp5"],
            'sac5' : fila["segundoApellido_comp5"],
            'sexoc5' : fila["sexo_comp5"],
            'domc5' : fila["domicilio_comp5"],
            'estc5' : fila["estadocivil_comp5"],

            'cedula_comp6' : fila["cedula_comprador6"],
            'pnc6' : fila["primerNombre_comp6"],
            'snc6' : fila["segundoNombre_comp6"],
            'pac6' : fila["primerApellido_comp6"],
            'sac6' : fila["segundoApellido_comp6"],
            'sexoc6' : fila["sexo_comp6"],
            'domc6' : fila["domicilio_comp6"],
            'estc6' : fila["estadocivil_comp6"],


            

            'CCV1' : fila["CC_V1"],
            'CCV2' : fila["CC_V2"],
            'CCV3' : fila["CC_V3"],
            'CCV4' : fila["CC_V4"],
            'CCV5' : fila["CC_V5"],
            'CCV6' : fila["CC_V6"],

            'CCC1' : fila["CC_C1"],
            'CCC2' : fila["CC_C2"],
            'CCC3' : fila["CC_C3"],
            'CCC4' : fila["CC_C4"],
            'CCC5' : fila["CC_C5"],
            'CCC6' : fila["CC_C6"],

            'GV1' : fila["GENERV1"],
            'GV2' : fila["GENERV2"],
            'GV3' : fila["GENERV3"],
            'GV4' : fila["GENERV4"],
            'GV5' : fila["GENERV5"],
            'GV6' : fila["GENERV6"],

            'GC1' : fila["GENERC1"],
            'GC2' : fila["GENERC2"],
            'GC3' : fila["GENERC3"],
            'GC4' : fila["GENERC4"],
            'GC5' : fila["GENERC5"],
            'GC6' : fila["GENERC6"],



            'domiV1' : fila["domiciliado_domiciliada1"],
            'domiV2' : fila["domiciliado_domiciliada2"],
            'domiV3' : fila["domiciliado_domiciliada3"],
            'domiV4' : fila["domiciliado_domiciliada4"],
            'domiV5' : fila["domiciliado_domiciliada5"],
            'domiV6' : fila["domiciliado_domiciliada6"],

            'domiC1' : fila["domiciliado_domiciliadaC1"],
            'domiC2' : fila["domiciliado_domiciliadaC2"],
            'domiC3' : fila["domiciliado_domiciliadaC3"],
            'domiC4' : fila["domiciliado_domiciliadaC4"],
            'domiC5' : fila["domiciliado_domiciliadaC5"],
            'domiC6' : fila["domiciliado_domiciliadaC6"],

            'comV1' : fila["comaV1"],
            'comV2' : fila["comaV2"],
            'comV3' : fila["comaV3"],
            'comV4' : fila["comaV4"],
            'comV5' : fila["comaV5"],
            'comV6' : fila["comaV6"],

            'comC1' : fila["comaC1"],
            'comC2' : fila["comaC2"],
            'comC3' : fila["comaC3"],
            'comC4' : fila["comaC4"],
            'comC5' : fila["comaC5"],
            'comC6' : fila["comaC6"],

            'pcom2' : fila["puntoycoma2"],
            'pcom3' : fila["puntoycoma3"],
            'pcom4' : fila["puntoycoma4"],
            'pcom5' : fila["puntoycoma5"],
            'pcom6' : fila["puntoycoma6"],

            'pcomc2' : fila["puntoycomaC2"],
            'pcomc3' : fila["puntoycomaC3"],
            'pcomc4' : fila["puntoycomaC4"],
            'pcomc5' : fila["puntoycomaC5"],
            'pcomc6' : fila["puntoycomaC6"],



            'ecv1' : fila["ECV1"],
            'ecv2' : fila["ECV2"],
            'ecv3' : fila["ECV3"],
            'ecv4' : fila["ECV4"],
            'ecv5' : fila["ECV5"],
            'ecv6' : fila["ECV6"],

            'ecc1' : fila["ECC1"],
            'ecc2' : fila["ECC2"],
            'ecc3' : fila["ECC3"],
            'ecc4' : fila["ECC4"],
            'ecc5' : fila["ECC5"],
            'ecc6' : fila["ECC6"],

            'hv1' : fila["Hechosv1"],
            'hv2' : fila["Hechosv2"],
            'hv3' : fila["Hechosv3"],
            'hv4' : fila["Hechosv4"],
            'hv5' : fila["Hechosv5"],
            'hv6' : fila["Hechosv6"],

            'hc1' : fila["Hechosc1"],
            'hc2' : fila["Hechosc2"],
            'hc3' : fila["Hechosc3"],
            'hc4' : fila["Hechosc4"],
            'hc5' : fila["Hechosc5"],
            

            'dv' : fila["despuessingpluralvendedores"],
            'dc' : fila["despuessingcompradores"],
            'spv' : fila["SingularPluralVendedores1AA1"],
            'spc' : fila["SingularPluralCompradores1A1"],
            'tspv' : fila["SingularPluralVendedores1AA"],
            'tspc' : fila["SingularPluralCompradores1A"],

            'obv' : fila["obrandovendedores"],
            'obc' : fila["obrandocompradores"],
            'trs' : fila["TRANSFER"],
            'grs' : fila["garantizar1"],
            'etg' : fila["entregar1"],
            'prt' : fila["presentesCompradores1"],

            
            'manV' : fila["manifiestoV1x"],
            'manC' : fila["manifiestoC1xxx"],

            'manVA' : fila["manifiestoV1A"],
            'manCA' : fila["manifiestoC1A"],
            'nudo' : fila["nudosC"],

            'usu' : fila["usufructo"],
            'usu2' : fila["usufructo2"],
            'sella' : fila["ellaellosellas"],
            'hayc' : fila["hayanC"],
            'hayv' : fila["hayanV"],
            'idtc' : fila["identificarC"],


            'mc' : fila["Cmanifiestan1"],

            'da' : fila["dia_actual"],
            'ma' : fila["mes_actual"],
            'aa' : fila["año_actual"],
            'rvc' : fila["RelacionVenCom1"],

            'acl' : fila["año_letra"],
            'dcl' : fila["dia_letra"],
            'mcl' : fila["mes1_letras"],

            'ne' : fila["numeroescr"],
            'nel' : fila["numeroescr_letras"],
            
            'dirv1' : fila["direccionv1"],
            'dirv2' : fila["direccionv2"],
            'dirv3' : fila["direccionv3"],
            'dirv4' : fila["direccionv4"],
            'dirv5' : fila["direccionv5"],
            'dirv6' : fila["direccionv6"],

            'telv1' : fila["telefonov1"],
            'telv2' : fila["telefonov2"],
            'telv3' : fila["telefonov3"],
            'telv4' : fila["telefonov4"],
            'telv5' : fila["telefonov5"],
            'telv6' : fila["telefonov6"],

            'emailv1' : fila["emailv1"],
            'emailv2' : fila["emailv2"],
            'emailv3' : fila["emailv3"],
            'emailv4' : fila["emailv4"],
            'emailv5' : fila["emailv5"],
            'emailv6' : fila["emailv6"],

            'estadov1' : fila["estadocivilv1"],
            'estadov2' : fila["estadocivilv2"],
            'estadov3' : fila["estadocivilv3"],
            'estadov4' : fila["estadocivilv4"],
            'estadov5' : fila["estadocivilv5"],
            'estadov6' : fila["estadocivilv6"],

            'dirc1' : fila["direc1"],
            'dirc2' : fila["direc2"],
            'dirc3' : fila["direc3"],
            'dirc4' : fila["direc4"],
            'dirc5' : fila["direc5"],
            'dirc6' : fila["direc6"],

            'telec1' : fila["telefonoc1"],
            'telec2' : fila["telefonoc2"],
            'telec3' : fila["telefonoc3"],
            'telec4' : fila["telefonoc4"],
            'telec5' : fila["telefonoc5"],
            'telec6' : fila["telefonoc6"],

            'emailc1' : fila["emailc1"],
            'emailc2' : fila["emailc2"],
            'emailc3' : fila["emailc3"],
            'emailc4' : fila["emailc4"],
            'emailc5' : fila["emailc5"],
            'emailc6' : fila["emailc6"],

            'estadoc1' : fila["estadoc1"],
            'estadoc2' : fila["estadoc2"],
            'estadoc3' : fila["estadoc3"],
            'estadoc4' : fila["estadoc4"],
            'estadoc5' : fila["estadoc5"],
            'estadoc6' : fila["estadoc6"],
            'ejerce' : fila["qejerce"],
            'linderos' : fila["Linderos"],
            'pazsalvo' : fila["pazysalvo"],
            'declaracion' : fila["Declaracion"],
            'vur' : fila["valor_urbano_rural"],
            'npag' : fila["n_paginas"],

            'dep1' : fila["notario_notaria_1"],
            'mun1' : fila["municipioA"],
            'usuario' : fila["username"],
            'numynot' : fila["notaria_actual_1"],

            'DPV' : fila["SingularPluralLosDeudores"],
            'DPC' : fila["SingularPluralLosAcredores"],
            'HPV' : fila["SingularPluralHipotecante"],
            'PPC' : fila["SingularPluralApoderados"],
            'PPV' : fila["SingularPluralPoderdantes"],
            'CPV' : fila["SingularPluralConstituyentes"],
            'FPC' : fila["SingularPluralFidecomisarios"],
            'FIV' : fila["SingularPluralFideicomitentes"],
            'CEPC' : fila["SingularPluralCesionarios"],
            'CEPV' : fila["SingularPluralCedentes"],
            'BPC' : fila["SingularPluralBeneficiarios"],
            'CTE' : fila["SingularPluralConstituyen"],
            'OTE' : fila["SingularPluralOtorgan"],

            'chp1' : fila["chip1"],
            'chpv' : fila["chip_valor"],




            
                 

                 
        } 
    
                   

       
       
        

       
        doc.render(datos)
        doc.save(f"archivos/escritura.docx")
        documento = Document(f"archivos/escritura.docx")

        
        

        

            
        for i, paragraph in enumerate(documento.paragraphs):
            print(f'Índice: {i}, Texto: {paragraph.text}')

        


        


        indices = [33, 52, 54, 55, 56, 57, 60] # índices de los párrafos que desea modificar

        for i, paragraph in enumerate(documento.paragraphs):
            if i in indices:
                lines = paragraph.text.split('\n')
                new_lines = []
                for line in lines:
                    words = line.split()
                    new_line = ' '.join(words)
                    new_lines.append(new_line)

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    

                paragraph.clear()

                new_run = paragraph.add_run('\n'.join(new_lines))
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                




        for paragraph in documento.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)    
                            
        
        documento.save(f"archivos/{ID_registro.get()}.docx")
        messagebox.showinfo(message="Escritura generada exitosamente", title="Aviso")
        ruta = os.path.join(os.getcwd(), 'archivos', f'{ID_registro.get()}.docx')
        os.startfile(ruta) 





ventana1 = Tk()
ventana1.title("Digiplus scriptures")
ventana1.geometry("300x400")
ventana1.iconbitmap("Logo.ico")


image=PhotoImage(file="Logo.gif")


label=Label(image=image)
label.pack()

Label(text="Acceso a Digiplus scriptures", bg="deep sky blue", fg="white", width=300, height=3, font=("calibri", 15)).pack()
Label(text="").pack()

boton=ttk.Button(ventana1, text="Iniciar Sesion", width=30, cursor="hand2", command=inicio_sesion).pack()
Label(text="").pack()
boton2=ttk.Button(ventana1, text="Registrar", width=30, cursor="hand2", command=registrarse).pack()

ventana1.mainloop()
           

      
    
    










        
            

        
    
   

    